# ==============================================================================
# 📄 Arquivo: gerador_word.py
# 📏 Linhas: ~200
# 🎯 Função: Exportação NATIVA para Microsoft Word (.docx)
# 📅 Versão: 2.0 (PRO Elite - Dossiê de Turma por Período Integrado)
# ==============================================================================

import os
import io
import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import get_alunos_por_turma, get_presencas_dia, get_diario_dia, get_midias_diario
from utils.imagem import baixar_imagem_http as baixar_imagem

def adicionar_rodape_padrao(doc):
    """Adiciona o rodapé oficial do Instituto Muda Brasil ao documento Word."""
    section = doc.sections[0]
    footer = section.footer

    if len(footer.paragraphs) == 0:
        p_footer = footer.add_paragraph()
    else:
        p_footer = footer.paragraphs[0]

    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_inst = p_footer.add_run("Instituto Muda Brasil | PROJETO: ESPORTE E SAÚDE NA COMUNIDADE FASE 2\n")
    run_inst.font.size = Pt(8)

    run_cnpj = p_footer.add_run("CNPJ: 08.817.519/0001-79 | Núcleo 2 - Campo Belo; Rua Sapoti, 20, Bairro Jardim Aeroporto - São Paulo\n")
    run_cnpj.font.size = Pt(8)

    run_links = p_footer.add_run("Instagram: @institutomudabrasil | Linktree: https://linktr.ee/institutomudabrasil\n")
    run_links.font.size = Pt(8)
    run_links.font.bold = True
    run_links.font.color.rgb = RGBColor(0, 102, 204) 

    run_motto = p_footer.add_run('"Impacto social. Transformação real. Porque mudar o Brasil começa com mudar histórias!"')
    run_motto.font.size = Pt(8)
    run_motto.font.italic = True


# ==============================================================================
# 1. RELATÓRIO: DOSSIÊ DE UMA AULA ÚNICA (Mantido e Refinado)
# ==============================================================================
def criar_documento_dossie(data_aula, turma):
    doc = Document()

    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Logo
    if os.path.exists('logo-imbra.png'):
        doc.add_picture('logo-imbra.png', width=Inches(0.75))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif os.path.exists('logo-imbra.jpg'):
        doc.add_picture('logo-imbra.jpg', width=Inches(0.75))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_tit.add_run('DOSSIÊ OFICIAL DE AULA')
    run_tit.bold = True
    run_tit.font.size = Pt(14)

    subtitulo = doc.add_paragraph('PROJETO: ESPORTE E SAÚDE NA COMUNIDADE FASE 2')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].bold = True

    p_cabecalho = doc.add_paragraph()
    p_cabecalho.add_run(f'Turma: ').bold = True
    p_cabecalho.add_run(f'{turma}\n')
    p_cabecalho.add_run(f'Data da Aula: ').bold = True
    p_cabecalho.add_run(f'{data_aula.strftime("%d/%m/%Y")}')

    diario = get_diario_dia(data_aula, turma)

    doc.add_heading('1. Objetivo da Sessão', level=1)
    if diario and diario.get('objetivo_geral'):
        doc.add_paragraph(diario['objetivo_geral']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.add_paragraph("Nenhum objetivo técnico registado para esta sessão.")

    doc.add_heading('2. Exercícios Executados', level=1)
    if diario and diario.get('exercicios_executados'):
        doc.add_paragraph(diario['exercicios_executados']).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        doc.add_paragraph("Nenhum detalhe de exercício registado para esta sessão.")

    doc.add_heading('3. Registo de Frequência', level=1)
    df_alunos = get_alunos_por_turma(turma)

    if not df_alunos.empty:
        ids_turma = df_alunos['id'].tolist()
        presencas = get_presencas_dia(data_aula, ids_turma)

        presentes = [row['nome'] for _, row in df_alunos.iterrows() if presencas.get(row['id'], False)]
        ausentes = [row['nome'] for _, row in df_alunos.iterrows() if not presencas.get(row['id'], False)]

        p_pres = doc.add_paragraph()
        p_pres.add_run(f'✔️ Presentes ({len(presentes)}): ').bold = True
        p_pres.add_run("; ".join([f"{i+1}. {n}" for i, n in enumerate(presentes)]) if presentes else "Nenhum aluno presente.")

        p_aus = doc.add_paragraph()
        p_aus.add_run(f'❌ Ausentes ({len(ausentes)}): ').bold = True
        p_aus.add_run("; ".join([f"{i+1}. {n}" for i, n in enumerate(ausentes)]) if ausentes else "Nenhum aluno ausente.")
    else:
        doc.add_paragraph("Nenhum aluno cadastrado nesta turma.")

    doc.add_heading('4. Evidências Fotográficas', level=1)
    if diario:
        todas_fotos = []
        url_grupo = diario.get('url_foto_grupo') or diario.get('foto_grupo') or diario.get('url_foto')
        if url_grupo:
            todas_fotos.append({'url': url_grupo, 'desc': 'Foto Oficial do Grupo'})

        midias = get_midias_diario(diario.get('id'))
        if midias:
            for m in midias:
                img_url = m.get('url_midia') or m.get('url') or m.get('midia_url')
                if img_url:
                    descricao = m.get('descricao_objetivo') or m.get('descricao') or 'Exercício / Atividade'
                    todas_fotos.append({'url': img_url, 'desc': descricao})

        if todas_fotos:
            table = doc.add_table(rows=0, cols=2)
            row_cells = None
            for i, foto in enumerate(todas_fotos):
                if i % 2 == 0: row_cells = table.add_row().cells
                cell = row_cells[i % 2]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_stream = baixar_imagem(foto['url'])
                if img_stream:
                    try:
                        p.add_run().add_picture(img_stream, width=Inches(2.8))
                        cell.add_paragraph(foto['desc']).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except: cell.add_paragraph("[Erro na imagem]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Nenhum Diário de Bordo registado.")

    adicionar_rodape_padrao(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ==============================================================================
# 1B. DOSSIÊ CLÍNICO INDIVIDUAL DO ALUNO — Word (.docx)
# ==============================================================================
def criar_documento_aluno_word(aluno_data, avaliacoes, historico, estatisticas):
    """Gera o Dossiê Clínico e Desportivo do Aluno em formato Word (.docx)."""
    import datetime
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- CABEÇALHO COM LOGOS ---
    from utils.identidade import get_config as _gcfg
    _cfg = _gcfg()
    logo_sec = _cfg.get("logo_secretaria", "logo-secretaria.png")
    logo_pri = _cfg.get("logo_principal", "logo-imbra.png")
    titulo_proj = _cfg.get("titulo_projeto", "ESPORTE E SAÚDE NA COMUNIDADE")

    tbl_hdr = doc.add_table(rows=1, cols=3)
    tbl_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_hdr.columns[0].width = Inches(1.4)
    tbl_hdr.columns[1].width = Inches(4.3)
    tbl_hdr.columns[2].width = Inches(1.4)

    c_esq = tbl_hdr.rows[0].cells[0]
    if os.path.exists(logo_sec):
        try:
            run_le = c_esq.paragraphs[0].add_run()
            run_le.add_picture(logo_sec, width=Inches(1.1))
        except Exception:
            pass

    c_mid = tbl_hdr.rows[0].cells[1]
    p_mid = c_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_mid.add_run(titulo_proj + "\n")
    run_t.bold = True
    run_t.font.size = Pt(13)
    run_t.font.color.rgb = RGBColor(30, 136, 229)
    run_sub = p_mid.add_run("DOSSIÊ CLÍNICO E DESPORTIVO DO ALUNO")
    run_sub.bold = True
    run_sub.font.size = Pt(11)

    c_dir = tbl_hdr.rows[0].cells[2]
    if os.path.exists(logo_pri):
        try:
            p_dir = c_dir.paragraphs[0]
            p_dir.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_ld = p_dir.add_run()
            run_ld.add_picture(logo_pri, width=Inches(1.1))
        except Exception:
            pass

    doc.add_paragraph()

    # --- 1. PERFIL PESSOAL ---
    def _titulo_secao(txt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1E88E5")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        p._p.get_or_add_pPr().append(shading)

    def _campo(label, valor):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r_l = p.add_run(f"{label}: ")
        r_l.bold = True
        r_l.font.size = Pt(11)
        r_v = p.add_run(str(valor) if valor else "Não informado")
        r_v.font.size = Pt(11)

    _titulo_secao(" 1. Perfil Pessoal e Biométrico")

    nasc = aluno_data.get("data_nascimento")
    nasc_str = "Não informado"
    if nasc and str(nasc).strip() not in ("", "nan", "None", "null"):
        try:
            import pandas as _pd
            dt_nasc = _pd.to_datetime(nasc).date()
            idade = datetime.date.today().year - dt_nasc.year
            nasc_str = f"{dt_nasc.strftime('%d/%m/%Y')} ({idade} anos)"
        except Exception:
            nasc_str = str(nasc)

    presencas = estatisticas.get("presentes", 0) if isinstance(estatisticas, dict) else 0

    _campo("Nome", aluno_data.get("nome", ""))
    _campo("Turma", aluno_data.get("turma", ""))
    _campo("Nascimento", nasc_str)
    _campo("Peso", f"{aluno_data.get('peso', 'N/A')} kg")
    _campo("Altura", f"{aluno_data.get('altura', 'N/A')} m")

    p_pres = doc.add_paragraph()
    r_lp = p_pres.add_run("Total de Presenças: ")
    r_lp.bold = True
    r_lp.font.size = Pt(11)
    r_vp = p_pres.add_run(f"{presencas} aulas concluídas")
    r_vp.bold = True
    r_vp.font.size = Pt(12)
    r_vp.font.color.rgb = RGBColor(30, 136, 229)

    # Foto do aluno (opcional — à direita do perfil via anchor não suportado facilmente, colocamos inline)
    url_foto = aluno_data.get("url_foto")
    if url_foto and str(url_foto).strip().lower() not in ("", "nan", "none", "null"):
        img_stream = baixar_imagem(url_foto)
        if img_stream:
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_img = p_img.add_run()
                run_img.add_picture(img_stream, width=Inches(1.2))
            except Exception:
                pass

    doc.add_paragraph()

    # --- 2. HISTÓRICO CLÍNICO ---
    _titulo_secao(" 2. Histórico Clínico e Biofeedback")
    import pandas as _pd2
    if isinstance(avaliacoes, _pd2.DataFrame):
        avaliacoes = avaliacoes.to_dict("records")

    if avaliacoes and len(avaliacoes) > 0:
        for av in avaliacoes:
            dt_av = av.get("data_avaliacao", "Data não informada")
            if isinstance(dt_av, (datetime.date, datetime.datetime)):
                dt_av = dt_av.strftime("%d/%m/%Y")
            dor = av.get("nivel_dor", "N/A")
            tug = av.get("tug", "N/A")
            simetria = av.get("simetria_forca", "N/A")
            p_av = doc.add_paragraph()
            r_dt = p_av.add_run(f"Avaliação de {dt_av}  |  ")
            r_dt.bold = True
            r_dt.font.size = Pt(11)
            p_av.add_run(f"Dor: {dor}/10   Equilíbrio (TUG): {tug}s   Simetria: {simetria}").font.size = Pt(11)
    else:
        doc.add_paragraph().add_run("Nenhuma avaliação clínica registada.").font.size = Pt(11)

    doc.add_paragraph()

    # --- 3. DIÁRIO DE BORDO ---
    _titulo_secao(" 3. Diário de Bordo (Aulas e Exercícios)")
    if isinstance(historico, _pd2.DataFrame):
        historico = historico.to_dict("records")

    if historico and len(historico) > 0:
        for h in historico:
            data_aula = h.get("data_aula", "")
            if isinstance(data_aula, (datetime.date, datetime.datetime)):
                data_aula = data_aula.strftime("%d/%m/%Y")
            p_aula = doc.add_paragraph()
            r_data = p_aula.add_run(f"Aula de {data_aula}:")
            r_data.bold = True
            r_data.font.size = Pt(11)
            r_data.font.color.rgb = RGBColor(30, 136, 229)
            obj = h.get("objetivo_geral")
            exc = h.get("exercicios_executados")
            if obj:
                p_obj = doc.add_paragraph()
                r_lo = p_obj.add_run("Objetivo: ")
                r_lo.bold = True
                r_lo.font.size = Pt(11)
                p_obj.add_run(str(obj)).font.size = Pt(11)
            if exc:
                p_exc = doc.add_paragraph()
                r_le = p_exc.add_run("Exercícios: ")
                r_le.bold = True
                r_le.font.size = Pt(11)
                p_exc.add_run(str(exc)).font.size = Pt(11)
            doc.add_paragraph()
    else:
        doc.add_paragraph().add_run("Nenhuma participação com diário registada.").font.size = Pt(11)

    adicionar_rodape_padrao(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ==============================================================================
# 2. NOVO: DOSSIÊ DA TURMA POR PERÍODO (Diários + Assiduidade)
# ==============================================================================
def criar_dossie_turma_periodo_word(turma, data_inicio, data_fim, diarios, df_estatisticas):
    doc = Document()

    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Logo e Cabeçalho
    if os.path.exists('logo-imbra.png'):
        doc.add_picture('logo-imbra.png', width=Inches(0.75))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit = p_tit.add_run('DOSSIÊ DA TURMA POR PERÍODO\n')
    run_tit.bold = True
    run_tit.font.size = Pt(14)
    run_tit.font.color.rgb = RGBColor(10, 37, 64)

    periodo_str = f"{data_inicio.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}"
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.add_run(f"Turma: {turma}   |   Período: {periodo_str}").bold = True

    doc.add_paragraph() # Espaçamento

    # --- SECÇÃO 1: DIÁRIOS DE BORDO ---
    doc.add_heading('1. Diários de Bordo (Conteúdo Programático)', level=1)

    if diarios and len(diarios) > 0:
        for d in diarios:
            # Tratamento da data
            data_aula = d.get("data_aula")
            if isinstance(data_aula, datetime.date) or isinstance(data_aula, datetime.datetime):
                data_aula = data_aula.strftime("%d/%m/%Y")
            elif isinstance(data_aula, str) and "-" in data_aula:
                try: data_aula = datetime.datetime.strptime(data_aula, "%Y-%m-%d").strftime("%d/%m/%Y")
                except: pass

            p_data = doc.add_paragraph()
            run_data = p_data.add_run(f"▶ Aula de {data_aula}:")
            run_data.bold = True
            run_data.font.color.rgb = RGBColor(30, 136, 229) # Azul MoveRight

            obj = d.get("objetivo_geral", "Não informado.")
            exe = d.get("exercicios_executados", "Não informado.")

            p_obj = doc.add_paragraph()
            p_obj.add_run("Objetivo: ").bold = True
            p_obj.add_run(str(obj))
            p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            p_exe = doc.add_paragraph()
            p_exe.add_run("Exercícios: ").bold = True
            p_exe.add_run(str(exe))
            p_exe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph() # Espaçamento extra entre diários
    else:
        doc.add_paragraph("Nenhum diário registrado neste período.")

    doc.add_page_break()

    # --- SECÇÃO 2: TABELA DE EVASÃO E ASSIDUIDADE ---
    doc.add_heading('2. Relatório de Evasão e Assiduidade', level=1)

    if not df_estatisticas.empty:
        # Criação de uma tabela nativa do Word
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid' # Estilo com bordas

        # Configura o Cabeçalho
        hdr_cells = table.rows[0].cells
        colunas = ['Nome do Aluno', 'Presenças', 'Faltas', 'Total Aulas', 'Taxa (%)']
        for i, nome_col in enumerate(colunas):
            hdr_cells[i].text = nome_col
            # Coloca em negrito
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            if i > 0: hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Preenche os dados
        df_estatisticas = df_estatisticas.sort_values(by="nome")

        for _, row in df_estatisticas.iterrows():
            row_cells = table.add_row().cells

            nome = str(row.get('nome', ''))[:45]
            presencas = int(row.get('presencas', 0))
            faltas = int(row.get('faltas', 0))
            total = presencas + faltas
            taxa = f"{(presencas/total * 100):.1f}%" if total > 0 else "0.0%"

            row_cells[0].text = nome

            row_cells[1].text = str(presencas)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[2].text = str(faltas)
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if faltas > 0:
                for r in row_cells[2].paragraphs[0].runs: r.font.color.rgb = RGBColor(185, 28, 28) # Vermelho

            row_cells[3].text = str(total)
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row_cells[4].text = taxa
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("Sem dados de frequência para esta turma no período.")

    adicionar_rodape_padrao(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()