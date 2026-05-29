# ==============================================================================
# 📄 Arquivo: views/relatorio_view.py
# 🏷️ VERSÃO: 8.40 PRIMEMAX (A MURALHA - PDF Nativo, Excel e Sincronização Anti-Furo)
# 📏 LINHAS: ~900
# 👤 DESENVOLVEDOR: Parceiro de Programação Gemini & Marcos Barbosa
# ⚙️ FUNÇÃO: Relatórios, B.I., Auditoria Interativa (PDF) e Prestação Pedagógica
# ==============================================================================

import streamlit as st
import pandas as pd
import datetime
import plotly.express as px
import io
import os
import re
from views.relatorio_identificacao_view import renderizar_aba_caracracha
from gerador_pdf import (
    criar_prestacao_diaria_pdf,
    criar_prestacao_periodo_pdf,
    criar_pdf_alerta_frequencia,
    gerar_pdf_monitoramento_clinico,
)

from database import (
    get_relatorio_periodo,
    buscar_alunos_geral,
    get_todas_turmas,
    get_alunos_por_turma,
    get_diarios_periodo,
    get_avaliacoes_aluno,
    get_midias_diario,
    get_ultima_presenca_batch,
    get_presentes_dia_todos,
    get_presentes_periodo_todos,
)

# 🚀 IMPORTAÇÃO DO MOTOR NATIVO DO WORD
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

# 🚀 IMPORTAÇÃO DO MOTOR NATIVO DE PDF
try:
    from xhtml2pdf import pisa

    XHTML_DISPONIVEL = True
except ImportError:
    XHTML_DISPONIVEL = False


# --- FUNÇÕES DE APOIO E IDENTIDADE ---
from utils.imagem import get_base64_image


def abrir_ficha_aluno(dados_aluno):
    """Callback executado apenas quando o botão 'Abrir Ficha' é clicado."""
    st.session_state.aluno_prontuario = dados_aluno
    st.session_state.menu_atual = "Portal do Aluno"


# ==============================================================================
# 📊 MOTOR EXCEL: PLANILHA, KPI, B.I. E RODAPÉ PADRÃO MOVERIGHT™
# ==============================================================================
def gerar_excel_planilha_frequencia(
    df_grid,
    turma_nome,
    periodo_str,
    caminho_logo_muda,
    caminho_logo_sec,
    total_alunos,
    total_presencas_geral,
    total_aulas,
):
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_grid.to_excel(
            writer, sheet_name="Prestação de Contas", startrow=12, index=False
        )
        workbook = writer.book
        worksheet = writer.sheets["Prestação de Contas"]

        # --- ESTILOS ---
        f_tit = workbook.add_format(
            {"bold": True, "font_size": 13, "align": "left", "valign": "vcenter"}
        )
        f_sub = workbook.add_format(
            {
                "font_size": 10,
                "align": "left",
                "valign": "vcenter",
                "font_color": "#555555",
            }
        )
        f_dst = workbook.add_format(
            {
                "bold": True,
                "font_size": 10,
                "align": "left",
                "valign": "vcenter",
                "font_color": "#0056b3",
            }
        )
        f_cab_c = workbook.add_format(
            {
                "bold": True,
                "bottom": 1,
                "bg_color": "#F8FAFC",
                "align": "center",
                "valign": "vcenter",
            }
        )
        f_cab_e = workbook.add_format(
            {
                "bold": True,
                "bottom": 1,
                "bg_color": "#F8FAFC",
                "align": "left",
                "valign": "vcenter",
            }
        )
        f_dat_c = workbook.add_format({"align": "center", "valign": "vcenter"})
        f_tot_c = workbook.add_format(
            {
                "bold": True,
                "bottom": 2,
                "bg_color": "#E2E8F0",
                "align": "center",
                "valign": "vcenter",
            }
        )
        f_tot_e = workbook.add_format(
            {
                "bold": True,
                "bottom": 2,
                "bg_color": "#E2E8F0",
                "align": "left",
                "valign": "vcenter",
            }
        )
        f_rodape = workbook.add_format(
            {"font_size": 8, "font_color": "#64748B", "align": "left"}
        )
        f_rodape_bold = workbook.add_format(
            {"font_size": 8, "font_color": "#0A2540", "bold": True, "align": "left"}
        )

        from utils.identidade import (
            get_config as _gcfg_xls,
            get_logo_data_url as _gld_xls,
        )

        _xcfg = _gcfg_xls()
        import os as _os_xls

        _lp = _xcfg.get("logo_principal", "logo-imbra.png")
        _ls = _xcfg.get("logo_secundaria", "logo-secretaria.png")
        try:
            if _os_xls.path.exists(_lp):
                worksheet.insert_image("A1", _lp, {"x_scale": 0.16, "y_scale": 0.16})
        except:
            pass
        try:
            if _os_xls.path.exists(_ls):
                worksheet.insert_image("C1", _ls, {"x_scale": 0.35, "y_scale": 0.35})
        except:
            pass

        worksheet.merge_range(
            "A6:E6",
            f"{_xcfg.get('titulo_projeto', 'PROJETO ESPORTE E SAÚDE NA COMUNIDADE')} - PLANILHA DE FREQUÊNCIA",
            f_tit,
        )
        worksheet.merge_range("A7:E7", f"Período: {periodo_str}", f_sub)
        worksheet.merge_range("A8:E8", f"Escopo: {turma_nome}", f_sub)
        worksheet.merge_range(
            "A9:E9", f"Total de Aulas Realizadas no Período: {total_aulas}", f_dst
        )
        worksheet.merge_range(
            "A10:E10", f"Total de Presenças no Período: {total_presencas_geral}", f_dst
        )
        worksheet.merge_range(
            "A11:E11", f"Total de Alunos no Período: {total_alunos}", f_dst
        )

        def _safe(v):
            """Converte NaN/Inf para string vazia — xlsxwriter não aceita float inválido."""
            if isinstance(v, float) and (v != v or v == float("inf") or v == float("-inf")):
                return ""
            return v

        for col_num, value in enumerate(df_grid.columns.values):
            fmt = f_cab_e if value in ["Aluno", "Turma"] else f_cab_c
            worksheet.write(12, col_num, value, fmt)
            val_topo = _safe(df_grid.iloc[0, col_num])
            fmt_t = f_tot_e if value in ["Aluno", "Turma"] else f_tot_c
            worksheet.write(13, col_num, val_topo, fmt_t)

            for row_num in range(1, len(df_grid)):
                val = _safe(df_grid.iloc[row_num, col_num])
                if value in ["Aluno", "Turma"]:
                    worksheet.write(13 + row_num, col_num, val)
                else:
                    worksheet.write(13 + row_num, col_num, val, f_dat_c)

        worksheet.set_column(0, 0, 8)
        worksheet.set_column(1, 1, 35)
        worksheet.set_column(2, 2, 18)
        if len(df_grid.columns) > 3:
            worksheet.set_column(3, len(df_grid.columns) - 1, 7.5)

        linha_rodape = 13 + len(df_grid) + 2
        worksheet.write(
            linha_rodape,
            0,
            "Sistema Esporte e Saúde - Gestão Inteligente Moveright™",
            f_rodape_bold,
        )
        worksheet.write(
            linha_rodape + 1,
            0,
            f"{_xcfg.get('nome_organizacao', 'Instituto Muda Brasil')} | CNPJ: {_xcfg.get('cnpj', '08.817.519/0001-79')}",
            f_rodape,
        )
        worksheet.write(
            linha_rodape + 2,
            0,
            f"Site: {_xcfg.get('site', 'imbra.org.br')} | Instagram: {_xcfg.get('instagram', '@institutomudabrasil')}",
            f_rodape,
        )
        worksheet.write(
            linha_rodape + 3, 0, f"Endereço: {_xcfg.get('endereco', '')}", f_rodape
        )
        if _xcfg.get("telefone"):
            worksheet.write(linha_rodape + 4, 0, f"Tel: {_xcfg['telefone']}", f_rodape)

        # ======================================================================
        # ABA B.I.: DASHBOARD MILIMÉTRICO
        # ======================================================================
        ws_bi = workbook.add_worksheet("Dashboard B.I.")
        df_al = df_grid.iloc[1:]
        tp = df_al["Total P"].sum() if "Total P" in df_al.columns else 0
        tf = df_al["Total F"].sum() if "Total F" in df_al.columns else 0
        r_cnt = len(df_al[df_al.get("Total F", 0) > df_al.get("Total P", 0)])

        ws_bi.write("A1", "Métrica", f_cab_e)
        ws_bi.write("B1", "Quantidade", f_cab_c)
        ws_bi.write("A2", "Presenças")
        ws_bi.write("B2", tp, f_dat_c)
        ws_bi.write("A3", "Faltas")
        ws_bi.write("B3", tf, f_dat_c)
        ws_bi.write("A5", "Regulares")
        ws_bi.write("B5", total_alunos - r_cnt, f_dat_c)
        ws_bi.write("A6", "Em Risco")
        ws_bi.write("B6", r_cnt, f_dat_c)

        taxa = tp / (tp + tf) if (tp + tf) > 0 else 0
        fmt_st = workbook.add_format(
            {
                "bold": True,
                "align": "center",
                "bg_color": "#D1FAE5" if taxa >= 0.65 else "#FEE2E2",
            }
        )
        ws_bi.merge_range(
            "A9:B9",
            f"{'🟢 SAUDÁVEL' if taxa >= 0.65 else '🔴 ALERTA'} ({taxa:.1%})",
            fmt_st,
        )

        dias_bi = [
            c
            for c in df_grid.columns
            if c
            not in [
                "Ordem",
                "Aluno",
                "Turma",
                "Total P",
                "Total F",
                "Total J",
                "% Presença",
            ]
        ]
        ws_bi.write(10, 0, "Data Aula", f_cab_e)
        ws_bi.write(10, 1, "Presentes", f_cab_c)
        for i, d in enumerate(dias_bi):
            ws_bi.write(11 + i, 0, d)
            ws_bi.write(11 + i, 1, df_grid.iloc[0][d], f_dat_c)

        c1 = workbook.add_chart({"type": "pie"})
        c1.add_series(
            {
                "categories": ["Dashboard B.I.", 1, 0, 2, 0],
                "values": ["Dashboard B.I.", 1, 1, 2, 1],
                "points": [
                    {"fill": {"color": "#10B981"}},
                    {"fill": {"color": "#EF4444"}},
                ],
            }
        )
        c1.set_title({"name": "Assiduidade Global"})
        ws_bi.insert_chart("D1", c1, {"x_scale": 0.6, "y_scale": 0.6})

        c2 = workbook.add_chart({"type": "pie"})
        c2.add_series(
            {
                "categories": ["Dashboard B.I.", 4, 0, 5, 0],
                "values": ["Dashboard B.I.", 4, 1, 5, 1],
                "points": [
                    {"fill": {"color": "#3B82F6"}},
                    {"fill": {"color": "#F59E0B"}},
                ],
            }
        )
        c2.set_title({"name": "Raio-X Turma"})
        ws_bi.insert_chart("H1", c2, {"x_scale": 0.6, "y_scale": 0.6})

        if dias_bi:
            c3 = workbook.add_chart({"type": "line"})
            c3.add_series(
                {
                    "categories": ["Dashboard B.I.", 11, 0, 11 + len(dias_bi) - 1, 0],
                    "values": ["Dashboard B.I.", 11, 1, 11 + len(dias_bi) - 1, 1],
                    "marker": {"type": "circle"},
                    "trendline": {
                        "type": "linear",
                        "line": {"color": "#EF4444", "dash_type": "long_dash"},
                    },
                    "data_labels": {"value": True},
                }
            )
            c3.set_title({"name": "Curva Diária de Frequência (Evolução)"})
            ws_bi.insert_chart("C13", c3, {"x_scale": 1.6, "y_scale": 1.1})

        ws_bi.set_column(0, 0, 25)
    return output.getvalue()


# ==============================================================================
# 🖨️ MOTOR PDF NATIVO: AUDITORIA OFICIAL (NADA DE HTML)
# ==============================================================================
def gerar_pdf_auditoria_core(falhas, contagem_falhas, turma_aud):
    """Gera PDF de auditoria usando fpdf (sem dependência de Pillow/reportlab)."""
    try:
        from fpdf import FPDF
    except ImportError:
        return None

    from utils.identidade import get_config as _gcfg_pdf
    _pcfg = _gcfg_pdf()
    org   = _pcfg.get("nome_organizacao", "INSTITUTO MUDA BRASIL").upper()
    proj  = _pcfg.get("titulo_projeto",   "ESPORTE E SAUDE NA COMUNIDADE")
    hoje  = datetime.date.today().strftime("%d/%m/%Y")

    # helper: limpa acentos que fpdf core não renderiza com Helvetica
    def _a(txt):
        import unicodedata
        return unicodedata.normalize("NFC", str(txt or "")).encode("latin-1", "replace").decode("latin-1")

    _DSEM_PDF = ["seg","ter","qua","qui","sex","sab","dom"]

    def _ult(f):
        raw = f.get("Ultima Presenca") or f.get("\u00daltima Presen\u00e7a") or ""
        v = str(raw).strip()
        if not v or v in ("", "None", "-", "\u2014", "nan"):
            return "Sem registro"
        try:
            _d = datetime.datetime.strptime(v, "%d/%m/%y").date()
            return f"{v} {_DSEM_PDF[_d.weekday()]}"
        except Exception:
            try:
                _d = datetime.datetime.strptime(v, "%d/%m/%Y").date()
                return f"{v} {_DSEM_PDF[_d.weekday()]}"
            except Exception:
                return v

    # ── layout A4 paisagem ─────────────────────────────────────────────────────
    pdf = FPDF(orientation="L", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    pdf.set_margins(14, 10, 14)
    W = pdf.w - 28  # largura útil

    # ── CABEÇALHO ──────────────────────────────────────────────────────────────
    pdf.set_fill_color(10, 37, 64)
    pdf.rect(14, 8, W, 18, "F")
    pdf.set_xy(14, 10)
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(W, 7, _a(org), align="C", ln=True)
    pdf.set_x(14)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(180, 210, 255)
    pdf.cell(W, 5, _a(proj), align="C", ln=True)

    pdf.ln(4)
    pdf.set_text_color(10, 37, 64)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(W, 7, "Relatorio Oficial de Auditoria de Cadastros e Documentos", ln=True)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(100, 116, 139)
    pdf.cell(W, 5, f"Emissao: {hoje}   |   Escopo: {_a(turma_aud)}", ln=True)

    # linha separadora
    pdf.set_draw_color(203, 213, 225)
    pdf.set_line_width(0.4)
    pdf.line(14, pdf.get_y(), 14 + W, pdf.get_y())
    pdf.ln(3)

    # ── RESUMO DE TOTALIZADORES ────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(10, 37, 64)
    pdf.cell(W, 5, "Resumo de Irregularidades:", ln=True)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(51, 65, 85)
    metricas = [(k, v) for k, v in contagem_falhas.items() if v > 0]
    colunas_met = 4
    chunk = [metricas[i:i+colunas_met] for i in range(0, len(metricas), colunas_met)]
    cw = W / colunas_met
    for linha in chunk:
        for label, cnt in linha:
            pdf.cell(cw, 5, _a(f"  {label}: {cnt}"), border=0)
        pdf.ln()
    pdf.ln(2)

    # ── TABELA ─────────────────────────────────────────────────────────────────
    # Colunas: Foto(14) | Nome + Ult.Presenca(85) | Turma(50) | Pendencias(resto)
    C_FOTO  = 14
    C_NOME  = 85
    C_TURMA = 50
    C_PEND  = W - C_FOTO - C_NOME - C_TURMA

    # cabeçalho da tabela
    pdf.set_fill_color(10, 37, 64)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_line_width(0.3)
    pdf.cell(C_FOTO,  6, "Foto",           border=1, fill=True, align="C")
    pdf.cell(C_NOME,  6, "Nome / Ult. Presenca", border=1, fill=True)
    pdf.cell(C_TURMA, 6, "Turma",          border=1, fill=True)
    pdf.cell(C_PEND,  6, "Pendencias Identificadas", border=1, fill=True, ln=True)

    # helpers para download de imagem
    import tempfile, urllib.request, os as _os

    def _baixar_foto(url, timeout=6):
        """Baixa a imagem da URL para um arquivo temporário. Retorna o caminho ou None."""
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                raw = resp.read()
            suffix = ".png" if raw[:8] == b'\x89PNG\r\n\x1a\n' else ".jpg"
            tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
            tmp.write(raw); tmp.close()
            return tmp.name
        except Exception:
            return None

    # linhas de dados
    pdf.set_font("Helvetica", "", 8)
    for idx, f in enumerate(falhas):
        # altura da linha: aumentada para acomodar foto
        LH = 6.5
        LINHA_H = LH * 2   # altura total da linha (13 mm)

        # cor zebrada
        if idx % 2 == 0:
            pdf.set_fill_color(248, 250, 252)
        else:
            pdf.set_fill_color(255, 255, 255)

        x0 = pdf.get_x()
        y0 = pdf.get_y()

        # ── célula FOTO — tenta embutir a imagem real ─────────────────────────
        url_foto = str(f.get("url_foto") or "").strip()
        tem_foto = url_foto.startswith("http")
        foto_embutida = False
        tmp_foto = None

        if tem_foto:
            tmp_foto = _baixar_foto(url_foto)
            if tmp_foto:
                try:
                    # borda da célula
                    pdf.set_draw_color(203, 213, 225)
                    pdf.set_line_width(0.3)
                    pdf.rect(x0, y0, C_FOTO, LINHA_H)
                    # imagem preenche o quadrante inteiro
                    pdf.image(tmp_foto, x=x0 + 0.5, y=y0 + 0.5,
                              w=C_FOTO - 1, h=LINHA_H - 1)
                    foto_embutida = True
                    # avança cursor para início da célula Nome
                    pdf.set_xy(x0 + C_FOTO, y0)
                except Exception:
                    foto_embutida = False
                finally:
                    if tmp_foto and _os.path.exists(tmp_foto):
                        _os.unlink(tmp_foto)

        if not foto_embutida:
            if tem_foto:
                pdf.set_fill_color(220, 252, 231); pdf.set_text_color(22, 101, 52)
                foto_txt = "SIM"
            else:
                pdf.set_fill_color(254, 226, 226); pdf.set_text_color(153, 27, 27)
                foto_txt = "NAO"
            pdf.set_font("Helvetica", "B", 7)
            pdf.cell(C_FOTO, LINHA_H, foto_txt, border=1, fill=True, align="C")

        # restaura fill para zebra
        if idx % 2 == 0:
            pdf.set_fill_color(248, 250, 252)
        else:
            pdf.set_fill_color(255, 255, 255)

        # ── célula NOME + última presença
        pdf.set_text_color(15, 23, 42)
        pdf.set_font("Helvetica", "B", 8)
        nome_x = pdf.get_x()
        nome_y = pdf.get_y()
        nome_txt = _a(f.get("Aluno", ""))
        ult_txt  = _a(_ult(f))
        # cada metade da linha de texto ocupa LH dentro de LINHA_H total
        LH_NOME = LINHA_H / 2
        pdf.cell(C_NOME, LH_NOME, nome_txt, border="LRT", fill=True)
        pdf.set_xy(nome_x, nome_y + LH_NOME)
        pdf.set_font("Helvetica", "I", 7)
        pdf.set_text_color(100, 116, 139)
        pdf.cell(C_NOME, LH_NOME, f"Ult. presenca: {ult_txt}", border="LRB", fill=True)

        # ── célula TURMA
        pdf.set_xy(nome_x + C_NOME, nome_y)
        pdf.set_text_color(15, 23, 42)
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(C_TURMA, LINHA_H, _a(f.get("Turma", "")), border=1, fill=True, align="C")

        # ── célula PENDÊNCIAS (multi_cell com posição manual)
        pdf.set_xy(nome_x + C_NOME + C_TURMA, nome_y)
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(51, 65, 85)
        pend_txt = _a(f.get("Pendências", ""))
        pdf.multi_cell(C_PEND, LH_NOME, pend_txt, border=1, fill=True)

        # garante que próxima linha começa na margem esquerda
        proximo_y = max(pdf.get_y(), nome_y + LINHA_H)
        pdf.set_xy(14, proximo_y)

    # ── RODAPÉ ─────────────────────────────────────────────────────────────────
    pdf.ln(4)
    pdf.set_draw_color(203, 213, 225)
    pdf.line(14, pdf.get_y(), 14 + W, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(W, 5, "Sistema Esporte e Saude - Gestao Inteligente Moveright(tm) - Documento Oficial de Auditoria", align="C")

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ==============================================================================
# 🏆 MOTOR WORD NATIVO: PRESTAÇÃO PEDAGÓGICA (PYTHON-DOCX)
# ==============================================================================
def gerar_word_prestacao_contas(
    turma, mes_nome, ano, engajamento, diarios, clinico, is_global=False
):
    if not DOCX_DISPONIVEL:
        return None
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # CABEÇALHO COM LOGOS LOCAIS
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.6)
    table.columns[1].width = Inches(3.4)
    table.columns[2].width = Inches(1.2)

    c_sec, c_txt, c_imb = table.rows[0].cells
    c_sec.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_imb.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_sec = c_sec.paragraphs[0]
    p_sec.alignment = WD_ALIGN_PARAGRAPH.LEFT
    from utils.identidade import get_config as _gcfg_word

    _wcfg = _gcfg_word()
    _wlogo_s = _wcfg.get("logo_secundaria", "logo-secretaria.png")
    _wlogo_p = _wcfg.get("logo_principal", "logo-imbra.png")
    if os.path.exists(_wlogo_s):
        p_sec.add_run().add_picture(_wlogo_s, width=Inches(1.6))

    p_txt = c_txt.paragraphs[0]
    p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t1 = p_txt.add_run(
        f"{_wcfg.get('nome_organizacao', 'INSTITUTO MUDA BRASIL').upper()}\n"
    )
    run_t1.bold = True
    run_t1.font.size = Pt(14)
    run_t1.font.color.rgb = RGBColor(10, 37, 64)
    run_t2 = p_txt.add_run(
        f"PROJETO: {_wcfg.get('titulo_projeto', 'ESPORTE E SAÚDE NA COMUNIDADE - FASE 2')}"
    )
    run_t2.bold = True
    run_t2.font.size = Pt(10)
    run_t2.font.color.rgb = RGBColor(100, 116, 139)

    p_imb = c_imb.paragraphs[0]
    p_imb.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(_wlogo_p):
        p_imb.add_run().add_picture(_wlogo_p, width=Inches(1.2))

    doc.add_paragraph("_" * 68).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # TÍTULO E ESCOPO
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.space_before = Pt(12)
    rt = p_title.add_run("RELATÓRIO MENSAL DE ATIVIDADES CLÍNICAS E PEDAGÓGICAS\n")
    rt.bold = True
    rt.font.size = Pt(13)
    rt.font.color.rgb = RGBColor(10, 37, 64)

    escopo = "Todas as Turmas do Polo" if is_global else turma
    rsub = p_title.add_run(
        f"Mês de Referência: {mes_nome.upper()} / {ano}  |  Escopo: {escopo}"
    )
    rsub.font.size = Pt(10)
    rsub.font.color.rgb = RGBColor(100, 116, 139)

    # ENGAJAMENTO
    p_e = doc.add_paragraph()
    p_e.space_before = Pt(15)
    re = p_e.add_run("1. MÉTRICAS DE ENGAJAMENTO")
    re.bold = True
    re.font.size = Pt(12)
    re.font.color.rgb = RGBColor(10, 37, 64)

    t_eng = doc.add_table(rows=2, cols=3)
    t_eng.style = "Table Grid"
    t_eng.alignment = WD_TABLE_ALIGNMENT.CENTER
    h1, h2, h3 = t_eng.rows[0].cells
    h1.text, h2.text, h3.text = "Alunos Ativos", "Aulas Lecionadas", "Assiduidade Média"
    v1, v2, v3 = t_eng.rows[1].cells
    v1.text = str(engajamento["total_alunos"])
    v2.text = str(engajamento["total_aulas"])
    v3.text = f"{engajamento['assiduidade']:.1f}%"
    for row in t_eng.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # DIÁRIO DE BORDO
    p_d = doc.add_paragraph()
    p_d.space_before = Pt(15)
    rd = p_d.add_run("2. RESUMO PEDAGÓGICO (Diário de Bordo)")
    rd.bold = True
    rd.font.size = Pt(12)
    rd.font.color.rgb = RGBColor(10, 37, 64)

    if diarios:
        for d in diarios[:20]:
            try:
                dt_obj = pd.to_datetime(str(d["data_aula"])).date()
                data_fmt = dt_obj.strftime("%d/%m/%Y")
            except:
                data_fmt = str(d["data_aula"])
            prefixo_turma = f"[{d.get('turma', '')[:5]}] " if is_global else ""
            p_aula = doc.add_paragraph()
            p_aula.style = "List Bullet"
            r_bold = p_aula.add_run(f"Aula de {data_fmt} {prefixo_turma}: ")
            r_bold.bold = True
            p_aula.add_run(f"{d.get('objetivo_geral', 'Sem objetivo definido.')}")
    else:
        doc.add_paragraph(
            "Nenhum registo de aula preenchido no sistema durante este período."
        )

    # BIOINDICADORES
    p_c = doc.add_paragraph()
    p_c.space_before = Pt(15)
    rc = p_c.add_run("3. IMPACTO CLÍNICO E SAÚDE")
    rc.bold = True
    rc.font.size = Pt(12)
    rc.font.color.rgb = RGBColor(10, 37, 64)

    doc.add_paragraph(
        f"Durante o mês de {mes_nome}, foram realizadas reavaliações clínicas. Com base nos dados consolidados, a média geral foi:"
    )
    b1 = doc.add_paragraph(style="List Bullet")
    b1.add_run("Esforço Médio Global (Borg): ")
    b1.add_run(f"{clinico['borg']}").bold = True
    b2 = doc.add_paragraph(style="List Bullet")
    b2.add_run("Nível de Dor Médio Reportado: ")
    b2.add_run(f"{clinico['dor']:.1f} / 10").bold = True
    b3 = doc.add_paragraph(style="List Bullet")
    b3.add_run("Padrão de Saúde Intestinal (Bristol): ")
    b3.add_run(f"{clinico['bristol']}").bold = True
    b4 = doc.add_paragraph(style="List Bullet")
    b4.add_run("Nível de Hidratação (Urina): ")
    b4.add_run(f"{clinico['urina']}").bold = True

    # ASSINATURA OFICIAL CORRIGIDA
    doc.add_paragraph("\n")
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs1 = p_sig.add_run("SISTEMA\nESPORTE E SAÚDE\n\n")
    rs1.bold = True
    rs1.font.size = Pt(11)
    rs1.font.color.rgb = RGBColor(10, 37, 64)
    rs2 = p_sig.add_run("Coordenador")
    rs2.font.size = Pt(10)
    rs2.font.color.rgb = RGBColor(10, 37, 64)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ==============================================================================
# MOTOR PDF: RELATÓRIO PRIME DE FREQUÊNCIA
# Gráficos em HTML/CSS puro — sem kaleido, sem matplotlib, 100% xhtml2pdf.
# REGRAS: sem emojis no HTML, sem display:table-cell, só <table> para layout.
# ==============================================================================
def _html_barras_rel(dados, titulo, cor="#0056b3", max_val=None):
    """Gráfico de barras horizontais em HTML puro — compatível xhtml2pdf."""
    if not dados:
        return ""
    try:
        _max = float(max_val) if max_val else max(
            float(str(v).replace(",", ".")) for _, v in dados
        ) or 1
    except Exception:
        _max = 1
    linhas = ""
    for label, val in dados:
        try:
            v = float(str(val).replace(",", "."))
        except Exception:
            v = 0
        pct = int(v / _max * 82) if _max > 0 else 0
        cor_b = "#10B981" if v >= _max * 0.75 else ("#F59E0B" if v >= _max * 0.5 else cor)
        linhas += (
            "<tr>"
            f"<td style='width:22%;font-size:7.5pt;text-align:right;padding:2px 5px 2px 0;"
            f"white-space:nowrap;color:#475569;'>{label}</td>"
            f"<td style='width:66%;padding:2px 0;vertical-align:middle;'>"
            f"<div style='background:{cor_b};height:11px;width:{pct}%;'></div></td>"
            f"<td style='width:12%;font-size:8pt;font-weight:700;color:#0A2540;"
            f"padding-left:4px;'>{val}</td>"
            "</tr>"
        )
    return (
        f"<p style='margin:10px 0 3px;font-size:9pt;font-weight:700;color:#0056b3;'>{titulo}</p>"
        f"<table style='width:100%;border-collapse:collapse;margin-bottom:6px;'>{linhas}</table>"
    )


def _gerar_pdf_relatorio_prime(
    df_mat, cols_datas, periodo_str, turma_str,
    n_alunos, n_aulas, tp_geral,
    media_p_qtd, taxa_media,
    n_exc, n_reg, n_at, n_crit, n_risco,
    d_i, d_f,
):
    """
    Gera PDF Prime completo com gráficos, resumo por turma, rankings e
    Planilha Detalhada (P/F/J colorida) em página landscape.
    Compatível com xhtml2pdf — sem emojis no HTML.
    """
    try:
        from xhtml2pdf import pisa
        from utils.identidade import get_config as _cfg, get_logo_data_url as _gld

        cfg      = _cfg()
        titulo   = cfg.get("titulo_projeto", "")
        subtit   = cfg.get("subtitulo_projeto", "")
        nome_org = cfg.get("nome_organizacao", "")
        cnpj     = cfg.get("cnpj", "")
        site     = cfg.get("site", "")
        insta    = cfg.get("instagram", "")
        endereco = cfg.get("endereco", "")

        logo_p = _gld(cfg.get("logo_principal", ""))
        logo_s = _gld(cfg.get("logo_secundaria", ""))
        img_p  = (f'<img src="data:image/png;base64,{logo_p}" style="max-width:90px;max-height:55px;" />'
                  if logo_p else f"<b>{nome_org}</b>")
        img_s  = (f'<img src="data:image/png;base64,{logo_s}" style="max-width:110px;max-height:55px;" />'
                  if logo_s else "")

        agora        = datetime.datetime.now().strftime("%d/%m/%Y as %H:%M")
        rodape_linha = " | ".join(p for p in [nome_org, f"CNPJ: {cnpj}" if cnpj else "",
                                               site, insta, endereco] if p)

        _DIAS_PT_PDF = {
            "Monday": "Segunda", "Tuesday": "Terca", "Wednesday": "Quarta",
            "Thursday": "Quinta", "Friday": "Sexta", "Saturday": "Sabado", "Sunday": "Domingo",
        }
        _ORDEM_SEM_PDF = ["Segunda", "Terca", "Quarta", "Quinta", "Sexta", "Sabado", "Domingo"]

        def _parse_dt(s, ini, fim):
            for yr in sorted({ini.year, fim.year}):
                try:
                    dt = datetime.datetime.strptime(f"{s}/{yr}", "%d/%m/%Y").date()
                    if ini <= dt <= fim:
                        return dt
                except Exception:
                    pass
            return None

        # ── Gráfico: presença diária ─────────────────────────────────────────
        tl_dados = []
        wd_acc   = {}
        _dt_map  = {}
        for col in cols_datas:
            pres = int((df_mat[col] == "P").sum())
            dt   = _parse_dt(col, d_i, d_f)
            wd   = _DIAS_PT_PDF.get(dt.strftime("%A"), "?") if dt else "?"
            label = f"{col} {wd[:3]}"
            tl_dados.append((label, pres))
            _dt_map[label] = dt or datetime.date.max
            if wd != "?":
                acc = wd_acc.setdefault(wd, [0, 0])
                acc[0] += pres
                acc[1] += 1
        tl_dados.sort(key=lambda x: _dt_map.get(x[0], datetime.date.max))
        chart_diario = _html_barras_rel(tl_dados, "Presencas Diarias no Periodo", cor="#0056b3")

        # ── Gráfico: média por dia da semana ─────────────────────────────────
        sem_dados = [
            (d, f"{wd_acc[d][0]/wd_acc[d][1]:.1f}")
            for d in _ORDEM_SEM_PDF if d in wd_acc and wd_acc[d][1] > 0
        ]
        max_sem = max((float(v) for _, v in sem_dados), default=1) if sem_dados else 1
        chart_semana = _html_barras_rel(
            sem_dados, "Padrao de Comparecimento por Dia da Semana",
            cor="#0056b3", max_val=max_sem,
        )

        # ── Resumo por turma + gráfico taxa ─────────────────────────────────
        linhas_turma = ""
        taxa_turma_dados = []
        if "Turma" in df_mat.columns and "Total P" in df_mat.columns:
            _col_aluno = "Aluno" if "Aluno" in df_mat.columns else "Nome"
            grp = (
                df_mat.groupby("Turma")
                .agg(Alunos=(_col_aluno, "count"), Aulas=("Total Aulas", "first"),
                     Presencas=("Total P", "sum"))
                .reset_index()
            )
            grp["Taxa"]  = (grp["Presencas"] / (grp["Alunos"] * grp["Aulas"]).replace(0, 1) * 100).round(1)
            grp["Media"] = (grp["Presencas"] / grp["Alunos"].replace(0, 1)).round(1)
            for _, r in grp.iterrows():
                taxa_turma_dados.append((str(r["Turma"]), f"{r['Taxa']:.1f}%"))
                linhas_turma += (
                    f"<tr><td>{r['Turma']}</td>"
                    f"<td style='text-align:center;'>{r['Alunos']}</td>"
                    f"<td style='text-align:center;'>{r['Aulas']}</td>"
                    f"<td style='text-align:center;font-weight:700;'>{r['Presencas']}</td>"
                    f"<td style='text-align:center;'>{r['Media']:.1f}</td>"
                    f"<td style='text-align:center;font-weight:700;color:#0056b3;'>{r['Taxa']:.1f}%</td>"
                    f"</tr>"
                )

        chart_taxa_turma = _html_barras_rel(
            taxa_turma_dados, "Taxa de Presenca por Turma (%)",
            cor="#0056b3",
            max_val=max((float(str(v).replace("%","")) for _, v in taxa_turma_dados), default=1) if taxa_turma_dados else 1,
        ) if taxa_turma_dados else ""

        # ── Gráfico: distribuição de assiduidade ────────────────────────────
        dist_dados = [
            ("Excelente (90%+)", n_exc),
            ("Regular (75-90%)", n_reg),
            ("Atencao (50-75%)", n_at),
            ("Critico (<50%)",   n_crit),
        ]
        chart_dist = _html_barras_rel(
            dist_dados, "Distribuicao de Assiduidade — Alunos por Faixa",
            cor="#0056b3",
            max_val=max(n_exc, n_reg, n_at, n_crit, 1),
        )

        # ── Rankings ────────────────────────────────────────────────────────
        def _perc_pdf(s):
            try:
                return float(str(s).replace("%", "").strip())
            except Exception:
                return 0.0

        col_perc = "% Presença" if "% Presença" in df_mat.columns else "% Presenca"
        _col_aluno = "Aluno" if "Aluno" in df_mat.columns else "Nome"
        df_r = df_mat[[_col_aluno, "Turma", "Total P", col_perc]].copy() if col_perc in df_mat.columns else df_mat[[_col_aluno, "Turma", "Total P"]].copy()
        df_r["_tx"] = df_r[col_perc].apply(_perc_pdf) if col_perc in df_r.columns else df_mat.get("_taxa_num", pd.Series(dtype=float))
        top10 = df_r.nlargest(10, "_tx")
        atenc = df_r[df_r["_tx"] < 75].nsmallest(20, "_tx")

        linhas_top = ""
        for i, (_, r) in enumerate(top10.iterrows(), 1):
            linhas_top += (
                f"<tr><td style='text-align:center;'>{i}</td>"
                f"<td>{r[_col_aluno]}</td><td>{r['Turma']}</td>"
                f"<td style='text-align:center;'>{int(r['Total P'])}</td>"
                f"<td style='text-align:center;font-weight:700;color:#10B981;'>{r['_tx']:.1f}%</td>"
                f"</tr>"
            )
        linhas_at = ""
        for _, r in atenc.iterrows():
            cor_tx = "#EF4444" if r["_tx"] < 50 else "#F59E0B"
            linhas_at += (
                f"<tr><td>{r[_col_aluno]}</td><td>{r['Turma']}</td>"
                f"<td style='text-align:center;'>{int(r['Total P'])}</td>"
                f"<td style='text-align:center;font-weight:700;color:{cor_tx};'>{r['_tx']:.1f}%</td>"
                f"</tr>"
            )

        sec_turma_taxa = ""
        if linhas_turma:
            sec_turma_taxa = f"""
<table style="border-collapse:collapse;margin-bottom:0;">
  <tr>
    <td style="width:52%;vertical-align:top;border-bottom:0;padding-right:8px;">
      <p style="margin:8px 0 3px;font-size:9pt;font-weight:700;color:#0056b3;border-bottom:2px solid #0056b3;padding-bottom:2px;">Resumo por Turma</p>
      <table>
        <tr><th>Turma</th><th>Alunos</th><th>Aulas</th><th>Presencas</th><th>Media</th><th>Taxa</th></tr>
        {linhas_turma}
      </table>
    </td>
    <td style="width:48%;vertical-align:top;border-bottom:0;padding-left:8px;">
      {chart_taxa_turma}
    </td>
  </tr>
</table>"""

        sec_rank = f"""
<table style="border-collapse:collapse;margin-bottom:0;">
  <tr>
    <td style="width:50%;vertical-align:top;border-bottom:0;padding-right:8px;">
      <p style="margin:8px 0 3px;font-size:9pt;font-weight:700;color:#10B981;border-bottom:2px solid #10B981;padding-bottom:2px;">Top 10 Mais Assíduos</p>
      <table>
        <tr><th>#</th><th>Aluno</th><th>Turma</th><th>Pres.</th><th>Taxa</th></tr>
        {linhas_top}
      </table>
    </td>
    <td style="width:50%;vertical-align:top;border-bottom:0;padding-left:8px;">
      <p style="margin:8px 0 3px;font-size:9pt;font-weight:700;color:#F59E0B;border-bottom:2px solid #F59E0B;padding-bottom:2px;">Requerem Atencao (abaixo 75%)</p>
      {"<table><tr><th>Aluno</th><th>Turma</th><th>Pres.</th><th>Taxa</th></tr>" + linhas_at + "</table>"
        if linhas_at else "<p style='color:#94A3B8;font-size:9pt;margin-top:8px;'>Todos os alunos acima de 75%.</p>"}
    </td>
  </tr>
</table>"""

        # ── Gráficos diário + semana lado a lado ────────────────────────────
        sec_graficos = f"""
<table style="border-collapse:collapse;margin-bottom:0;">
  <tr>
    <td style="width:60%;vertical-align:top;border-bottom:0;padding-right:8px;">{chart_diario}</td>
    <td style="width:40%;vertical-align:top;border-bottom:0;padding-left:8px;">{chart_semana}</td>
  </tr>
</table>"""

        # ══════════════════════════════════════════════════════════════════════
        # PLANILHA DETALHADA — Página 2 (Landscape)
        # ══════════════════════════════════════════════════════════════════════
        _col_aluno = "Aluno" if "Aluno" in df_mat.columns else "Nome"
        col_perc   = "% Presença" if "% Presença" in df_mat.columns else "% Presenca"

        n_datas = len(cols_datas)
        # Calcula % de largura para cada coluna de data
        # Fixas: Ordem(3%), Aluno(17%), Turma(9%), %Pres(5%), TotP(3.5%), TotAulas(3.5%) = 41%
        pct_data = max(1.0, round(59.0 / n_datas, 2)) if n_datas > 0 else 2.0

        # Cabeçalho das datas (encurtado para caber)
        cabecalho_datas = "".join(
            f"<th style='text-align:center;font-size:5pt;padding:1px;width:{pct_data}%;'>"
            f"{c}</th>"
            for c in cols_datas
        )

        # Linha de total por dia
        total_dia_cells = ""
        for col in cols_datas:
            cnt = int((df_mat[col] == "P").sum())
            total_dia_cells += (
                f"<td style='text-align:center;font-size:5.5pt;padding:1px;"
                f"background:#1D4ED8;color:#fff;font-weight:700;'>{cnt if cnt > 0 else '-'}</td>"
            )

        tj_total = int(df_mat["Total J"].sum()) if "Total J" in df_mat.columns else 0
        tf_total = int(df_mat["Total F"].sum()) if "Total F" in df_mat.columns else 0

        # Linhas dos alunos
        df_plan = df_mat.sort_values(_col_aluno).reset_index(drop=True)
        linhas_plan = ""
        for idx, (_, row) in enumerate(df_plan.iterrows()):
            bg_row = "#FFFFFF" if idx % 2 == 0 else "#F8FAFC"
            nome_aluno = str(row.get(_col_aluno, ""))[:32]
            turma_aluno = str(row.get("Turma", ""))[:18]
            perc_str  = str(row.get(col_perc, "-"))
            tot_p_val = str(row.get("Total P", "-"))
            tot_a_val = str(row.get("Total Aulas", "-"))

            cells_datas = ""
            for col in cols_datas:
                val = str(row.get(col, "")).strip()
                if val in ("P", "p"):
                    bg_c = "#DCFCE7"; fg_c = "#15803D"; txt = "P"
                elif val in ("F", "f"):
                    bg_c = "#FEE2E2"; fg_c = "#DC2626"; txt = "F"
                elif val in ("J", "j", "T", "t"):
                    bg_c = "#FEF9C3"; fg_c = "#B45309"; txt = "J"
                else:
                    bg_c = bg_row; fg_c = "#CBD5E1"; txt = "-"
                cells_datas += (
                    f"<td style='text-align:center;font-size:5.5pt;padding:1px;"
                    f"background:{bg_c};color:{fg_c};font-weight:700;'>{txt}</td>"
                )

            cor_taxa = "#10B981" if _perc_pdf(perc_str) >= 75 else ("#F59E0B" if _perc_pdf(perc_str) >= 50 else "#EF4444")
            linhas_plan += f"""
<tr style="background:{bg_row};">
  <td style="text-align:center;font-size:6pt;padding:1px 2px;">{idx+1}</td>
  <td style="font-size:6pt;padding:1px 3px;font-weight:700;">{nome_aluno}</td>
  <td style="font-size:5.5pt;padding:1px 2px;color:#475569;">{turma_aluno}</td>
  <td style="text-align:center;font-size:6pt;padding:1px 2px;font-weight:700;color:{cor_taxa};">{perc_str}</td>
  <td style="text-align:center;font-size:6pt;padding:1px 2px;font-weight:700;color:#0056b3;">{tot_p_val}</td>
  <td style="text-align:center;font-size:6pt;padding:1px 2px;color:#475569;">{tot_a_val}</td>
  {cells_datas}
</tr>"""

        tabela_planilha = f"""
<table style="border-collapse:collapse;width:100%;table-layout:fixed;">
  <colgroup>
    <col style="width:3%;"/>
    <col style="width:17%;"/>
    <col style="width:9%;"/>
    <col style="width:5%;"/>
    <col style="width:3.5%;"/>
    <col style="width:3.5%;"/>
  </colgroup>
  <thead>
    <tr style="background:#0056b3;color:#fff;">
      <th style="font-size:6pt;padding:2px 2px;text-align:center;">No.</th>
      <th style="font-size:6pt;padding:2px 3px;text-align:left;">Aluno</th>
      <th style="font-size:6pt;padding:2px 2px;">Turma</th>
      <th style="font-size:6pt;padding:2px 2px;text-align:center;">% Pres</th>
      <th style="font-size:6pt;padding:2px 2px;text-align:center;">Tot P</th>
      <th style="font-size:6pt;padding:2px 2px;text-align:center;">Aulas</th>
      {cabecalho_datas}
    </tr>
    <tr style="background:#1D4ED8;color:#fff;">
      <td style="font-size:5.5pt;padding:1px 2px;text-align:center;">-</td>
      <td style="font-size:6pt;padding:1px 3px;font-weight:700;" colspan="2">TOTAL PRESENCAS / DIA</td>
      <td style="font-size:5.5pt;padding:1px 2px;text-align:center;">-</td>
      <td style="font-size:6pt;padding:1px 2px;text-align:center;font-weight:700;">{tp_geral}</td>
      <td style="font-size:5.5pt;padding:1px 2px;text-align:center;">-</td>
      {total_dia_cells}
    </tr>
  </thead>
  <tbody>{linhas_plan}</tbody>
</table>"""

        # ── HTML Final ──────────────────────────────────────────────────────
        html = f"""<!DOCTYPE html>
<html><head>
<meta charset="UTF-8"/>
<style>
  @page portrait_page {{
    size: 210mm 297mm;
    margin: 14mm 18mm 14mm 18mm;
  }}
  @page landscape_page {{
    size: 297mm 210mm;
    margin: 8mm 10mm 8mm 10mm;
  }}
  body {{ font-family: Arial, sans-serif; font-size: 9pt; color: #1e293b; }}
  table {{ width: 100%; border-collapse: collapse; margin-bottom: 8px; }}
  th {{ background: #0056b3; color: #fff; padding: 4px 6px; text-align: left; font-size: 7.5pt; }}
  td {{ padding: 3px 6px; border-bottom: 1px solid #E2E8F0; font-size: 7.5pt; vertical-align: middle; }}
  tr:nth-child(even) td {{ background: #F8FAFC; }}
  .kv {{ font-size: 15pt; font-weight: 900; color: #0056b3; line-height: 1.1; }}
  .kl {{ font-size: 6.5pt; color: #475569; font-weight: 700; text-transform: uppercase; letter-spacing:.3px; }}
  .ks {{ font-size: 6pt; color: #94A3B8; }}
  .rod {{ margin-top: 14px; text-align: center; font-size: 6.5pt; color: #94A3B8;
          border-top: 1px solid #E2E8F0; padding-top: 5px; }}
  .cabec-logo td {{ border-bottom: 0 !important; padding: 4px 6px; }}
</style>
</head>
<body>

<!-- ═══════════ PÁGINA 1 — ANÁLISE PRIME (Portrait) ═══════════ -->

<table class="cabec-logo" style="border-bottom:3px solid #0056b3;margin-bottom:10px;">
  <tr>
    <td style="width:20%;text-align:left;">{img_s}</td>
    <td style="width:60%;text-align:center;">
      <p style="margin:0;font-size:9pt;font-weight:900;color:#0A2540;">{titulo}</p>
      <p style="margin:1px 0;font-size:8pt;color:#475569;">{subtit}</p>
      <p style="margin:0;font-size:9pt;font-weight:700;color:#0056b3;">Relatorio Prime — Analise de Frequencia</p>
      <p style="margin:1px 0 0;font-size:7.5pt;color:#64748B;">Periodo: {periodo_str} &nbsp;|&nbsp; Turma: {turma_str}</p>
      <p style="margin:1px 0 0;font-size:7pt;color:#94A3B8;">Gerado em: {agora}</p>
    </td>
    <td style="width:20%;text-align:right;">{img_p}</td>
  </tr>
</table>

<p style="margin:0 0 4px;font-size:8.5pt;font-weight:700;color:#0056b3;border-bottom:2px solid #0056b3;padding-bottom:2px;">Indicadores do Periodo</p>
<table style="margin-bottom:8px;">
  <tr>
    <td style="width:20%;text-align:center;background:#EFF6FF;border:1.5px solid #BFDBFE;padding:5px;">
      <div class="kv">{n_alunos}</div><div class="kl">Alunos</div>
    </td>
    <td style="width:20%;text-align:center;background:#EFF6FF;border:1.5px solid #BFDBFE;padding:5px;">
      <div class="kv">{n_aulas}</div><div class="kl">Aulas Realizadas</div>
    </td>
    <td style="width:20%;text-align:center;background:#EFF6FF;border:1.5px solid #BFDBFE;padding:5px;">
      <div class="kv">{tp_geral}</div><div class="kl">Total Presencas</div>
    </td>
    <td style="width:20%;text-align:center;background:#EFF6FF;border:1.5px solid #BFDBFE;padding:5px;">
      <div class="kv">{media_p_qtd:.1f}</div>
      <div class="kl">Media P / Aluno</div>
      <div class="ks">de {n_aulas} aulas</div>
    </td>
    <td style="width:20%;text-align:center;background:#EFF6FF;border:1.5px solid #BFDBFE;padding:5px;">
      <div class="kv">{taxa_media:.1f}%</div>
      <div class="kl">Taxa Media Individual</div>
      <div class="ks">meta: 75%</div>
    </td>
  </tr>
</table>

<table style="margin-bottom:10px;">
  <tr>
    <td style="width:25%;text-align:center;background:#DCFCE7;border:1px solid #BBF7D0;padding:4px;">
      <div class="kv" style="color:#10B981;">{n_exc}</div><div class="kl">Excelente (90%+)</div>
    </td>
    <td style="width:25%;text-align:center;background:#EFF6FF;border:1px solid #BFDBFE;padding:4px;">
      <div class="kv" style="color:#3B82F6;">{n_reg}</div><div class="kl">Regular (75-90%)</div>
    </td>
    <td style="width:25%;text-align:center;background:#FEF9C3;border:1px solid #FDE68A;padding:4px;">
      <div class="kv" style="color:#F59E0B;">{n_at}</div><div class="kl">Atencao (50-75%)</div>
    </td>
    <td style="width:25%;text-align:center;background:#FEE2E2;border:1px solid #FECACA;padding:4px;">
      <div class="kv" style="color:#EF4444;">{n_crit}</div><div class="kl">Critico (&lt;50%)</div>
    </td>
  </tr>
</table>

{sec_graficos}

{sec_turma_taxa}

{chart_dist}

{sec_rank}

<div class="rod">{rodape_linha}</div>

<!-- ═══════════ PÁGINA 2 — PLANILHA DETALHADA (Landscape) ═══════════ -->
<pdf:nexttemplate name="landscape_page" />
<pdf:nextpage />

<table class="cabec-logo" style="border-bottom:2px solid #0056b3;margin-bottom:6px;">
  <tr>
    <td style="width:18%;text-align:left;">{img_s}</td>
    <td style="width:64%;text-align:center;">
      <p style="margin:0;font-size:8.5pt;font-weight:900;color:#0A2540;">{titulo}</p>
      <p style="margin:1px 0;font-size:7.5pt;font-weight:700;color:#0056b3;">Planilha Detalhada — Frequencia Individual</p>
      <p style="margin:0;font-size:7pt;color:#64748B;">Periodo: {periodo_str} &nbsp;|&nbsp; Turma: {turma_str} &nbsp;|&nbsp; {n_alunos} alunos &nbsp;|&nbsp; {n_aulas} aulas</p>
      <p style="margin:1px 0 0;font-size:6.5pt;color:#94A3B8;">P = Presente (verde) &nbsp;|&nbsp; F = Falta (vermelho) &nbsp;|&nbsp; J = Justificada (amarelo)</p>
    </td>
    <td style="width:18%;text-align:right;">{img_p}</td>
  </tr>
</table>

{tabela_planilha}

<div class="rod">{rodape_linha} | Gerado em: {agora}</div>

</body></html>"""

        buf = io.BytesIO()
        result = pisa.CreatePDF(html.encode("utf-8"), dest=buf)
        if result.err:
            st.error(f"Erro ao renderizar PDF ({result.err}). Verifique o conteudo e tente novamente.")
            return None
        return buf.getvalue() or None
    except Exception as e:
        st.error(f"Erro ao gerar PDF do Relatorio Prime: {e}")
        import traceback
        st.caption(traceback.format_exc())
        return None


# ==============================================================================
# 📋 PRESTAÇÃO DIÁRIA — Lista de Presença por Período (PDF)
# ==============================================================================

# ── Feriados nacionais brasileiros ────────────────────────────────────────────
def _pascoa(ano: int) -> datetime.date:
    """Algoritmo anônimo gregoriano para calcular a Páscoa."""
    a = ano % 19
    b, c = divmod(ano, 100)
    d, e = divmod(b, 4)
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i, k = divmod(c, 4)
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    mes = (h + l - 7 * m + 114) // 31
    dia = (h + l - 7 * m + 114) % 31 + 1
    return datetime.date(ano, mes, dia)


def _feriados_nacionais_br(anos: set) -> set:
    """
    Retorna set de datetime.date com feriados nacionais brasileiros
    para todos os anos solicitados. Inclui fixos + móveis (Páscoa-based).
    """
    feriados = set()
    FIXOS = [(1, 1), (4, 21), (5, 1), (9, 7), (10, 12), (11, 2), (11, 15), (12, 25)]
    for ano in anos:
        for mes, dia in FIXOS:
            try:
                feriados.add(datetime.date(ano, mes, dia))
            except ValueError:
                pass
        pascoa = _pascoa(ano)
        feriados.add(pascoa - datetime.timedelta(days=48))  # Carnaval 2ª-feira
        feriados.add(pascoa - datetime.timedelta(days=47))  # Carnaval 3ª-feira
        feriados.add(pascoa - datetime.timedelta(days=2))   # Sexta-feira Santa
        feriados.add(pascoa + datetime.timedelta(days=60))  # Corpus Christi
    return feriados


def _dia_da_semana_pt(d: datetime.date) -> str:
    nomes = ["Segunda-feira", "Terça-feira", "Quarta-feira",
             "Quinta-feira", "Sexta-feira", "Sábado", "Domingo"]
    return nomes[d.weekday()]


def _bloco_preview_dia(data_fmt, nomes, img_s, img_p, nome_org, titulo_proj, hoje_fmt, numero_dia=None, total_dias=None):
    """Retorna HTML de um bloco de preview para um dia de presença."""
    linhas = ""
    for i, nome in enumerate(nomes, 1):
        bg = "#F8FAFC" if i % 2 == 0 else "#FFFFFF"
        linhas += (
            f"<tr style='background:{bg};'>"
            f"<td style='padding:4px 8px;text-align:center;width:42px;"
            f"font-weight:700;color:#475569;font-size:11px;border:1px solid #E2E8F0;'>{i}</td>"
            f"<td style='padding:4px 10px;font-size:12px;font-weight:600;"
            f"color:#0F172A;text-transform:uppercase;border:1px solid #E2E8F0;'>{nome}</td>"
            f"</tr>"
        )
    pag_info = f" &nbsp;·&nbsp; Página {numero_dia} de {total_dias}" if numero_dia else ""
    return f"""
<div style='background:#fff;border:1px solid #CBD5E1;border-radius:8px;
            box-shadow:0 4px 18px rgba(0,0,0,.08);padding:24px 28px;
            max-width:720px;margin:0 auto 28px auto;font-family:Arial,sans-serif;'>
  <div style='display:grid;grid-template-columns:76px 1fr 76px;
              align-items:center;gap:10px;
              border-bottom:3px solid #1E3A5F;padding-bottom:10px;margin-bottom:12px;'>
    <div style='text-align:center;'>{img_s}</div>
    <div style='text-align:center;'>
      <div style='font-size:10.5px;font-weight:900;color:#0A2540;
                  text-transform:uppercase;line-height:1.35;'>{nome_org}<br>{titulo_proj}</div>
      <div style='font-size:8.5px;color:#64748B;font-weight:600;margin-top:2px;'>
        PLANILHA DE FREQUÊNCIA DIÁRIA
      </div>
    </div>
    <div style='text-align:center;'>{img_p}</div>
  </div>
  <div style='background:#1E3A5F;color:#fff;padding:6px 10px;
              font-weight:900;font-size:12px;margin-bottom:8px;'>
    Lista de Presença — {data_fmt}
  </div>
  <div style='font-size:10px;color:#1E3A5F;font-weight:700;margin-bottom:7px;'>
    Total de alunos presentes: <span style='background:#1E88E5;color:#fff;
    padding:1px 8px;border-radius:10px;'>{len(nomes)}</span>
  </div>
  <table style='width:100%;border-collapse:collapse;'>
    <thead>
      <tr style='background:#1E3A5F;'>
        <th style='padding:5px 8px;text-align:center;font-size:9.5px;font-weight:800;
                   color:#fff;width:42px;border:1px solid #1E3A5F;'>#</th>
        <th style='padding:5px 10px;text-align:left;font-size:9.5px;font-weight:800;
                   color:#fff;text-transform:uppercase;border:1px solid #1E3A5F;'>
          NOME DO ALUNO
        </th>
      </tr>
    </thead>
    <tbody>{linhas}</tbody>
  </table>
  <div style='margin-top:12px;text-align:center;font-size:7.5px;color:#94A3B8;
              border-top:1px solid #E2E8F0;padding-top:7px;'>
    Sistema Esporte e Saúde — Gestão Inteligente MoveRight™ &nbsp;|&nbsp;
    Emitido em: {hoje_fmt} &nbsp;|&nbsp; Data de referência: {data_fmt}{pag_info}
  </div>
</div>"""


def _renderizar_aba_prestacao_diaria():
    from utils.identidade import get_config as _gcfg
    from utils.imagem import get_base64_image

    st.markdown("""
        <div style='background:#EFF6FF;border-left:4px solid #1E88E5;
                    padding:12px 16px;border-radius:6px;margin-bottom:18px;'>
            <strong style='color:#1E3A5F;'>📋 Prestação de Contas Diária</strong><br>
            <span style='color:#1D4ED8;font-size:13px;'>
                Selecione um período para gerar a lista de presença de cada dia útil.
                Sábados, domingos e feriados nacionais são excluídos automaticamente.
                Dias úteis sem frequência lançada são sinalizados em alerta separado.
                Suporta até 1 ano de período por consulta.
            </span>
        </div>
    """, unsafe_allow_html=True)

    hoje = datetime.date.today()
    c_ini, c_fim, c_btn = st.columns([2, 2, 1], vertical_alignment="bottom")
    data_ini = c_ini.date_input(
        "📅 Data Inicial:", value=hoje, format="DD/MM/YYYY", key="pd_data_ini"
    )
    data_fim = c_fim.date_input(
        "📅 Data Final:", value=hoje, format="DD/MM/YYYY", key="pd_data_fim"
    )
    gerar = c_btn.button("🔍 Buscar Presenças", type="primary",
                         use_container_width=True, key="pd_buscar")

    if not gerar:
        st.caption("Selecione o período e clique em **Buscar Presenças**.")
        return

    if data_fim < data_ini:
        st.error("⚠️ A data final não pode ser anterior à data inicial.")
        return

    if (data_fim - data_ini).days > 366:
        st.error("⚠️ Período máximo permitido: 1 ano por vez.")
        return

    # ── Calcula feriados e dias úteis do período ──────────────────────────────
    anos_no_periodo = {data_ini.year + i for i in range((data_fim - data_ini).days // 365 + 2)
                       if (data_ini.year + i) <= data_fim.year}
    feriados = _feriados_nacionais_br(anos_no_periodo)

    def eh_dia_util(d: datetime.date) -> bool:
        return d.weekday() < 5 and d not in feriados  # 0=Seg … 4=Sex

    dias_uteis_range = sorted(
        data_ini + datetime.timedelta(days=i)
        for i in range((data_fim - data_ini).days + 1)
        if eh_dia_util(data_ini + datetime.timedelta(days=i))
    )

    # ── Busca presenças do período ─────────────────────────────────────────────
    with st.spinner(f"Consultando presenças de {data_ini.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}…"):
        por_dia_raw = get_presentes_periodo_todos(str(data_ini), str(data_fim))

    # Remove fins de semana e feriados da base retornada (segurança extra)
    por_dia = {
        iso: nomes
        for iso, nomes in por_dia_raw.items()
        if eh_dia_util(datetime.date.fromisoformat(iso))
    }

    # ── Dias úteis SEM frequência lançada ─────────────────────────────────────
    datas_com_frequencia = {datetime.date.fromisoformat(iso) for iso in por_dia}
    dias_sem_frequencia = [d for d in dias_uteis_range if d not in datas_com_frequencia]

    sufixo = (
        data_ini.strftime("%Y%m%d")
        if data_ini == data_fim else
        f"{data_ini.strftime('%Y%m%d')}_a_{data_fim.strftime('%Y%m%d')}"
    )
    ini_fmt = data_ini.strftime("%d/%m/%Y")
    fim_fmt = data_fim.strftime("%d/%m/%Y")

    # ══════════════════════════════════════════════════════════════════════════
    # ALERTA — dias úteis sem frequência
    # ══════════════════════════════════════════════════════════════════════════
    if dias_sem_frequencia:
        nomes_dias_semana = {
            "Segunda-feira": "Seg", "Terça-feira": "Ter", "Quarta-feira": "Qua",
            "Quinta-feira": "Qui", "Sexta-feira": "Sex",
        }
        linhas_alerta = ""
        itens_alerta = []
        for i, d in enumerate(dias_sem_frequencia, 1):
            dsem = _dia_da_semana_pt(d)
            d_fmt = d.strftime("%d/%m/%Y")
            item_str = f"{d_fmt} ({dsem})"
            itens_alerta.append(item_str)
            bg = "#FFF5F5" if i % 2 == 1 else "#FFECEC"
            linhas_alerta += f"""
            <tr style='background:{bg};'>
              <td style='padding:5px 10px;font-weight:700;color:#7F1D1D;
                         font-size:12px;border:1px solid #FECACA;text-align:center;
                         width:36px;'>{i}</td>
              <td style='padding:5px 12px;font-size:12px;font-weight:600;
                         color:#991B1B;border:1px solid #FECACA;'>{d_fmt}</td>
              <td style='padding:5px 12px;font-size:12px;color:#7F1D1D;
                         border:1px solid #FECACA;'>{dsem}</td>
              <td style='padding:5px 12px;font-size:11px;color:#B91C1C;font-style:italic;
                         border:1px solid #FECACA;'>Sem frequência lançada</td>
            </tr>"""

        st.markdown(f"""
<div style='border:2px solid #FCA5A5;border-radius:8px;background:#FEF2F2;
            padding:18px 20px;margin-bottom:18px;'>
  <div style='font-size:14px;font-weight:900;color:#991B1B;margin-bottom:10px;'>
    ⚠️ ATENÇÃO — {len(dias_sem_frequencia)} dia(s) útil(is) sem frequência registrada
  </div>
  <div style='font-size:12px;color:#7F1D1D;margin-bottom:12px;'>
    Os dias abaixo são dias úteis (seg–sex, excluindo feriados) sem nenhuma presença lançada
    no sistema. Verifique se houve aula e corrija se necessário.
  </div>
  <table style='width:100%;border-collapse:collapse;'>
    <thead>
      <tr style='background:#DC2626;'>
        <th style='padding:6px 10px;color:#fff;font-size:10px;border:1px solid #DC2626;
                   width:36px;'>#</th>
        <th style='padding:6px 12px;color:#fff;font-size:10px;text-align:left;
                   border:1px solid #DC2626;'>Data</th>
        <th style='padding:6px 12px;color:#fff;font-size:10px;text-align:left;
                   border:1px solid #DC2626;'>Dia da Semana</th>
        <th style='padding:6px 12px;color:#fff;font-size:10px;text-align:left;
                   border:1px solid #DC2626;'>Situação</th>
      </tr>
    </thead>
    <tbody>{linhas_alerta}</tbody>
  </table>
</div>
        """, unsafe_allow_html=True)

        # Botão de download do alerta em PDF
        with st.spinner("Preparando PDF do alerta…"):
            pdf_alerta = criar_pdf_alerta_frequencia(itens_alerta, ini_fmt, fim_fmt)
        st.download_button(
            label=f"🖨️ Imprimir Alerta — {len(dias_sem_frequencia)} dia(s) sem frequência",
            data=pdf_alerta,
            file_name=f"Alerta_Frequencia_{sufixo}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="pd_alerta_download",
        )
        st.markdown("<hr style='margin:16px 0;border-color:#E2E8F0;'>", unsafe_allow_html=True)

    # ══════════════════════════════════════════════════════════════════════════
    # RELATÓRIO PRINCIPAL — dias com frequência
    # ══════════════════════════════════════════════════════════════════════════
    if not por_dia:
        st.info("Nenhum dia útil com frequência lançada no período selecionado.")
        return

    total_dias   = len(por_dia)
    total_alunos = sum(len(v) for v in por_dia.values())

    m1, m2, m3 = st.columns(3)
    m1.metric("📅 Dias com aula", total_dias)
    m2.metric("👥 Total de presenças", total_alunos)
    m3.metric("📊 Média por dia", f"{total_alunos / total_dias:.1f}")

    st.markdown("<hr style='margin:12px 0;border-color:#E2E8F0;'>", unsafe_allow_html=True)

    with st.spinner("Gerando PDF…"):
        pdf_bytes = criar_prestacao_periodo_pdf(por_dia)

    st.download_button(
        label=f"📄 Baixar PDF — {total_dias} dia(s) / {total_alunos} presenças",
        data=pdf_bytes,
        file_name=f"Presenca_Periodo_{sufixo}.pdf",
        mime="application/pdf",
        type="primary",
        use_container_width=True,
        key="pd_download",
    )

    st.markdown("<hr style='margin:14px 0 10px 0;border-color:#E2E8F0;'>", unsafe_allow_html=True)

    # ── Preview visual — um bloco por dia ────────────────────────────────────
    st.markdown(f"#### 🖥️ Preview do Relatório — {total_dias} página(s)")

    cfg      = _gcfg()
    logo_p   = get_base64_image(cfg.get("logo_principal",  "logo-imbra.png"))
    logo_s   = get_base64_image(cfg.get("logo_secundaria", "logo-secretaria.png"))
    img_p    = (f"<img src='data:image/png;base64,{logo_p}' style='height:48px;object-fit:contain;'>") if logo_p else ""
    img_s    = (f"<img src='data:image/png;base64,{logo_s}' style='height:48px;object-fit:contain;'>") if logo_s else ""
    nome_org    = cfg.get("nome_organizacao",  "INSTITUTO MUDA BRASIL").upper()
    titulo_proj = cfg.get("titulo_projeto", "ESPORTE E SAÚDE NA COMUNIDADE - FASE 2").upper()
    hoje_fmt    = hoje.strftime("%d/%m/%Y")

    for idx, (data_iso, nomes) in enumerate(por_dia.items(), 1):
        try:
            data_fmt = datetime.date.fromisoformat(data_iso).strftime("%d/%m/%Y")
        except Exception:
            data_fmt = data_iso

        bloco = _bloco_preview_dia(
            data_fmt, nomes, img_s, img_p, nome_org, titulo_proj, hoje_fmt,
            numero_dia=idx, total_dias=total_dias,
        )
        st.markdown(bloco, unsafe_allow_html=True)


# ==============================================================================
# ABA 7: MONITORAMENTO CLÍNICO — B.I. DA SAÚDE
# ==============================================================================
def _extrair_secao_hashtag(texto: str, hashtag: str) -> str:
    """Extrai o conteúdo de uma seção #Hashtag: de um texto multilinha."""
    if not texto:
        return ""
    pattern = rf"#{re.escape(hashtag)}:\s*(.+?)(?=\n\n#|\Z)"
    m = re.search(pattern, texto, re.IGNORECASE | re.DOTALL)
    return m.group(1).strip() if m else ""


PALAVRAS_ALERTA = ["cirurgia", "cardiaco", "cardíaco", "pressao", "pressão", "infarto",
                   "marcapasso", "insuficiencia", "coagulação", "anticoagulante"]

BORG_OPCOES = ["0 — Nenhum", "1 — Muito leve", "2 — Leve", "3 — Moderado",
               "4 — Ligeiramente pesado", "5 — Pesado", "6", "7 — Muito pesado",
               "8", "9", "10 — Máximo / Emergência"]


def _tem_alerta(texto: str) -> bool:
    t = texto.lower() if texto else ""
    return any(p in t for p in PALAVRAS_ALERTA)


def _renderizar_monitoramento_clinico():
    st.markdown("""
        <div style='background:linear-gradient(135deg,#FFF1F2,#FEF2F2);
                    border-left:5px solid #EF4444;padding:18px 20px;
                    border-radius:8px;margin-bottom:20px;'>
            <h3 style='margin:0 0 4px 0;color:#991B1B;'>🏥 Monitoramento Clínico — B.I. da Saúde</h3>
            <p style='margin:0;color:#B91C1C;font-size:13px;'>
                Painel de classificação de risco por aluno. Use a coluna <strong>Borg / Risco</strong>
                para registrar o nível percebido e identificar alertas clínicos automaticamente.
                <br>⚠️ <em>Dados confidenciais — acesso restrito à equipe técnica.</em>
            </p>
        </div>
    """, unsafe_allow_html=True)

    # ── Filtros ────────────────────────────────────────────────────────────────
    c_turma, c_alerta, c_busca = st.columns([2, 1, 2])
    with c_turma:
        df_turmas = get_todas_turmas(ativas_apenas=True)
        opcoes_turma = ["Todas as Turmas"] + (df_turmas["nome"].tolist() if not df_turmas.empty else [])
        turma_filtro = st.selectbox("Turma:", opcoes_turma, key="mc_turma")
    with c_alerta:
        so_alertas = st.checkbox("🔴 Apenas Alertas", key="mc_alertas")
    with c_busca:
        busca_mc = st.text_input("🔍 Buscar aluno:", key="mc_busca", placeholder="Nome…")

    # ── Carrega alunos ─────────────────────────────────────────────────────────
    with st.spinner("Carregando dados clínicos…"):
        df_raw = buscar_alunos_geral("", incluir_inativos=False)

    if df_raw is None or df_raw.empty:
        st.warning("Nenhum aluno ativo encontrado.")
        return

    df_raw = df_raw[df_raw.get("status", "Ativo") != "Inativo"] if "status" in df_raw.columns else df_raw

    # ── Aplica filtros ─────────────────────────────────────────────────────────
    if turma_filtro != "Todas as Turmas":
        df_raw = df_raw[df_raw["turma"] == turma_filtro]

    if busca_mc and len(busca_mc.strip()) >= 2:
        from utils.texto import remover_acentos as _ra
        mask = df_raw["nome"].apply(_ra).str.contains(_ra(busca_mc), case=False, na=False)
        df_raw = df_raw[mask]

    # ── Monta tabela de monitoramento ─────────────────────────────────────────
    registros = []
    for _, r in df_raw.iterrows():
        ps = str(r.get("problemas_saude") or "")
        patologias  = _extrair_secao_hashtag(ps, "Patologias")    or ps[:80] if ps else ""
        restricoes  = _extrair_secao_hashtag(ps, "Restrições_Físicas")
        alergias    = _extrair_secao_hashtag(ps, "Alergias")
        incomodos   = _extrair_secao_hashtag(ps, "Incômodos_Físicos")
        medicament  = _extrair_secao_hashtag(ps, "Uso_Contínuo_Medicamentos")

        alerta = _tem_alerta(patologias) or _tem_alerta(ps)

        # Borg salvo em session_state (editável na sessão)
        borg_key = f"borg_mc_{r['id']}"
        borg_atual = st.session_state.get(borg_key, 0)

        registros.append({
            "id":             str(r["id"]),
            "🔴":             "🔴" if alerta else "",
            "Nome":           str(r.get("nome", "")),
            "Turma":          str(r.get("turma") or ""),
            "Patologias":     patologias[:120],
            "Restrições":     restricoes[:80],
            "Alergias":       alergias[:60],
            "Incômodos":      incomodos[:60],
            "Medicamentos":   medicament[:80],
            "Ct. Emergência": str(r.get("contato_emergencia", "") or "")[:100],
            "Borg/Risco":     borg_atual,
        })

    df_monitor = pd.DataFrame(registros)

    if so_alertas:
        df_monitor = df_monitor[df_monitor["🔴"] == "🔴"]

    if df_monitor.empty:
        st.info("Nenhum aluno encontrado com os filtros selecionados.")
        return

    # ── Legenda Borg ──────────────────────────────────────────────────────────
    with st.expander("📖 Legenda da Escala Borg (0–10)", expanded=False):
        st.markdown("""
        | Valor | Significado |
        |-------|-------------|
        | 0 | Nenhum esforço / Sem dor |
        | 1–2 | Muito leve |
        | 3–4 | Moderado — atenção |
        | 5–6 | Pesado — monitorar de perto |
        | 7–8 | Muito pesado — considerar adaptações |
        | 9–10 | **Máximo / Emergência — intervir imediatamente** |
        """)

    # ── Editor Borg (manter editável para classificações) ────────────────────
    st.markdown(f"**{len(df_monitor)} aluno(s) exibidos** — edite a coluna **Borg/Risco** (0–10) e clique em 💾 Salvar:")

    col_config = {
        "id":               st.column_config.NumberColumn("ID", disabled=True, width="small"),
        "🔴":               st.column_config.TextColumn("⚠️", width="small"),
        "Nome":             st.column_config.TextColumn("Nome", disabled=True, width="medium"),
        "Turma":            st.column_config.TextColumn("Turma", disabled=True, width="small"),
        "Patologias":       st.column_config.TextColumn("Patologias / Saúde", disabled=True, width="large"),
        "Restrições":       st.column_config.TextColumn("Restrições Físicas", disabled=True, width="medium"),
        "Alergias":         st.column_config.TextColumn("Alergias", disabled=True, width="medium"),
        "Incômodos":        st.column_config.TextColumn("Incômodos", disabled=True, width="medium"),
        "Medicamentos":     st.column_config.TextColumn("Medicamentos", disabled=True, width="medium"),
        "Ct. Emergência":   st.column_config.TextColumn("🚨 Ct. Emergência", disabled=True, width="medium",
                                help="Contato de emergência — Nome e Telefone cadastrados na ficha do aluno"),
        "Borg/Risco":       st.column_config.NumberColumn(
            "Borg / Risco (0–10)", min_value=0, max_value=10, step=1, width="small",
            help="0 = sem risco  ·  10 = emergência",
        ),
    }

    df_editado = st.data_editor(
        df_monitor, column_config=col_config,
        use_container_width=True, hide_index=True,
        key="data_editor_borg", num_rows="fixed",
    )

    # ── Botões de ação ────────────────────────────────────────────────────────
    c_salvar, c_reset, c_pdf, _ = st.columns([2, 1, 2, 1])
    with c_salvar:
        if st.button("💾 Salvar Borg (sessão)", type="primary", use_container_width=True, key="mc_salvar"):
            for _, row_e in df_editado.iterrows():
                st.session_state[f"borg_mc_{row_e['id']}"] = int(row_e["Borg/Risco"])
            st.toast(f"✅ {len(df_editado)} classificações salvas!", icon="💾")
    with c_reset:
        if st.button("🔄 Limpar Borg", use_container_width=True, key="mc_reset"):
            for _, row_e in df_editado.iterrows():
                k = f"borg_mc_{row_e['id']}"
                if k in st.session_state:
                    del st.session_state[k]
            st.rerun()
    with c_pdf:
        with st.spinner("Preparando PDF…"):
            try:
                _pdf_mc = gerar_pdf_monitoramento_clinico(
                    df_editado.drop(columns=["id"], errors="ignore"),
                    turma_filtro=turma_filtro,
                )
                st.download_button(
                    label="📄 Baixar PDF",
                    data=_pdf_mc,
                    file_name=f"monitoramento_clinico_{datetime.date.today().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                    key="mc_download_pdf",
                )
            except Exception as _ep:
                st.error(f"Erro PDF: {_ep}")

    # ── Preview idêntico ao PDF ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### 🖥️ Preview do Relatório (idêntico ao PDF)")

    _hoje_prev = datetime.date.today().strftime("%d/%m/%Y")
    _hora_prev = datetime.datetime.now().strftime("%H:%M")
    _n_alertas = int((df_editado["🔴"] == "🔴").sum())
    _n_borg7   = int((df_editado["Borg/Risco"].fillna(0).astype(int) >= 7).sum())

    # Linhas da tabela HTML
    def _wa_emergencia_html(contato: str) -> str:
        """Retorna célula HTML com contato de emergência e link WhatsApp se detectar número."""
        if not contato or contato in ("nan", "None", "—", ""):
            return "<span style='color:#94A3B8;font-size:9px;'>Não informado</span>"
        # Tenta extrair número de telefone (sequências de 8–13 dígitos)
        grupos = re.findall(r"\d{8,13}", contato)
        wa_url = None
        if grupos:
            phone = grupos[-1]
            if len(phone) <= 11:
                phone = "55" + phone
            wa_url = f"https://wa.me/{phone}"
        # Exibe contato completo; se tiver número, adiciona botão wa.me
        texto_exib = contato[:60] + ("…" if len(contato) > 60 else "")
        if wa_url:
            return (
                f"<span style='font-size:9px;display:block;'>{texto_exib}</span>"
                f"<a href='{wa_url}' target='_blank' "
                f"style='font-size:9px;color:#fff;background:#25D366;padding:1px 6px;"
                f"border-radius:8px;text-decoration:none;font-weight:700;"
                f"display:inline-block;margin-top:2px;'>📲 WhatsApp</a>"
            )
        return f"<span style='font-size:9px;'>{texto_exib}</span>"

    _linhas_html = ""
    for _zi, (_, _row) in enumerate(df_editado.iterrows()):
        _alerta = str(_row.get("🔴", "")) == "🔴"
        _bv     = int(_row.get("Borg/Risco", 0) or 0)
        _borg7  = _bv >= 7
        if _alerta:
            _bg = "background:#FEE2E2;"
        elif _borg7:
            _bg = "background:#FEF3C7;"
        elif _zi % 2 == 1:
            _bg = "background:#F5F7FA;"
        else:
            _bg = "background:#FFFFFF;"
        _icone    = "🔴" if _alerta else ""
        _borg_txt = str(_bv) if _bv > 0 else "—"
        _ct_em_html = _wa_emergencia_html(str(_row.get("Ct. Emergência", "") or ""))
        _linhas_html += (
            f"<tr style='{_bg}'>"
            f"<td style='text-align:center;padding:3px 4px;border:1px solid #CBD5E1;font-size:10px;'>{_icone}</td>"
            f"<td style='padding:3px 5px;border:1px solid #CBD5E1;font-size:10px;font-weight:600;'>{_row.get('Nome','')}</td>"
            f"<td style='padding:3px 4px;border:1px solid #CBD5E1;font-size:10px;'>{_row.get('Turma','')}</td>"
            f"<td style='padding:3px 5px;border:1px solid #CBD5E1;font-size:9.5px;'>{_row.get('Patologias','')}</td>"
            f"<td style='padding:3px 4px;border:1px solid #CBD5E1;font-size:9.5px;'>{_row.get('Restrições','')}</td>"
            f"<td style='padding:3px 4px;border:1px solid #CBD5E1;font-size:9.5px;'>{_row.get('Alergias','')}</td>"
            f"<td style='padding:3px 4px;border:1px solid #CBD5E1;font-size:9.5px;'>{_row.get('Incômodos','')}</td>"
            f"<td style='padding:3px 4px;border:1px solid #CBD5E1;font-size:9.5px;'>{_row.get('Medicamentos','')}</td>"
            f"<td style='padding:3px 4px;border:1px solid #CBD5E1;font-size:9.5px;'>{_ct_em_html}</td>"
            f"<td style='text-align:center;padding:3px 4px;border:1px solid #CBD5E1;font-size:10px;font-weight:700;"
            f"color:{'#991B1B' if _borg7 else '#0A2540'};'>{_borg_txt}</td>"
            f"</tr>"
        )

    _preview_html = f"""
    <div style="font-family:Arial,sans-serif;background:#fff;border:1px solid #CBD5E1;
                border-radius:8px;overflow:hidden;padding:12px 14px;margin-top:4px;">
      <!-- Cabeçalho -->
      <div style="border-bottom:2px solid #0056b3;padding-bottom:8px;margin-bottom:6px;
                  display:flex;align-items:center;justify-content:space-between;">
        <div style="font-size:11px;color:#64748B;">Instituto Muda Brasil</div>
        <div style="text-align:center;">
          <div style="font-size:13px;font-weight:900;color:#0A2540;">Monitoramento Clínico — B.I. da Saúde</div>
          <div style="font-size:9px;color:#64748B;">Emitido em {_hoje_prev} às {_hora_prev} &nbsp;|&nbsp; Turma: {turma_filtro}</div>
        </div>
        <div style="font-size:11px;color:#64748B;text-align:right;">{len(df_editado)} alunos</div>
      </div>
      <!-- Metadados de alerta -->
      <div style="font-size:9.5px;color:#991B1B;font-weight:700;margin-bottom:5px;">
        CONFIDENCIAL — Alertas clínicos: {_n_alertas} &nbsp;|&nbsp; Borg ≥ 7: {_n_borg7}
      </div>
      <!-- Legenda -->
      <div style="font-size:8.5px;margin-bottom:6px;display:flex;gap:12px;">
        <span><span style="display:inline-block;width:10px;height:10px;background:#FEE2E2;border:1px solid #CBD5E1;vertical-align:middle;"></span> Alerta clínico</span>
        <span><span style="display:inline-block;width:10px;height:10px;background:#FEF3C7;border:1px solid #CBD5E1;vertical-align:middle;"></span> Borg ≥ 7</span>
        <span><span style="display:inline-block;width:10px;height:10px;background:#F5F7FA;border:1px solid #CBD5E1;vertical-align:middle;"></span> Normal (zebra)</span>
      </div>
      <!-- Tabela -->
      <table style="width:100%;border-collapse:collapse;">
        <thead>
          <tr style="background:#0A2540;color:#fff;">
            <th style="padding:4px 3px;font-size:9.5px;border:1px solid #1E3A5F;width:3%;">⚠️</th>
            <th style="padding:4px 5px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:14%;">Nome</th>
            <th style="padding:4px 4px;font-size:9.5px;border:1px solid #1E3A5F;width:7%;">Turma</th>
            <th style="padding:4px 5px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:18%;">Patologias / Saúde</th>
            <th style="padding:4px 4px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:10%;">Restrições</th>
            <th style="padding:4px 4px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:7%;">Alergias</th>
            <th style="padding:4px 4px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:8%;">Incômodos</th>
            <th style="padding:4px 4px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:11%;">Medicamentos</th>
            <th style="padding:4px 4px;font-size:9.5px;border:1px solid #1E3A5F;text-align:left;width:17%;">🚨 Ct. Emergência</th>
            <th style="padding:4px 3px;font-size:9.5px;border:1px solid #1E3A5F;width:5%;">Borg</th>
          </tr>
        </thead>
        <tbody>{_linhas_html}</tbody>
      </table>
      <!-- Rodapé -->
      <div style="border-top:1px solid #E2E8F0;margin-top:6px;padding-top:5px;
                  font-size:8.5px;color:#991B1B;font-weight:700;">
        Resumo: {len(df_editado)} alunos &nbsp;|&nbsp; {_n_alertas} alertas clínicos &nbsp;|&nbsp; {_n_borg7} com Borg ≥ 7
      </div>
      <div style="font-size:7.5px;color:#94A3B8;margin-top:2px;">
        Documento confidencial — uso restrito à equipe técnica. Proibida reprodução sem autorização.
      </div>
    </div>
    """
    st.markdown(_preview_html, unsafe_allow_html=True)

    # ── Alertas ativos ────────────────────────────────────────────────────────
    alertas_borg = df_editado[df_editado["Borg/Risco"] >= 7]
    alertas_clin = df_editado[df_editado["🔴"] == "🔴"]
    if not alertas_borg.empty or not alertas_clin.empty:
        st.markdown("### 🚨 Alertas Ativos")
        col_b, col_c = st.columns(2)
        with col_b:
            if not alertas_borg.empty:
                st.error(f"⚡ **{len(alertas_borg)} aluno(s) com Borg ≥ 7:**")
                for _, ar in alertas_borg.iterrows():
                    st.markdown(f"  - **{ar['Nome']}** ({ar['Turma']}) — Borg: `{int(ar['Borg/Risco'])}`")
        with col_c:
            if not alertas_clin.empty:
                st.warning(f"🔴 **{len(alertas_clin)} aluno(s) com alertas clínicos:**")
                for _, ar in alertas_clin.iterrows():
                    st.markdown(f"  - **{ar['Nome']}** ({ar['Turma']}) — {ar['Patologias'][:60]}")

    st.info(
        "💡 As classificações Borg são mantidas durante esta sessão. "
        "Para persistência permanente, execute o SQL: "
        "`ALTER TABLE alunos ADD COLUMN IF NOT EXISTS risco_borg integer DEFAULT 0;` "
        "e informe para que o sistema salve automaticamente."
    )


# ==============================================================================
# RENDERIZAÇÃO DA INTERFACE PRINCIPAL (ST)
# ==============================================================================
def tela_relatorio():
    st.markdown(
        "<div style='background:linear-gradient(135deg,#F8FAFC,#E0F2FE);padding:25px;border-radius:12px;border-left:6px solid #1E88E5;'><h2>📄 Central de Relatórios e Auditoria</h2><p>Planilhas de Frequência, B.I. Analítico e Prestação de Contas Oficial.</p></div>",
        unsafe_allow_html=True,
    )

    tab_f, tab_id, tab_a, tab_w, tab_diario, tab_sem_av, tab_clinico = st.tabs(
        [
            "📊 Plan. Frequência",
            "🪪 Cara-Crachá",
            "🔎 Auditoria",
            "🏆 Prestação Pedagógica",
            "📋 Prestação Diária",
            "🧪 Avaliações",
            "🏥 Monitoramento",
        ]
    )

    # ==============================================================================
    # --- ABA 1: FREQUÊNCIA (MOTOR ANTI-FURO IMPLEMENTADO) ---
    # ==============================================================================
    with tab_f:
        c1, c2, c3 = st.columns([1, 1, 2], vertical_alignment="bottom")
        d_i = c1.date_input(
            "Data de Início", datetime.date.today().replace(day=1), format="DD/MM/YYYY"
        )
        d_f = c2.date_input("Data de Fim", datetime.date.today(), format="DD/MM/YYYY")

        turmas = get_todas_turmas(ativas_apenas=True)
        t_sel = c3.selectbox(
            "Filtrar Turma",
            ["Todas as Turmas"] + turmas["nome"].tolist()
            if not turmas.empty
            else ["Todas as Turmas"],
        )

        if st.button(
            "🔍 Processar Frequência",
            type="primary",
            use_container_width=True,
        ):
            with st.spinner("Cruzando Diário de Aulas com registros de frequência..."):
                df_matriz = get_relatorio_periodo(
                    d_i, d_f, "" if t_sel == "Todas as Turmas" else t_sel
                )

                if df_matriz.empty:
                    st.warning(
                        "⚠️ Não foram encontradas aulas no Diário para este período."
                    )
                    return

                # Renomear para compatibilidade Excel
                if "Nome" in df_matriz.columns:
                    df_matriz.rename(columns={"Nome": "Aluno"}, inplace=True)

                # Colunas meta — excluindo "Total Aulas" que é metadado, não data
                _META = {"Ordem", "Aluno", "Turma", "Total Aulas",
                         "Total P", "Total F", "Total J", "% Presença"}
                # Filtra APENAS datas com ao menos 1 registro real (P/F/J) —
                # dias sem nenhuma aula registrada são excluídos de tudo:
                # cálculos, gráficos, planilha e totais.
                cols_data_reais = [
                    c for c in df_matriz.columns
                    if c not in _META
                    and df_matriz[c].isin(["P", "F", "J"]).any()
                ]

                # Ordena cronologicamente (evita datas fora de ordem no PDF/Excel)
                def _parse_col_dt(col, _di=d_i, _df=d_f):
                    for yr in sorted({_di.year, _df.year}):
                        try:
                            dt = datetime.datetime.strptime(f"{col}/{yr}", "%d/%m/%Y").date()
                            if _di <= dt <= _df:
                                return dt
                        except Exception:
                            pass
                    return datetime.date.max

                cols_data_reais = sorted(cols_data_reais, key=_parse_col_dt)

                # Totais globais
                tp_geral = int(df_matriz["Total P"].sum())
                tf_geral = int(df_matriz["Total F"].sum())
                tj_geral = int(df_matriz["Total J"].sum()) if "Total J" in df_matriz.columns else 0
                n_alunos  = len(df_matriz)

                # Total de aulas: soma por turma (cada turma tem seu próprio calendário)
                if "Total Aulas" in df_matriz.columns and "Turma" in df_matriz.columns:
                    n_aulas_por_turma = df_matriz.groupby("Turma")["Total Aulas"].first()
                    n_aulas = int(n_aulas_por_turma.sum())
                    if len(n_aulas_por_turma) > 1:
                        partes = " + ".join(str(v) for v in n_aulas_por_turma.values)
                        n_aulas_total_desc = f"{partes} = {n_aulas}"
                    else:
                        n_aulas_total_desc = str(n_aulas)
                else:
                    n_aulas = len(cols_data_reais)
                    n_aulas_total_desc = str(n_aulas)

                # Taxa geral de presença (base: P / (P+F+J) que é o total de aulas reais)
                total_registros = tp_geral + tf_geral + tj_geral
                taxa_geral = round(tp_geral / total_registros * 100, 1) if total_registros > 0 else 0.0

                # Alunos em risco: < 75 % de presença
                def _perc(s):
                    try:
                        return float(str(s).replace("%", "").strip())
                    except Exception:
                        return 0.0
                df_matriz["_taxa_num"] = df_matriz["% Presença"].apply(_perc)
                df_risco = df_matriz[df_matriz["_taxa_num"] < 75].copy()

                periodo_formatado = f"{d_i.strftime('%d/%m/%Y')} a {d_f.strftime('%d/%m/%Y')}"

                # ── Inserir Ordem agora (após cálculos)
                df_matriz.insert(0, "Ordem", range(1, 1 + len(df_matriz)))

                # ── Linha totalizadora por dia
                tot_d = {"Ordem": "-", "Aluno": "TOTAL PRESENÇAS / DIA", "Turma": "-",
                         "Total Aulas": "-"}
                for c in df_matriz.columns:
                    if c in tot_d:
                        continue
                    if c in cols_data_reais:
                        tot_d[c] = int((df_matriz[c] == "P").sum())
                    elif c == "Total P":
                        tot_d[c] = tp_geral
                    elif c == "Total F":
                        tot_d[c] = tf_geral
                    elif c == "Total J":
                        tot_d[c] = tj_geral
                    elif c == "_taxa_num":
                        tot_d[c] = ""
                    else:
                        tot_d[c] = "-"

                df_final = pd.concat([pd.DataFrame([tot_d]), df_matriz], ignore_index=True)
                # Garante tipos homogêneos na coluna Ordem (evita erro Arrow int/str misto)
                df_final["Ordem"] = df_final["Ordem"].astype(str)

                # ── Excel
                excel = gerar_excel_planilha_frequencia(
                    df_final.drop(columns=["_taxa_num"], errors="ignore"),
                    t_sel,
                    periodo_formatado,
                    "logo-imbra.png",
                    "logo-secretaria.png",
                    n_alunos,
                    tp_geral,
                    n_aulas,
                )

                # ── Gerar PDF Prime junto com o Excel ────────────────────
                taxa_media_pdf  = df_matriz["_taxa_num"].mean() if not df_matriz.empty else 0.0
                media_p_pdf     = tp_geral / n_alunos if n_alunos > 0 else 0
                n_exc_pdf  = int((df_matriz["_taxa_num"] >= 90).sum())  if not df_matriz.empty else 0
                n_reg_pdf  = int(((df_matriz["_taxa_num"] >= 75) & (df_matriz["_taxa_num"] < 90)).sum()) if not df_matriz.empty else 0
                n_at_pdf   = int(((df_matriz["_taxa_num"] >= 50) & (df_matriz["_taxa_num"] < 75)).sum()) if not df_matriz.empty else 0
                n_crit_pdf = int((df_matriz["_taxa_num"] < 50).sum())   if not df_matriz.empty else 0
                n_risco_pdf= int((df_matriz["_taxa_num"] < 75).sum())   if not df_matriz.empty else 0

                with st.spinner("Gerando PDF do Relatório Prime..."):
                    pdf_prime = _gerar_pdf_relatorio_prime(
                        df_matriz, cols_data_reais, periodo_formatado, t_sel,
                        n_alunos, n_aulas, tp_geral,
                        media_p_pdf, taxa_media_pdf,
                        n_exc_pdf, n_reg_pdf, n_at_pdf, n_crit_pdf, n_risco_pdf,
                        d_i, d_f,
                    )

                st.success(
                    f"✅ {n_alunos} alunos · {n_aulas_total_desc} aulas · período {periodo_formatado}"
                )

                _btn_xls, _btn_pdf = st.columns(2)
                with _btn_xls:
                    st.download_button(
                        "📥 PLANILHA DE FREQUÊNCIA (EXCEL)",
                        excel,
                        f"Frequencia_{t_sel}_{d_i.strftime('%d_%m_%Y')}.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                        type="primary",
                    )
                with _btn_pdf:
                    if pdf_prime:
                        st.download_button(
                            "📄 RELATÓRIO PRIME (PDF)",
                            pdf_prime,
                            f"RelPrime_{t_sel}_{d_i.strftime('%d_%m_%Y')}.pdf",
                            "application/pdf",
                            use_container_width=True,
                            type="primary",
                        )
                    else:
                        st.warning("PDF não disponível — verifique xhtml2pdf.")

                # ══════════════════════════════════════════════════════════════
                # RELATÓRIO PRIME — DASHBOARD ANALÍTICO DE FREQUÊNCIA
                # Foco: presença é o indicador. Falta aparece só como alerta.
                # ══════════════════════════════════════════════════════════════
                st.markdown("---")
                st.markdown(
                    "<div style='background:linear-gradient(135deg,#EFF6FF,#DBEAFE);"
                    "padding:14px 20px;border-radius:12px;border-left:6px solid #1D4ED8;"
                    "margin-bottom:18px;'>"
                    "<h3 style='margin:0;color:#1D4ED8;font-size:1.2rem;'>📊 Relatório Prime — Análise de Frequência</h3>"
                    f"<p style='margin:4px 0 0;color:#475569;font-size:12px;'>"
                    f"Período: {periodo_formatado} · {n_alunos} alunos · {n_aulas} aulas realizadas"
                    f"</p></div>",
                    unsafe_allow_html=True,
                )

                # ── Métricas avançadas ────────────────────────────────────────
                media_p_aluno_qtd = tp_geral / n_alunos if n_alunos > 0 else 0
                # Taxa média individual: média das taxas por aluno (pesa todos igualmente)
                taxa_media_ind = df_matriz["_taxa_num"].mean() if not df_matriz.empty else 0.0
                n_risco        = len(df_risco)
                n_excelentes   = int((df_matriz["_taxa_num"] >= 90).sum()) if not df_matriz.empty else 0
                n_regulares    = int(((df_matriz["_taxa_num"] >= 75) & (df_matriz["_taxa_num"] < 90)).sum()) if not df_matriz.empty else 0
                n_atencao      = int(((df_matriz["_taxa_num"] >= 50) & (df_matriz["_taxa_num"] < 75)).sum()) if not df_matriz.empty else 0
                n_critico      = int((df_matriz["_taxa_num"] < 50).sum()) if not df_matriz.empty else 0

                def _kpi(label, valor, sub="", cor="#1D4ED8"):
                    st.markdown(
                        f"<div style='background:#fff;border:1.5px solid #E2E8F0;"
                        f"border-radius:10px;padding:13px 14px;text-align:center;"
                        f"box-shadow:0 1px 4px rgba(0,0,0,0.06);'>"
                        f"<div style='font-size:8.5pt;color:#64748B;font-weight:700;"
                        f"text-transform:uppercase;letter-spacing:.5px;'>{label}</div>"
                        f"<div style='font-size:24px;font-weight:900;color:{cor};"
                        f"line-height:1.25;margin-top:4px;'>{valor}</div>"
                        f"<div style='font-size:7.5pt;color:#94A3B8;margin-top:3px;'>{sub}</div>"
                        f"</div>",
                        unsafe_allow_html=True,
                    )

                ck1, ck2, ck3, ck4, ck5 = st.columns(5)
                with ck1:
                    _kpi("Alunos no Período", n_alunos, "participantes ativos")
                with ck2:
                    _kpi("Aulas Realizadas", n_aulas,
                         n_aulas_total_desc if n_aulas_total_desc != str(n_aulas) else "sessões no período")
                with ck3:
                    _kpi(
                        "Média Presenças / Aluno",
                        f"{media_p_aluno_qtd:.1f}",
                        f"de {n_aulas} aulas · {tp_geral} presenças totais",
                    )
                with ck4:
                    cor_tx = "#10B981" if taxa_media_ind >= 75 else ("#F59E0B" if taxa_media_ind >= 50 else "#EF4444")
                    _kpi(
                        "Taxa Média Individual",
                        f"{taxa_media_ind:.1f}%",
                        f"meta 75% · cada aluno pesa igual",
                        cor=cor_tx,
                    )
                with ck5:
                    cor_ex = "#10B981" if n_excelentes > 0 else "#94A3B8"
                    _kpi(
                        "Excelente (≥ 90%)",
                        n_excelentes,
                        f"{n_regulares} regulares · {n_atencao} atenção · {n_critico} crítico",
                        cor=cor_ex,
                    )

                # ── Linha do tempo + Padrão por Dia da Semana ────────────────
                if cols_data_reais:
                    _DIAS_PT_R = {
                        "Monday": "Segunda", "Tuesday": "Terça", "Wednesday": "Quarta",
                        "Thursday": "Quinta", "Friday": "Sexta",
                        "Saturday": "Sábado", "Sunday": "Domingo",
                    }
                    _ORDEM_SEM = ["Segunda", "Terça", "Quarta", "Quinta", "Sexta", "Sábado", "Domingo"]

                    def _parse_col_dt(s, ini, fim):
                        for yr in sorted({ini.year, fim.year}):
                            try:
                                dt = datetime.datetime.strptime(f"{s}/{yr}", "%d/%m/%Y").date()
                                if ini <= dt <= fim:
                                    return dt
                            except Exception:
                                pass
                        return None

                    weekday_acc = {}  # dia → [soma_pres, n_dias]
                    tl_rows = []
                    for col in cols_data_reais:
                        pres = int((df_matriz[col] == "P").sum())
                        dt   = _parse_col_dt(col, d_i, d_f)
                        wd   = _DIAS_PT_R.get(dt.strftime("%A"), "?") if dt else "?"
                        tl_rows.append({"Data": col, "Presenças": pres, "Dia": wd})
                        if wd != "?":
                            acc = weekday_acc.setdefault(wd, [0, 0])
                            acc[0] += pres
                            acc[1] += 1

                    df_tl    = pd.DataFrame(tl_rows)
                    media_tl = df_tl["Presenças"].mean() if not df_tl.empty else 0

                    df_sem_r = pd.DataFrame([
                        {
                            "Dia": d,
                            "Media": round(weekday_acc[d][0] / weekday_acc[d][1], 1) if weekday_acc.get(d, [0, 0])[1] > 0 else 0,
                            "Aulas": weekday_acc[d][1] if d in weekday_acc else 0,
                        }
                        for d in _ORDEM_SEM
                        if d in weekday_acc
                    ])

                    ct1, ct2 = st.columns([3, 2], gap="large")
                    with ct1:
                        with st.container(border=True):
                            st.markdown("##### 📈 Presença Diária no Período")
                            fig_tl = px.bar(
                                df_tl, x="Data", y="Presenças",
                                color="Presenças",
                                color_continuous_scale=["#93C5FD", "#1D4ED8"],
                                text="Presenças",
                                custom_data=["Dia"],
                            )
                            fig_tl.add_hline(
                                y=media_tl, line_dash="dot", line_color="#10B981",
                                annotation_text=f"Média {media_tl:.1f}",
                                annotation_position="top left",
                            )
                            fig_tl.update_traces(
                                textposition="outside",
                                textfont_size=8,
                                hovertemplate="<b>%{x}</b> (%{customdata[0]})<br>Presenças: %{y}<extra></extra>",
                            )
                            fig_tl.update_layout(
                                height=290, coloraxis_showscale=False, showlegend=False,
                                paper_bgcolor="white", plot_bgcolor="white",
                                margin=dict(t=10, b=30, l=0, r=0),
                                xaxis=dict(showgrid=False, tickfont=dict(size=8)),
                                yaxis=dict(showgrid=True, gridcolor="#F1F5F9"),
                            )
                            st.plotly_chart(fig_tl, use_container_width=True)
                    with ct2:
                        with st.container(border=True):
                            st.markdown("##### 📅 Padrão por Dia da Semana")
                            if not df_sem_r.empty:
                                fig_wd = px.bar(
                                    df_sem_r, x="Dia", y="Media",
                                    text=df_sem_r["Media"].apply(lambda v: f"{v:.0f}"),
                                    color="Media",
                                    color_continuous_scale=["#BFDBFE", "#1D4ED8"],
                                    custom_data=["Aulas"],
                                )
                                fig_wd.update_traces(
                                    textposition="outside",
                                    hovertemplate="<b>%{x}</b><br>Média: %{y:.1f} presenças<br>%{customdata[0]} aula(s) no período<extra></extra>",
                                )
                                fig_wd.update_layout(
                                    height=230, coloraxis_showscale=False, showlegend=False,
                                    paper_bgcolor="white", plot_bgcolor="white",
                                    margin=dict(t=10, b=30, l=0, r=0),
                                    xaxis=dict(showgrid=False),
                                    yaxis=dict(showgrid=True, gridcolor="#F1F5F9"),
                                )
                                st.plotly_chart(fig_wd, use_container_width=True)
                                if len(df_sem_r) >= 2:
                                    pico_d = df_sem_r.loc[df_sem_r["Media"].idxmax(), "Dia"]
                                    low_d  = df_sem_r.loc[df_sem_r["Media"].idxmin(), "Dia"]
                                    st.caption(f"Pico de comparecimento: **{pico_d}** · Menor fluxo: **{low_d}**")

                # ── Resumo por Turma ──────────────────────────────────────────
                if "Turma" in df_matriz.columns:
                    grp = (
                        df_matriz.groupby("Turma")
                        .agg(
                            Alunos=("Aluno", "count"),
                            Aulas=("Total Aulas", "first"),
                            Presenças=("Total P", "sum"),
                        )
                        .reset_index()
                    )
                    grp["Esperado"]     = grp["Alunos"] * grp["Aulas"]
                    grp["Taxa %"]       = (grp["Presenças"] / grp["Esperado"].replace(0, 1) * 100).round(1)
                    grp["Media P/Aluno"]= (grp["Presenças"] / grp["Alunos"].replace(0, 1)).round(1)
                    grp["Taxa fmt"]     = grp["Taxa %"].apply(lambda v: f"{v:.1f}%")

                    cr1, cr2 = st.columns([1, 1], gap="large")
                    with cr1:
                        with st.container(border=True):
                            st.markdown("##### 🏫 Resumo por Turma")
                            st.dataframe(
                                grp[["Turma", "Alunos", "Aulas", "Presenças", "Media P/Aluno", "Taxa fmt"]]
                                .rename(columns={"Taxa fmt": "Taxa Presença", "Media P/Aluno": "Média P/Aluno"}),
                                use_container_width=True, hide_index=True,
                            )
                    with cr2:
                        with st.container(border=True):
                            st.markdown("##### 📊 Taxa de Presença por Turma")
                            fig_turma = px.bar(
                                grp, x="Turma", y="Taxa %",
                                text="Taxa fmt",
                                color="Taxa %",
                                color_continuous_scale=["#FEE2E2", "#FEF3C7", "#DCFCE7"],
                                range_color=[0, 100],
                            )
                            fig_turma.add_hline(
                                y=75, line_dash="dash", line_color="#6366F1",
                                annotation_text="Meta 75%", annotation_position="top left",
                            )
                            fig_turma.update_traces(textposition="outside")
                            fig_turma.update_layout(
                                coloraxis_showscale=False, showlegend=False, height=260,
                                paper_bgcolor="white", plot_bgcolor="white",
                                margin=dict(t=10, b=30, l=0, r=0),
                            )
                            st.plotly_chart(fig_turma, use_container_width=True)

                # ── Distribuição de Assiduidade + Rankings ────────────────────
                cd1, cd2 = st.columns([1, 1], gap="large")
                with cd1:
                    with st.container(border=True):
                        st.markdown("##### 📊 Distribuição de Assiduidade")
                        _bins   = [0, 50, 75, 90, 100.01]
                        _labels = ["< 50% (crítico)", "50–75% (atenção)", "75–90% (regular)", "≥ 90% (excelente)"]
                        _cores  = ["#EF4444", "#F59E0B", "#3B82F6", "#10B981"]
                        df_dist = (
                            pd.cut(df_matriz["_taxa_num"], bins=_bins, labels=_labels, right=False)
                            .value_counts()
                            .reindex(_labels)
                            .fillna(0)
                            .reset_index()
                        )
                        df_dist.columns = ["Faixa", "Alunos"]
                        df_dist["Pct"] = (df_dist["Alunos"] / n_alunos * 100).round(1).apply(lambda v: f"{v:.0f}%")
                        fig_dist = px.bar(
                            df_dist, x="Faixa", y="Alunos",
                            text=df_dist.apply(lambda r: f"{int(r['Alunos'])} ({r['Pct']})", axis=1),
                            color="Faixa",
                            color_discrete_map=dict(zip(_labels, _cores)),
                        )
                        fig_dist.update_traces(textposition="outside")
                        fig_dist.update_layout(
                            showlegend=False, height=280,
                            paper_bgcolor="white", plot_bgcolor="white",
                            margin=dict(t=10, b=30, l=0, r=0),
                            xaxis=dict(showgrid=False, tickfont=dict(size=9)),
                            yaxis=dict(showgrid=True, gridcolor="#F1F5F9"),
                        )
                        st.plotly_chart(fig_dist, use_container_width=True)
                with cd2:
                    with st.container(border=True):
                        st.markdown("##### 🏆 Rankings de Presença")
                        df_rank = (
                            df_matriz[["Aluno", "Turma", "Total P", "_taxa_num", "Total Aulas"]]
                            .copy()
                            .rename(columns={"Total P": "Presenças", "_taxa_num": "Taxa %", "Total Aulas": "Aulas"})
                        )
                        df_rank["Taxa"] = df_rank["Taxa %"].apply(lambda v: f"{v:.1f}%")
                        tab_top, tab_at = st.tabs(["⭐ Mais Assíduos", "⚠️ Precisa Atenção"])
                        with tab_top:
                            top_df = (
                                df_rank.nlargest(10, "Taxa %")[["Aluno", "Turma", "Presenças", "Aulas", "Taxa"]]
                                .reset_index(drop=True)
                            )
                            top_df.index += 1
                            st.dataframe(top_df, use_container_width=True)
                        with tab_at:
                            if n_risco > 0:
                                at_df = (
                                    df_rank[df_rank["Taxa %"] < 75]
                                    .nsmallest(len(df_rank), "Taxa %")[["Aluno", "Turma", "Presenças", "Aulas", "Taxa"]]
                                    .reset_index(drop=True)
                                )
                                at_df.index += 1
                                st.dataframe(at_df, use_container_width=True)
                                st.caption(f"{n_risco} aluno(s) abaixo de 75% — presença é hábito, não obrigação: acolha, não puna.")
                            else:
                                st.success("Todos os alunos acima de 75% — turma saudável!")

                # ── Planilha de Frequência Detalhada ─────────────────────────
                st.markdown("---")
                st.markdown("#### 📅 Planilha Detalhada — Frequência Individual")
                st.caption("Colunas prioritárias: Taxa e Total de Presenças aparecem antes das datas.")

                # Reordenação: % Presença e Total P logo após Ordem/Aluno/Turma
                _priority_cols = [c for c in ["Ordem", "Aluno", "Turma", "% Presença", "Total P", "Total Aulas"] if c in df_final.columns]
                _date_cols     = cols_data_reais
                _extra_cols    = [c for c in df_final.columns if c not in _priority_cols and c not in _date_cols and c not in {"_taxa_num", "Total F", "Total J"}]
                _col_order     = _priority_cols + _date_cols + _extra_cols
                _col_order     = [c for c in _col_order if c in df_final.columns]
                df_show        = df_final[_col_order]

                def colorir_status(val):
                    if val == "P":
                        return "color:#10B981;font-weight:bold;background-color:#D1FAE5;"
                    if val == "F":
                        return "color:#EF4444;font-weight:bold;background-color:#FEE2E2;"
                    if val == "J":
                        return "color:#F59E0B;font-weight:bold;background-color:#FEF3C7;"
                    return ""

                df_st = df_show.style.map(colorir_status).set_properties(
                    subset=[c for c in df_show.columns if c not in ["Aluno", "Turma"]],
                    **{"text-align": "center"},
                )
                st.dataframe(df_st, use_container_width=True, hide_index=True)
    # ==============================================================================
    # --- ABA 1.5: RELATÓRIO CARA-CRACHÁ ---
    # ==============================================================================
    with tab_id:
        renderizar_aba_caracracha()

    # ==============================================================================
    # --- ABA 2: AUDITORIA COM PDF NATIVO E GRID INTERATIVO ---
    # ==============================================================================
    with tab_a:
        st.markdown("### 🔎 Auditoria de Cadastros e Documentos")
        st.write(
            "Identifique pendências documentais e clique no botão para acessar e corrigir a ficha do aluno instantaneamente."
        )

        df_aud = buscar_alunos_geral("")
        if not df_aud.empty:
            with st.container(border=True):
                c_aud1, _ = st.columns(2)
                turma_aud = c_aud1.selectbox(
                    "Turma para Auditoria",
                    ["Todas"]
                    + sorted(
                        [
                            t
                            for t in df_aud["turma"].unique().tolist()
                            if isinstance(t, str)
                        ]
                    ),
                )
                if turma_aud != "Todas":
                    df_aud = df_aud[df_aud["turma"] == turma_aud]

            if st.button(
                "🚀 INICIAR VERIFICAÇÃO DE INTEGRIDADE", use_container_width=True
            ):
                # Nomes de coluna reais na tabela `alunos` do Supabase
                # CPF e RG são verificados em conjunto: basta um dos dois estar preenchido
                checks = {
                    "url_foto":            "📸 Foto",
                    "url_rg":              "🪪 Documento Oficial",
                    "data_nascimento":     "🎂 Nasc.",
                    "whatsapp":            "📱 WhatsApp",
                    "url_atestado_medico": "⚕️ Atestado Médico",
                }
                LABEL_IDENTIFICACAO = "🆔 CPF / RG"
                falhas = []
                contagem_falhas = {label: 0 for label in checks.values()}
                contagem_falhas[LABEL_IDENTIFICACAO] = 0

                # Strings que devem ser tratadas como vazio
                _VAZIOS = {"", "none", "null", "undefined", "nan", "não informado",
                           "nao informado", "n/a", "na", "-", "0"}

                def _ausente(val):
                    """Retorna True se o valor indica campo não preenchido."""
                    if val is None:
                        return True
                    try:
                        if pd.isna(val):
                            return True
                    except Exception:
                        pass
                    return str(val).strip().lower() in _VAZIOS

                with st.spinner("Verificando cadastros… aguarde."):
                    for _, r in df_aud.iterrows():
                        missing = []
                        for col, label in checks.items():
                            if _ausente(r.get(col)):
                                missing.append(label)
                                contagem_falhas[label] += 1
                        # CPF e RG: só falha se ambos estiverem ausentes
                        if _ausente(r.get("cpf")) and _ausente(r.get("rg")):
                            missing.append(LABEL_IDENTIFICACAO)
                            contagem_falhas[LABEL_IDENTIFICACAO] += 1
                        if missing:
                            falhas.append(
                                {
                                    "Aluno": r["nome"],
                                    "Turma": r["turma"],
                                    "Pendências": ", ".join(missing),
                                    "id_aluno": str(r.get("id", "")).split(".")[0],
                                    "url_foto": r.get("url_foto") or "",
                                    "dict_aluno": r.to_dict(),
                                }
                            )

                # Busca última presença de todos os alunos com pendência em 1 query
                if falhas:
                    _ids_batch = tuple(f["id_aluno"] for f in falhas)
                    _ult_pres = get_ultima_presenca_batch(_ids_batch)
                    for f in falhas:
                        f["Última Presença"] = _ult_pres.get(f["id_aluno"], "—")

                if falhas:
                    st.error(
                        f"⚠️ Identificadas irregularidades em {len(falhas)} alunos da turma {turma_aud}."
                    )
                    st.markdown("#### 📊 Totalizadores de Pendências")
                    cols_metric = st.columns(len(contagem_falhas))
                    for idx, (label, count) in enumerate(contagem_falhas.items()):
                        cols_metric[idx].metric(label.split(" ", 1)[-1], count)

                    # 🚀 REPOSICIONAMENTO E EXPORTAÇÃO PARA PDF REAL
                    st.markdown("#### 🖨️ Opções de Exportação Oficial")
                    c_dw1, c_dw2 = st.columns(2)

                    # Exportação Excel — espelha exatamente o grid na tela
                    def _foto_label(row):
                        url = str(row.get("url_foto") or "").strip()
                        return "Sim" if url.startswith("http") else "Nao"

                    _DSEM = ["seg","ter","qua","qui","sex","sáb","dom"]

                    def _fmt_dsem(raw):
                        v = str(raw or "").strip()
                        if not v or v in ("", "None", "-", "—", "nan"):
                            return "Sem registro"
                        try:
                            _d = datetime.datetime.strptime(v, "%d/%m/%y").date()
                            return f"{v} {_DSEM[_d.weekday()]}"
                        except Exception:
                            try:
                                _d = datetime.datetime.strptime(v, "%d/%m/%Y").date()
                                return f"{v} {_DSEM[_d.weekday()]}"
                            except Exception:
                                return v

                    def _ult_label(row):
                        raw = row.get("Última Presença") or row.get("Ultima Presenca") or ""
                        return _fmt_dsem(raw)

                    df_export_rows = [
                        {
                            "Foto": _foto_label(f),
                            "Nome do Aluno": f.get("Aluno", ""),
                            "Turma": f.get("Turma", ""),
                            "Últ. Presença": _ult_label(f),
                            "Pendências Identificadas": f.get("Pendências", ""),
                        }
                        for f in falhas
                    ]
                    df_export = pd.DataFrame(df_export_rows)
                    output_aud = io.BytesIO()
                    with pd.ExcelWriter(output_aud, engine="xlsxwriter") as writer:
                        pd.DataFrame(
                            list(contagem_falhas.items()),
                            columns=["Documento/Campo", "Quantidade em Falta"],
                        ).to_excel(writer, index=False, sheet_name="Resumo_Auditoria")
                        ws_res = writer.sheets["Resumo_Auditoria"]
                        ws_res.set_column(0, 0, 30)
                        ws_res.set_column(1, 1, 20)
                        df_export.to_excel(
                            writer, index=False, sheet_name="Detalhamento"
                        )
                        ws_det = writer.sheets["Detalhamento"]
                        ws_det.set_column(0, 0, 8)   # Foto
                        ws_det.set_column(1, 1, 35)  # Nome
                        ws_det.set_column(2, 2, 20)  # Turma
                        ws_det.set_column(3, 3, 15)  # Últ. Presença
                        ws_det.set_column(4, 4, 55)  # Pendências

                    c_dw1.download_button(
                        "📥 Exportar Lista Completa (Excel)",
                        output_aud.getvalue(),
                        f"Auditoria_{datetime.date.today().strftime('%d_%m_%Y')}.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                    )

                    # 🚀 Exportação PDF Real
                    pdf_auditoria = gerar_pdf_auditoria_core(
                        falhas, contagem_falhas, turma_aud
                    )
                    if pdf_auditoria:
                        c_dw2.download_button(
                            "📕 Imprimir Resumo (PDF Oficial)",
                            pdf_auditoria,
                            f"Auditoria_Oficial_{datetime.date.today().strftime('%d_%m_%Y')}.pdf",
                            "application/pdf",
                            use_container_width=True,
                        )
                    else:
                        c_dw2.error(
                            "⚠️ Biblioteca PDF (xhtml2pdf) não encontrada no servidor."
                        )

                    # Grid Interativo com Botões de Ação
                    st.markdown("<hr/>", unsafe_allow_html=True)
                    st.markdown("#### 📋 Detalhamento e Ação Rápida")
                    st.markdown(
                        "<div style='background-color: #F8FAFC; padding: 10px; border-radius: 8px; margin-bottom: 10px;'>",
                        unsafe_allow_html=True,
                    )
                    ch0, ch1, ch2, ch3, ch4, ch5 = st.columns(
                        [1, 2.2, 1.6, 3.8, 1.4, 1.6], vertical_alignment="center"
                    )
                    ch0.markdown("**Foto**")
                    ch1.markdown("**Nome do Aluno**")
                    ch2.markdown("**Turma**")
                    ch3.markdown("**Falta Preencher**")
                    ch4.markdown("**Últ. Presença**")
                    ch5.markdown("**Ação**")
                    st.markdown("</div>", unsafe_allow_html=True)

                    for f in falhas:
                        with st.container():
                            c0, c1, c2, c3, c4, c5 = st.columns(
                                [1, 2.2, 1.6, 3.8, 1.4, 1.6], vertical_alignment="center"
                            )
                            # Foto do aluno
                            _foto = f.get("url_foto", "")
                            if _foto and str(_foto).startswith("http"):
                                try:
                                    c0.image(_foto, width="stretch")
                                except Exception:
                                    c0.markdown("👤")
                            else:
                                c0.markdown("👤")
                            c1.write(f["Aluno"])
                            c2.write(f["Turma"])
                            c3.write(f["Pendências"])
                            c4.write(_fmt_dsem(f.get("Última Presença", "—")))
                            with c5:
                                st.button(
                                    "🩺 Abrir Ficha",
                                    key=f"aud_{f['id_aluno']}",
                                    on_click=abrir_ficha_aluno,
                                    args=(f["dict_aluno"],),
                                    use_container_width=True,
                                )
                            st.markdown(
                                "<hr style='margin: 0px 0px 10px 0px; border-color: #E2E8F0;'/>",
                                unsafe_allow_html=True,
                            )
                else:
                    st.success(
                        "🎉 Todos os alunos desta turma possuem documentos e cadastros 100% completos!"
                    )

    # ==============================================================================
    # --- ABA 3: PRESTAÇÃO PEDAGÓGICA (MOTOR NATIVO .DOCX) ---
    # ==============================================================================
    with tab_w:
        st.markdown("### 🏆 Prestação de Conta Pedagógica")
        st.info(
            "Geração do documento oficial consolidando métricas bioindicadoras e de engajamento baseadas nos dados preenchidos no Diário de Bordo."
        )
        if not DOCX_DISPONIVEL:
            st.error(
                "⚠️ Biblioteca 'python-docx' não instalada no servidor. O relatório nativo não pode ser gerado."
            )
            return

        cw1, cw2, cw3 = st.columns(3)
        meses_dict = {
            1: "Janeiro",
            2: "Fevereiro",
            3: "Março",
            4: "Abril",
            5: "Maio",
            6: "Junho",
            7: "Julho",
            8: "Agosto",
            9: "Setembro",
            10: "Outubro",
            11: "Novembro",
            12: "Dezembro",
        }
        hoje_mes = datetime.date.today().month
        m_w_nome = cw1.selectbox(
            "Mês de Referência",
            list(meses_dict.values()),
            index=hoje_mes - 1 if hoje_mes - 1 < 12 else 0,
        )
        m_w_num = [k for k, v in meses_dict.items() if v == m_w_nome][0]
        a_w = cw2.selectbox("Ano de Referência", [2025, 2026, 2027], index=1)
        t_w_sel = cw3.selectbox(
            "Escopo do Relatório",
            ["Global Polo"] + turmas["nome"].tolist()
            if not turmas.empty
            else ["Global Polo"],
        )

        if st.button(
            "🚀 GERAR ARQUIVO WORD (.docx)", type="primary", use_container_width=True
        ):
            with st.spinner(
                "Compilando prontuários e diários com Motor Nativo Word..."
            ):
                data_ini_w = datetime.date(a_w, m_w_num, 1)
                if m_w_num == 12:
                    data_fim_w = datetime.date(a_w, 12, 31)
                else:
                    data_fim_w = datetime.date(
                        a_w, m_w_num + 1, 1
                    ) - datetime.timedelta(days=1)

                turma_query = "" if t_w_sel == "Global Polo" else t_w_sel
                if t_w_sel == "Global Polo":
                    df_alunos_w = buscar_alunos_geral()
                    if not df_alunos_w.empty:
                        df_alunos_w = df_alunos_w[df_alunos_w["status"] != "Inativo"]
                else:
                    df_alunos_w = get_alunos_por_turma(t_w_sel)

                total_alunos_w = len(df_alunos_w) if not df_alunos_w.empty else 0
                df_freq_w = get_relatorio_periodo(data_ini_w, data_fim_w, turma_query)

                # A nova função já nos dá as aulas úteis perfeitamente no DataFrame
                if not df_freq_w.empty:
                    _META_W = {"Nome", "Turma", "Total Aulas",
                               "Total P", "Total F", "Total J", "% Presença"}
                    cols_data = [c for c in df_freq_w.columns if c not in _META_W]
                    total_aulas_w = len(cols_data)
                else:
                    total_aulas_w = 0

                if not df_freq_w.empty and total_aulas_w > 0:
                    presencas_totais_w = df_freq_w["Total P"].sum()
                    possiveis = total_alunos_w * total_aulas_w
                    assiduidade_w = (
                        (presencas_totais_w / possiveis * 100) if possiveis > 0 else 0.0
                    )
                else:
                    assiduidade_w = 0.0

                eng_w = {
                    "total_alunos": total_alunos_w,
                    "total_aulas": total_aulas_w,
                    "assiduidade": assiduidade_w,
                }
                diarios_w = get_diarios_periodo(data_ini_w, data_fim_w, turma_query)
                if isinstance(diarios_w, pd.DataFrame):
                    diarios_w = diarios_w.to_dict("records")

                total_avals, soma_dor = 0, 0
                borg_list, bristol_list, urina_list = [], [], []

                if not df_alunos_w.empty:
                    for _, aluno in df_alunos_w.iterrows():
                        avals = get_avaliacoes_aluno(aluno["id"])
                        avals_lista = (
                            avals.to_dict("records")
                            if isinstance(avals, pd.DataFrame)
                            else (avals or [])
                        )
                        for av in avals_lista:
                            try:
                                dt_av = pd.to_datetime(av["data_avaliacao"]).date()
                                if data_ini_w <= dt_av <= data_fim_w:
                                    total_avals += 1
                                    soma_dor += float(
                                        av.get("nivel_dor", av.get("dor_nivel", 0))
                                    )
                                    if av.get("borg"):
                                        borg_list.append(av.get("borg"))
                                    if av.get("bristol"):
                                        bristol_list.append(av.get("bristol"))
                                    if av.get("urina"):
                                        urina_list.append(av.get("urina"))
                            except:
                                pass

                def get_moda(lista, default="Dados Insuficientes"):
                    return max(set(lista), key=lista.count) if lista else default

                clin_w = {
                    "total_avaliacoes": total_avals,
                    "dor": (soma_dor / total_avals) if total_avals > 0 else 0.0,
                    "borg": get_moda(borg_list),
                    "bristol": get_moda(bristol_list),
                    "urina": get_moda(urina_list),
                }

                doc_word = gerar_word_prestacao_contas(
                    t_w_sel, m_w_nome, a_w, eng_w, diarios_w, clin_w
                )
                if doc_word:
                    st.success("✅ Relatório Clínico/Pedagógico compilado com sucesso!")
                    st.download_button(
                        "📥 BAIXAR RELATÓRIO EXECUTIVO (.docx)",
                        doc_word,
                        f"Prestacao_Contas_{m_w_nome}_{a_w}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

    # ==============================================================================
    # --- ABA 5: PRESTAÇÃO DIÁRIA (Lista de Presença por Dia - PDF) ---
    # ==============================================================================
    with tab_diario:
        _renderizar_aba_prestacao_diaria()

    # ==============================================================================
    # --- ABA 6: AVALIAÇÕES PENDENTES ---
    # ==============================================================================
    with tab_sem_av:
        from views.sem_avaliacao_view import renderizar_aba_sem_avaliacao
        renderizar_aba_sem_avaliacao()

    # ==============================================================================
    # --- ABA 7: MONITORAMENTO CLÍNICO (B.I. DA SAÚDE) ---
    # ==============================================================================
    with tab_clinico:
        _renderizar_monitoramento_clinico()

    st.markdown(
        "<br><p style='text-align:center; color:#94a3b8; font-size:10px;'>Moveright™ Gestão Inteligente - Projeto Esporte e Saúde Community Phase 2 - v8.40 PRIMEMAX</p>",
        unsafe_allow_html=True,
    )
