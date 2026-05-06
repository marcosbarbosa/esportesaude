# ==============================================================================
# 📄 Arquivo: views/ficha_aluno_view.py
# 🏷️ VERSÃO: 14.1 (Atualização UI - Documento Único Dinâmico RG/CPF)
# 👤 AUTOR: Marcos Barbosa - MoveRight (c)
# ⚙️ FUNÇÃO: Busca, Filtros em Lote e Geração de Fichas de Matrícula (PDF).
# ==============================================================================

import streamlit as st
import datetime
import base64
import io
from database import buscar_alunos_geral
from utils.texto import normalizar_fonetica
from utils.imagem import get_base64_image

try:
    from st_keyup import st_keyup
    HAS_KEYUP = True
except ImportError:
    HAS_KEYUP = False

try:
    import qrcode
    HAS_QRCODE = True
except ImportError:
    HAS_QRCODE = False

try:
    from xhtml2pdf import pisa
    XHTML_DISPONIVEL = True
except ImportError:
    XHTML_DISPONIVEL = False

try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt as _Pt, RGBColor as _RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    DOCX_FICHA_OK = True
except Exception:
    DOCX_FICHA_OK = False


def gerar_qr_code_b64(texto_auditoria):
    if not HAS_QRCODE:
        return None
    try:
        qr = qrcode.QRCode(version=1, box_size=4, border=0)
        qr.add_data(texto_auditoria)
        qr.make(fit=True)
        img = qr.make_image(fill_color="#0A2540", back_color="white")
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        b64_string = base64.b64encode(buffer.getvalue()).decode()
        return f"data:image/png;base64,{b64_string}"
    except Exception:
        return None


# ==============================================================================
# 🚀 MOTOR UNIFICADO MULTI-PAGE (HTML PREVIEW)
# ==============================================================================
def gerar_html_fichas(lista_alunos, host_url):
    from utils.identidade import get_config as _get_cfg_ficha, get_logo_data_url as _gld
    _cfg = _get_cfg_ficha()
    _logo_p = _gld(_cfg.get("logo_principal", "logo-imbra.png"))
    _logo_s = _gld(_cfg.get("logo_secundaria", "logo-secretaria.png"))
    html_logo = (
        f'<img src="{_logo_p}" style="max-width: 120px; max-height: 80px; width: auto; height: auto;" alt="{_cfg.get("nome_organizacao","")}">'
        if _logo_p else "INSERIR LOGO<br>INSTITUTO"
    )
    html_logo_sec = (
        f'<img src="{_logo_s}" style="max-width: 120px; max-height: 80px; width: auto; height: auto;" alt="Parceiro Institucional">'
        if _logo_s else "INSERIR LOGO<br>SECRETARIA"
    )

    data_hoje = datetime.date.today().strftime("%d/%m/%Y")

    html_base = f"""
    <!DOCTYPE html>
    <html lang="pt-br">
    <head>
        <meta charset="UTF-8">
        <title>Fichas de Matrícula</title>
        <style>
            @page {{ size: A4; margin: 1.2cm; }}
            * {{ box-sizing: border-box; -webkit-print-color-adjust: exact; }}
            body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 0; padding: 0; color: #2c3e50; line-height: 1.4; font-size: 10pt; background: #e2e8f0; }}

            .page-container {{ width: 100%; max-width: 210mm; margin: 20px auto; padding: 10px; background: white; box-shadow: 0 4px 10px rgba(0,0,0,0.1); }}

            /* 🚀 MAGIA DA QUEBRA DE PÁGINA (BULK PRINT) */
            .page-break {{ page-break-after: always; }}

            header {{ display: flex; align-items: center; justify-content: space-between; border-bottom: 2px solid #ff6b35; padding-bottom: 15px; margin-bottom: 20px; }}
            .logo-placeholder {{ width: 130px; height: 80px; background: transparent; display: flex; align-items: center; justify-content: center; font-size: 8pt; text-align: center; color: #64748b; }}
            .header-title {{ text-align: center; flex-grow: 1; padding: 0 5px; }}

            .header-title h1 {{ margin: 0; font-size: 12.5pt; color: #0a2540; text-transform: uppercase; font-weight: 900; white-space: nowrap; letter-spacing: -0.5px; }}
            .header-title h2 {{ margin: 5px 0 0 0; font-size: 10.5pt; font-weight: bold; color: #475569; }}

            .section {{ margin-bottom: 15px; }}
            .section-title {{ background: #f8fafc; padding: 6px 12px; font-weight: bold; text-transform: uppercase; font-size: 9pt; border-left: 4px solid #0056b3; margin-bottom: 10px; color: #0a2540; }}
            .data-grid {{ display: grid; grid-template-columns: repeat(2, 1fr); gap: 8px; }}
            .data-item {{ border-bottom: 1px solid #f1f5f9; padding: 4px 0; }}
            .label {{ font-weight: bold; font-size: 8pt; color: #64748b; display: block; text-transform: uppercase; }}
            .value {{ font-size: 10pt; color: #0f172a; text-transform: uppercase; font-weight: 600; }}
            .value-lower {{ font-size: 10pt; color: #0f172a; font-weight: 600; }}
            .health-alert {{ background: #fef2f2; border: 1px solid #fecaca; border-radius: 6px; padding: 12px; }}
            .legal-term {{ font-size: 8.5pt; text-align: justify; background: #f8fafc; border: 1px solid #e2e8f0; padding: 12px; font-style: italic; margin: 20px 0; color: #475569; }}
            .signature-area {{ margin-top: 50px; display: flex; justify-content: space-between; width: 100%; }}
            .sig-box {{ width: 45%; text-align: center; }}
            .sig-line {{ border-top: 1px solid #0a2540; margin-bottom: 8px; }}
            .sig-label {{ font-size: 8pt; font-weight: bold; color: #0a2540; }}
            .footer-info {{ margin-top: 30px; display: flex; justify-content: space-between; align-items: flex-end; font-size: 7.5pt; color: #94a3b8; border-top: 1px solid #e2e8f0; padding-top: 10px; }}
            .qr-placeholder {{ width: 60px; height: 60px; border: 1px solid #cbd5e1; display: flex; align-items: center; justify-content: center; font-size: 6pt; text-align: center; background: #f8fafc; }}

            .btn-imprimir {{ position: fixed; top: 20px; right: 20px; background: #0056b3; color: white; border: none; padding: 12px 24px; border-radius: 8px; font-size: 14px; font-weight: bold; cursor: pointer; box-shadow: 0 4px 15px rgba(0,86,179,0.3); transition: 0.3s; z-index: 1000; display: flex; align-items: center; gap: 8px; }}
            .btn-imprimir:hover {{ background: #004494; transform: translateY(-2px); }}

            @media print {{ 
                body {{ background: white; margin: 0; padding: 0; }} 
                .page-container {{ border: none; width: 100%; margin: 0; padding: 0; box-shadow: none; max-width: none; }} 
                .no-print {{ display: none !important; }} 
            }}
        </style>
    </head>
    <body>
        <button class="btn-imprimir no-print" onclick="window.print()">🖨️ IMPRIMIR TODAS AS FICHAS (A4)</button>
    """

    conteudo_paginas = ""
    total_alunos = len(lista_alunos)

    for index, aluno in enumerate(lista_alunos):
        nome = aluno.get("nome", "NÃO INFORMADO")
        nascimento = aluno.get("data_nascimento", "NÃO INFORMADO")
        telefone = aluno.get("whatsapp", "NÃO INFORMADO")

        email_raw = aluno.get("email", "NÃO INFORMADO")
        email = (
            email_raw.lower()
            if isinstance(email_raw, str)
            and email_raw.strip() != ""
            and email_raw != "NÃO INFORMADO"
            else "NÃO INFORMADO"
        )
        endereco = aluno.get("endereco", "NÃO INFORMADO")
        turma = aluno.get("turma", "Turma Padrão")
        id_aluno = aluno.get("id", "0000")

        # 🚀 LÓGICA DE PRIORIDADE RG -> CPF
        rg_raw = str(aluno.get("rg", ""))
        cpf_raw = str(aluno.get("cpf", ""))

        def is_clean(v):
            return v.strip() and v.strip().upper() not in ["NÃO INFORMADO", "NAN", "NONE"]

        if is_clean(rg_raw):
            doc_label, doc_val = "RG", rg_raw.strip()
        elif is_clean(cpf_raw):
            doc_label, doc_val = "CPF", cpf_raw.strip()
        else:
            doc_label, doc_val = "Documento Oficial", "NÃO INFORMADO"

        url_validacao = f"https://{host_url}/?rota=validar&id={id_aluno}"
        qr_b64 = gerar_qr_code_b64(url_validacao)
        html_qr_code = (
            f'<a href="{url_validacao}" target="_blank"><img src="{qr_b64}" style="width: 60px; height: 60px;"></a>'
            if qr_b64
            else '<div class="qr-placeholder">QR CODE</div>'
        )

        pagina = f"""
        <div class="page-container">
            <header>
                <div class="logo-placeholder">{html_logo_sec}</div>
                <div class="header-title">
                    <h1>{_cfg.get("titulo_projeto","ESPORTE E SAÚDE NA COMUNIDADE - FASE 2")}</h1>
                    <h2>FICHA DE MATRÍCULA E TERMO DE ADESÃO</h2>
                    <div style="font-size: 8.5pt; margin-top: 8px; color: #64748b;">
                        <strong>ID Protocolo:</strong> #{id_aluno} &nbsp;|&nbsp; 
                        <strong>Emissão:</strong> {data_hoje}
                    </div>
                </div>
                <div class="logo-placeholder">{html_logo}</div>
            </header>

            <div class="section">
                <div class="section-title">Informações Cadastrais do Aluno</div>
                <div class="data-grid">
                    <div class="data-item" style="grid-column: span 2;">
                        <span class="label">Nome Completo</span>
                        <span class="value">{nome}</span>
                    </div>
                    <div class="data-item">
                        <span class="label">{doc_label}</span>
                        <span class="value">{doc_val}</span>
                    </div>
                    <div class="data-item">
                        <span class="label">Data de Nascimento</span>
                        <span class="value">{nascimento}</span>
                    </div>
                    <div class="data-item">
                        <span class="label">Turma / Modalidade de Inscrição</span>
                        <span class="value">{turma}</span>
                    </div>
                    <div class="data-item">
                        <span class="label">Telefone / WhatsApp</span>
                        <span class="value">{telefone}</span>
                    </div>
                    <div class="data-item" style="grid-column: span 2;">
                        <span class="label">Endereço Residencial</span>
                        <span class="value">{endereco}</span>
                    </div>
                    <div class="data-item" style="grid-column: span 2;">
                        <span class="label">E-mail de Contato</span>
                        <span class="value-lower">{email}</span>
                    </div>
                </div>
            </div>

            <div class="section">
                <div class="section-title">Saúde e Segurança de Risco</div>
                <div class="health-alert">
                    <div class="data-grid">
                        <div class="data-item">
                            <span class="label" style="color: #991b1b;">Liberação Médica para Atividades</span>
                            <span class="value" style="color: #7f1d1d;">Sim, Apto para atividades físicas</span>
                        </div>
                        <div class="data-item">
                            <span class="label" style="color: #991b1b;">Uso de Medicamentos Contínuos</span>
                            <span class="value" style="color: #7f1d1d;">Não Informado / Conforme Atestado</span>
                        </div>
                        <div class="data-item" style="grid-column: span 2;">
                            <span class="label" style="color: #991b1b;">Alergias ou Lesões Prévias Graves</span>
                            <span class="value" style="color: #7f1d1d;">Nenhuma restrição grave informada via sistema</span>
                        </div>
                    </div>
                </div>
            </div>

            <div class="legal-term">
                Eu, <strong>{nome}</strong>, declaro estar ciente e concordo com minha inscrição no {_cfg.get("nome_organizacao","Instituto Muda Brasil")} para o projeto {_cfg.get("titulo_projeto","ESPORTE E SAÚDE NA COMUNIDADE - FASE 2")}. Declaro que as informações de saúde acima são verídicas e condizem com meu atestado físico. Autorizo expressamente o tratamento de meus dados pessoais, <strong>bem como o uso gratuito da minha imagem e voz captadas durante as atividades</strong>, para fins de gestão acadêmica, prestação de contas oficial, divulgação institucional e registros de frequência, em total conformidade com a Lei Geral de Proteção de Dados (Lei 13.709/2018 - LGPD) e a legislação vigente sobre direitos de imagem.
            </div>

            <div class="signature-area">
                <div class="sig-box">
                    <div class="sig-line"></div>
                    <div class="sig-label">ASSINATURA DO ALUNO (OU RESPONSÁVEL LEGAL)</div>
                    <div style="font-size: 7.5pt; color: #64748b; margin-top: 4px;">Assinado fisicamente no dia: ___/___/202___</div>
                </div>
                <div class="sig-box">
                    <div class="sig-line"></div>
                    <div class="sig-label">INSTITUTO MUDA BRASIL / SECRETARIA</div>
                    <div style="font-size: 7.5pt; color: #64748b; margin-top: 4px;">Rubrica do Coordenador</div>
                </div>
            </div>

            <div class="footer-info">
                <div>
                    <strong>Sistema Esporte e Saúde - Gestão Inteligente Moveright™</strong><br>
                    {_cfg.get("nome_organizacao","Instituto Muda Brasil")} | CNPJ: {_cfg.get("cnpj","08.817.519/0001-79")} | {_cfg.get("site","imbra.org.br")}
                </div>
                <div style="text-align: right; display: flex; align-items: center; gap: 10px;">
                    <span>Escaneie para Auditoria:<br>Documento Oficial Registrado.</span>
                    {html_qr_code}
                </div>
            </div>
        </div>
        """

        conteudo_paginas += pagina
        if index < total_alunos - 1:
            conteudo_paginas += '<div class="page-break"></div>'

    html_final = html_base + conteudo_paginas + "</body></html>"
    return html_final


# ==============================================================================
# 🖨️ GERADOR PDF NATIVO (xhtml2pdf)
# ==============================================================================
def _html_pagina_ficha_pdf(aluno, cfg, host_url, data_hoje, logo_p_url, logo_s_url):
    nome = aluno.get("nome", "NÃO INFORMADO")
    nascimento = aluno.get("data_nascimento", "NÃO INFORMADO")
    telefone = aluno.get("whatsapp", "NÃO INFORMADO")
    email_raw = aluno.get("email", "NÃO INFORMADO")
    email = (
        email_raw.lower()
        if isinstance(email_raw, str) and email_raw.strip() and email_raw != "NÃO INFORMADO"
        else "NÃO INFORMADO"
    )
    endereco = aluno.get("endereco", "NÃO INFORMADO")
    turma = aluno.get("turma", "Turma Padrão")
    id_aluno = aluno.get("id", "0000")

    # 🚀 LÓGICA DE PRIORIDADE RG -> CPF APLICADA AQUI TAMBÉM
    rg_raw = str(aluno.get("rg", ""))
    cpf_raw = str(aluno.get("cpf", ""))

    def is_clean(v):
        return v.strip() and v.strip().upper() not in ["NÃO INFORMADO", "NAN", "NONE"]

    if is_clean(rg_raw):
        doc_label, doc_val = "RG", rg_raw.strip()
    elif is_clean(cpf_raw):
        doc_label, doc_val = "CPF", cpf_raw.strip()
    else:
        doc_label, doc_val = "Documento Oficial", "NÃO INFORMADO"

    titulo_proj = cfg.get("titulo_projeto", "ESPORTE E SAÚDE NA COMUNIDADE - FASE 2")
    nome_org = cfg.get("nome_organizacao", "Instituto Muda Brasil")
    cnpj = cfg.get("cnpj", "08.817.519/0001-79")
    site = cfg.get("site", "imbra.org.br")

    url_val = f"https://{host_url}/?rota=validar&id={id_aluno}"
    qr_b64 = gerar_qr_code_b64(url_val)
    qr_img = (
        f'<img src="{qr_b64}" width="55" height="55" />'
        if qr_b64
        else "<span style='font-size:7pt;'>QR CODE</span>"
    )

    logo_p_html = (
        f'<img src="{logo_p_url}" style="width:100px;height:auto;max-height:65px;" />'
        if logo_p_url
        else ""
    )
    logo_s_html = (
        f'<img src="{logo_s_url}" style="width:100px;height:auto;max-height:65px;" />'
        if logo_s_url
        else ""
    )

    # Note que a estrutura em tabelas (<tr>) foi ajustada para o reequilíbrio:
    # Row 1: Nome. Row 2: Documento | Nascimento. Row 3: Turma | Telefone. Row 4: Endereço. Row 5: E-mail.
    return f"""
<table width="100%" border="0" cellspacing="0" cellpadding="4"
       style="border-bottom:2px solid #ff6b35;margin-bottom:10px;">
  <tr>
    <td width="18%" valign="middle" align="left">{logo_s_html}</td>
    <td width="64%" valign="middle" align="center">
      <b style="font-size:12pt;color:#0a2540;text-transform:uppercase;">{titulo_proj}</b><br/>
      <span style="font-size:10pt;color:#475569;font-weight:bold;">FICHA DE MATRÍCULA E TERMO DE ADESÃO</span><br/>
      <span style="font-size:8pt;color:#64748b;"><b>ID Protocolo:</b> #{id_aluno} &nbsp;|&nbsp; <b>Emissão:</b> {data_hoje}</span>
    </td>
    <td width="18%" valign="middle" align="right">{logo_p_html}</td>
  </tr>
</table>

<p style="background:#f8fafc;padding:5px 10px;font-weight:bold;font-size:9pt;
          border-left:4px solid #0056b3;color:#0a2540;text-transform:uppercase;margin-bottom:6px;">
  Informações Cadastrais do Aluno
</p>
<table width="100%" border="0" cellspacing="0" cellpadding="4" style="margin-bottom:10px;">
  <tr>
    <td colspan="2" style="border-bottom:1px solid #f1f5f9;padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">Nome Completo</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;text-transform:uppercase;">{nome}</span>
    </td>
  </tr>
  <tr>
    <td width="50%" style="border-bottom:1px solid #f1f5f9;padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">{doc_label}</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;">{doc_val}</span>
    </td>
    <td width="50%" style="border-bottom:1px solid #f1f5f9;padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">Data de Nascimento</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;">{nascimento}</span>
    </td>
  </tr>
  <tr>
    <td width="50%" style="border-bottom:1px solid #f1f5f9;padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">Turma / Modalidade</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;">{turma}</span>
    </td>
    <td width="50%" style="border-bottom:1px solid #f1f5f9;padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">Telefone / WhatsApp</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;">{telefone}</span>
    </td>
  </tr>
  <tr>
    <td colspan="2" style="border-bottom:1px solid #f1f5f9;padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">Endereço Residencial</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;text-transform:uppercase;">{endereco}</span>
    </td>
  </tr>
  <tr>
    <td colspan="2" style="padding:3px 2px;">
      <span style="font-size:8pt;font-weight:bold;color:#64748b;text-transform:uppercase;">E-mail de Contato</span><br/>
      <span style="font-size:10pt;color:#0f172a;font-weight:bold;">{email}</span>
    </td>
  </tr>
</table>

<p style="background:#f8fafc;padding:5px 10px;font-weight:bold;font-size:9pt;
          border-left:4px solid #0056b3;color:#0a2540;text-transform:uppercase;margin-bottom:6px;">
  Saúde e Segurança de Risco
</p>
<table width="100%" border="0" cellspacing="0" cellpadding="6"
       style="background:#fef2f2;border:1px solid #fecaca;margin-bottom:10px;">
  <tr>
    <td width="50%">
      <span style="font-size:8pt;font-weight:bold;color:#991b1b;text-transform:uppercase;">Liberação Médica para Atividades</span><br/>
      <span style="font-size:10pt;color:#7f1d1d;font-weight:bold;">Sim, Apto para atividades físicas</span>
    </td>
    <td width="50%">
      <span style="font-size:8pt;font-weight:bold;color:#991b1b;text-transform:uppercase;">Uso de Medicamentos Contínuos</span><br/>
      <span style="font-size:10pt;color:#7f1d1d;font-weight:bold;">Não Informado / Conforme Atestado</span>
    </td>
  </tr>
  <tr>
    <td colspan="2">
      <span style="font-size:8pt;font-weight:bold;color:#991b1b;text-transform:uppercase;">Alergias ou Lesões Prévias Graves</span><br/>
      <span style="font-size:10pt;color:#7f1d1d;font-weight:bold;">Nenhuma restrição grave informada via sistema</span>
    </td>
  </tr>
</table>

<p style="font-size:8.5pt;text-align:justify;background:#f8fafc;border:1px solid #e2e8f0;
          padding:10px;font-style:italic;margin:8px 0;color:#475569;">
  Eu, <b>{nome}</b>, declaro estar ciente e concordo com minha inscrição no {nome_org} para
  o projeto {titulo_proj}. Declaro que as informações de saúde acima são verídicas e condizem com
  meu atestado físico. Autorizo expressamente o tratamento de meus dados pessoais,
  <b>bem como o uso gratuito da minha imagem e voz captadas durante as atividades</b>,
  para fins de gestão acadêmica, prestação de contas oficial, divulgação institucional e
  registros de frequência, em total conformidade com a Lei Geral de Proteção de Dados
  (Lei 13.709/2018 - LGPD) e a legislação vigente sobre direitos de imagem.
</p>

<table width="100%" border="0" cellspacing="0" cellpadding="0" style="margin-top:28px;">
  <tr>
    <td width="45%" align="center" style="border-top:1px solid #0a2540;padding-top:6px;">
      <b style="font-size:8pt;">ASSINATURA DO ALUNO (OU RESPONSÁVEL LEGAL)</b><br/>
      <span style="font-size:7.5pt;color:#64748b;">Assinado fisicamente no dia: ___/___/202___</span>
    </td>
    <td width="10%">&nbsp;</td>
    <td width="45%" align="center" style="border-top:1px solid #0a2540;padding-top:6px;">
      <b style="font-size:8pt;">{nome_org.upper()} / SECRETARIA</b><br/>
      <span style="font-size:7.5pt;color:#64748b;">Rubrica do Coordenador</span>
    </td>
  </tr>
</table>

<table width="100%" border="0" cellspacing="0" cellpadding="2"
       style="margin-top:12px;border-top:1px solid #e2e8f0;">
  <tr>
    <td valign="bottom">
      <span style="font-size:7.5pt;color:#94a3b8;">
        <b>Sistema Esporte e Saúde - Gestão Inteligente Moveright™</b><br/>
        {nome_org} | CNPJ: {cnpj} | {site}
      </span>
    </td>
    <td align="right" valign="bottom">
      {qr_img}<br/>
      <span style="font-size:6.5pt;color:#94a3b8;">Escaneie para Auditoria</span>
    </td>
  </tr>
</table>
"""


def gerar_pdf_fichas_bytes(lista_alunos, host_url):
    """Gera um PDF multi-página com todas as fichas via xhtml2pdf."""
    if not XHTML_DISPONIVEL:
        return None

    from utils.identidade import get_config as _gcfg_f, get_logo_data_url as _gld_f
    cfg = _gcfg_f()
    logo_p_url = _gld_f(cfg.get("logo_principal", "logo-imbra.png"))
    logo_s_url = _gld_f(cfg.get("logo_secundaria", "logo-secretaria.png"))
    data_hoje = datetime.date.today().strftime("%d/%m/%Y")

    paginas_html = []
    for aluno in lista_alunos:
        paginas_html.append(
            _html_pagina_ficha_pdf(aluno, cfg, host_url, data_hoje, logo_p_url, logo_s_url)
        )

    conteudo = '<div style="page-break-after:always;"></div>'.join(paginas_html)

    html_completo = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"/>
<style>
  @page {{ size: A4 portrait; margin: 1.2cm; }}
  body {{ font-family: Helvetica, Arial, sans-serif; margin:0; padding:0;
         color:#2c3e50; font-size:10pt; }}
  table {{ border-collapse:collapse; }}
</style>
</head><body>
{conteudo}
</body></html>"""

    resultado = io.BytesIO()
    pisa.pisaDocument(io.StringIO(html_completo), resultado)
    return resultado.getvalue() if resultado.getvalue() else None


# ==============================================================================
# 📝 GERADOR WORD — FICHA DE MATRÍCULA INDIVIDUAL
# ==============================================================================
def gerar_word_ficha_bytes(aluno):
    """Gera um .docx formatado com os dados da ficha de matrícula do aluno."""
    if not DOCX_FICHA_OK:
        return None
    try:
        from utils.identidade import get_config as _gcfg
        cfg = _gcfg()
        nome_org   = cfg.get("nome_organizacao", "Instituto Muda Brasil")
        titulo_proj = cfg.get("titulo_projeto", "ESPORTE E SAÚDE NA COMUNIDADE - FASE 2")
        cnpj       = cfg.get("cnpj", "08.817.519/0001-79")
        site       = cfg.get("site", "imbra.org.br")

        doc = _DocxDocument()

        for sec in doc.sections:
            sec.top_margin    = _Pt(36)
            sec.bottom_margin = _Pt(36)
            sec.left_margin   = _Pt(50)
            sec.right_margin  = _Pt(50)

        h = doc.add_heading(titulo_proj, level=1)
        h.alignment = _WD_ALIGN.CENTER
        h.runs[0].font.size = _Pt(13)
        h.runs[0].font.color.rgb = _RGBColor(0x0A, 0x25, 0x40)

        sub = doc.add_paragraph("FICHA DE MATRÍCULA E TERMO DE ADESÃO")
        sub.alignment = _WD_ALIGN.CENTER
        sub.runs[0].font.size = _Pt(11)
        sub.runs[0].font.bold = True

        data_hoje = datetime.date.today().strftime("%d/%m/%Y")
        sub2 = doc.add_paragraph(f"Emissão: {data_hoje}  |  ID: #{aluno.get('id','—')}")
        sub2.alignment = _WD_ALIGN.CENTER
        sub2.runs[0].font.size = _Pt(9)
        sub2.runs[0].font.color.rgb = _RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

        def _secao(titulo):
            p = doc.add_paragraph()
            run = p.add_run(f"  {titulo}  ")
            run.font.bold = True
            run.font.size = _Pt(10)
            run.font.color.rgb = _RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0056B3')
            pPr.append(shd)
            return p

        def _campo(label, valor):
            p = doc.add_paragraph()
            run_l = p.add_run(f"{label}: ")
            run_l.font.bold = True
            run_l.font.size = _Pt(10)
            run_v = p.add_run(str(valor) if valor else "Não informado")
            run_v.font.size = _Pt(10)

        # 🚀 LÓGICA DE PRIORIDADE RG -> CPF PARA O WORD
        rg_raw = str(aluno.get("rg", ""))
        cpf_raw = str(aluno.get("cpf", ""))

        def is_clean(v):
            return v.strip() and v.strip().upper() not in ["NÃO INFORMADO", "NAN", "NONE"]

        if is_clean(rg_raw):
            doc_label, doc_val = "RG", rg_raw.strip()
        elif is_clean(cpf_raw):
            doc_label, doc_val = "CPF", cpf_raw.strip()
        else:
            doc_label, doc_val = "Documento Oficial", "Não informado"

        _secao("1. IDENTIFICAÇÃO DO ALUNO")
        _campo("Nome Completo",        aluno.get("nome"))
        _campo("Data de Nascimento",   aluno.get("data_nascimento") or aluno.get("nascimento"))
        _campo(doc_label,              doc_val)
        _campo("Turma",                aluno.get("turma"))
        _campo("Status",               aluno.get("status"))
        doc.add_paragraph()

        _secao("2. CONTATO")
        _campo("WhatsApp / Telefone",  aluno.get("whatsapp") or aluno.get("telefone"))
        _campo("E-mail",               aluno.get("email"))
        _campo("Endereço",             aluno.get("endereco"))
        doc.add_paragraph()

        _secao("3. CONTATO DE EMERGÊNCIA")
        _campo("Nome do Responsável",  aluno.get("contato_emergencia_nome") or aluno.get("responsavel"))
        _campo("Telefone Emergência",  aluno.get("contato_emergencia_tel") or aluno.get("tel_emergencia"))
        doc.add_paragraph()

        _secao("4. DADOS DE SAÚDE")
        _campo("Condições / Restrições", aluno.get("condicoes_saude") or aluno.get("observacoes_saude"))
        _campo("Medicamentos em uso",    aluno.get("medicamentos"))
        _campo("Tipo Sanguíneo",         aluno.get("tipo_sanguineo"))
        doc.add_paragraph()

        _secao("5. TERMO DE ADESÃO")
        termo = doc.add_paragraph(
            "Declaro que as informações acima são verdadeiras e autorizo o uso dos dados para fins "
            "de gestão do programa social. Estou ciente das normas de participação e comprometo-me "
            "a cumpri-las. Autorizo, ainda, o uso de imagem e voz para divulgação institucional "
            "não comercial do programa."
        )
        termo.runs[0].font.size = _Pt(9)

        doc.add_paragraph()
        doc.add_paragraph()

        # Assinaturas
        ass = doc.add_paragraph()
        run_a = ass.add_run(f"{'_'*35}          {'_'*35}")
        run_a.font.size = _Pt(10)
        leg = doc.add_paragraph(f"Assinatura do Aluno / Responsável          Assinatura da Coordenação")
        leg.runs[0].font.size = _Pt(9)
        leg.alignment = _WD_ALIGN.LEFT

        # Rodapé no texto
        doc.add_paragraph()
        rod = doc.add_paragraph(f"{nome_org} | CNPJ: {cnpj} | {site}")
        rod.alignment = _WD_ALIGN.CENTER
        rod.runs[0].font.size = _Pt(8)
        rod.runs[0].font.color.rgb = _RGBColor(0x94, 0xA3, 0xB8)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        return None


# ==============================================================================
# INTERFACE PRINCIPAL (TABS)
# ==============================================================================
def tela_impressao_ficha():
    st.markdown("### 🖨️ Central de Impressão de Fichas (Prestação de Contas)")

    try:
        host = st.context.headers.get("host", "esportesaude.riker.replit.dev")
    except:
        host = "esportesaude.riker.replit.dev"

    df_alunos = buscar_alunos_geral("")
    if df_alunos.empty:
        st.warning("Nenhum aluno encontrado na base de dados.")
        return

    tab_individual, tab_lote = st.tabs(
        ["👤 Impressão Individual", "🗂️ Impressão em Lote (Bulk Print)"]
    )

    with tab_individual:
        st.write(
            "Digite o nome do aluno abaixo. A lista atualiza automaticamente conforme digita."
        )
        with st.container(border=True):
            st.markdown(
                "<div style='font-size: 14px; font-weight: 700; color: #0056b3; margin-bottom: 10px;'>🔍 BUSCAR ALUNO</div>",
                unsafe_allow_html=True,
            )
            if HAS_KEYUP:
                termo_busca = st_keyup(
                    "Buscar Individual",
                    label_visibility="collapsed",
                    placeholder="🔍 Digite pelo menos 3 letras...",
                )
            else:
                termo_busca = st.text_input(
                    "Buscar Individual",
                    label_visibility="collapsed",
                    placeholder="🔍 Digite pelo menos 3 letras e prima ENTER...",
                )

        if termo_busca and len(termo_busca) >= 3:
            termo_norm = normalizar_fonetica(termo_busca)
            df_alunos["nome_norm"] = df_alunos["nome"].apply(normalizar_fonetica)
            df_view = df_alunos[
                df_alunos["nome_norm"].str.contains(termo_norm, case=False, na=False)
            ].sort_values("nome")

            if df_view.empty:
                st.warning("⚠️ Nenhum aluno encontrado.")
            else:
                st.markdown(f"**{len(df_view)} aluno(s) encontrado(s):**")
                for idx, row in df_view.iterrows():
                    with st.container(border=True):
                        c_info, c_acao = st.columns([4, 1], vertical_alignment="center")
                        with c_info:
                            st.markdown(f"**👤 {row['nome']}**")
                            st.caption(
                                f"CPF: {row.get('cpf', 'N/A')} | Turma: {row.get('turma', 'N/A')}"
                            )
                        with c_acao:
                            if st.button(
                                "🖨️ Gerar Ficha",
                                key=f"btn_ficha_{row['id']}",
                                use_container_width=True,
                            ):
                                st.session_state["aluno_ficha_selecionado"] = (
                                    row.to_dict()
                                )

        if st.session_state.get("aluno_ficha_selecionado"):
            aluno_dados = st.session_state["aluno_ficha_selecionado"]
            nome_aluno = aluno_dados.get("nome", "aluno")
            id_aluno   = aluno_dados.get("id", "00")
            st.markdown(
                "<hr style='border-top: 2px dashed #E2E8F0; margin: 25px 0;'>",
                unsafe_allow_html=True,
            )

            st.success(f"✅ Ficha gerada para: **{nome_aluno}**")

            html_gerado = gerar_html_fichas([aluno_dados], host)

            col_pdf, col_word, col_html = st.columns(3)

            with col_pdf:
                if XHTML_DISPONIVEL:
                    with st.spinner("Gerando PDF…"):
                        pdf_bytes = gerar_pdf_fichas_bytes([aluno_dados], host)
                    if pdf_bytes:
                        st.download_button(
                            label="📄 Baixar PDF",
                            data=pdf_bytes,
                            file_name=f"Ficha_{id_aluno}.pdf",
                            mime="application/pdf",
                            type="primary",
                            use_container_width=True,
                        )
                    else:
                        st.error("Erro ao gerar PDF.")
                else:
                    st.warning("PDF indisponível.")

            with col_word:
                if DOCX_FICHA_OK:
                    word_bytes = gerar_word_ficha_bytes(aluno_dados)
                    if word_bytes:
                        st.download_button(
                            label="📝 Baixar Word",
                            data=word_bytes,
                            file_name=f"Ficha_{id_aluno}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    else:
                        st.warning("Erro ao gerar Word.")
                else:
                    st.warning("Word indisponível.")

            with col_html:
                st.download_button(
                    label="🌐 Baixar HTML",
                    data=html_gerado,
                    file_name=f"Ficha_{id_aluno}.html",
                    mime="text/html",
                    use_container_width=True,
                    help="Abra no navegador e use Ctrl+P para imprimir/salvar como PDF.",
                )

            with st.expander("👀 Pré-visualização da Ficha", expanded=True):
                st.info("💡 **Dica:** Use os botões acima para baixar. Para imprimir direto, baixe o HTML e pressione Ctrl+P no navegador.")
                st.components.v1.html(html_gerado, height=950, scrolling=True)

    with tab_lote:
        st.write(
            "Filtre os alunos por Turma ou Inicial do Nome para gerar um único documento contendo dezenas de fichas simultaneamente."
        )

        with st.container(border=True):
            st.markdown(
                "<div style='font-size: 14px; font-weight: 700; color: #0056b3; margin-bottom: 10px;'>🗂️ FILTROS DE LOTE</div>",
                unsafe_allow_html=True,
            )

            c_filtro1, c_filtro2 = st.columns(2)

            turmas_disp = sorted(df_alunos["turma"].dropna().unique().tolist())
            turma_selecionada = c_filtro1.selectbox(
                "Filtrar por Turma:", ["Todas as Turmas"] + turmas_disp
            )

            alfabeto = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
            letras_selecionadas = c_filtro2.multiselect(
                "Filtrar por Inicial do Nome (Letra):",
                alfabeto,
                placeholder="Ex: A, B, C...",
            )

            apenas_ativos = st.checkbox("Imprimir apenas alunos ATIVOS", value=True)

        df_lote = df_alunos.copy()

        if apenas_ativos:
            df_lote = df_lote[df_lote["status"] != "Inativo"]

        if turma_selecionada != "Todas as Turmas":
            df_lote = df_lote[df_lote["turma"] == turma_selecionada]

        if letras_selecionadas:
            padrao = "^" + "|".join(letras_selecionadas)
            df_lote = df_lote[
                df_lote["nome"].str.contains(padrao, case=False, na=False)
            ]

        df_lote = df_lote.sort_values(by="nome")
        total_lote = len(df_lote)

        st.info(f"📊 O seu filtro encontrou **{total_lote} aluno(s)** para impressão.")

        if total_lote > 0:
            if total_lote > 150:
                st.warning(
                    "⚠️ Um lote com mais de 150 alunos pode demorar alguns segundos a processar. Por favor aguarde ao clicar no botão."
                )

            if not XHTML_DISPONIVEL:
                st.error("⚠️ xhtml2pdf não disponível no servidor — PDF indisponível.")
            elif st.button(
                f"🚀 GERAR LOTE COM {total_lote} FICHAS (PDF)",
                use_container_width=True,
                type="primary",
            ):
                with st.spinner(
                    f"A compilar {total_lote} fichas em PDF. Aguarde…"
                ):
                    lista_lote = df_lote.to_dict("records")
                    pdf_lote = gerar_pdf_fichas_bytes(lista_lote, host)

                if pdf_lote:
                    st.session_state["pdf_lote_gerado"] = pdf_lote
                    st.session_state["nome_arquivo_lote"] = (
                        f"Lote_Fichas_{turma_selecionada.split()[0]}_{total_lote}_Alunos.pdf"
                    )
                    st.success("✅ Lote processado com sucesso!")
                else:
                    st.error("❌ Erro ao gerar PDF do lote.")

        if st.session_state.get("pdf_lote_gerado"):
            st.markdown(
                "<hr style='border-top: 2px dashed #E2E8F0; margin: 25px 0;'>",
                unsafe_allow_html=True,
            )
            c_lote_msg, c_lote_btn = st.columns([2, 1], vertical_alignment="center")
            c_lote_msg.success(
                "PDF pronto. Cada ficha ocupa uma página A4 — imprima directamente pelo leitor de PDF."
            )
            with c_lote_btn:
                st.download_button(
                    label="📄 BAIXAR FICHAS EM LOTE (PDF)",
                    data=st.session_state["pdf_lote_gerado"],
                    file_name=st.session_state["nome_arquivo_lote"],
                    mime="application/pdf",
                    type="primary",
                    use_container_width=True,
                )