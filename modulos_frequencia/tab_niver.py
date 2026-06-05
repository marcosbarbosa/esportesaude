# ==============================================================================
# 📄 ARQUIVO: modulos_frequencia/tab_niver.py
# 🏷️ VERSÃO: 14.0 — Automação WhatsApp + E-mail + Status Parabenizado
# 👤 COPYRIGHT: © 2026 MoveRight Gestão Inteligente • Instituto Muda Brasil
# ⚙️ FUNÇÃO: Portal de Aniversários, Disparo WhatsApp/E-mail e Cartazes.
# ==============================================================================

import streamlit as st
import pandas as pd
import datetime
import urllib.parse
import base64
import io
import os
import requests
from PIL import Image, ImageOps
from database import buscar_alunos_geral
from utils.texto import formatar_whatsapp_link
from utils.identidade import get_config as _get_id_cfg, get_logo_data_url as _get_logo_url
from utils.niver_automatico import (
    get_config_niver,
    is_parabenizado,
    get_parabenizados_dict,
    marcar_parabenizado,
    desmarcar_parabenizado,
    montar_link_whatsapp,
    montar_mensagem_niver,
    status_niver_por_delta,
    enviar_email_aniversariantes,
    disparar_zapi_aniversariantes,
)

try:
    from xhtml2pdf import pisa
    XHTML_DISPONIVEL = True
except ImportError:
    XHTML_DISPONIVEL = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    DOCX_DISPONIVEL = True
    _DOCX_ERR = None
except Exception as _e:
    DOCX_DISPONIVEL = False
    _DOCX_ERR = str(_e)


# ==============================================================================
# 🗜️ FUNÇÕES UTILITÁRIAS E TRATAMENTO DE IMAGEM
# ==============================================================================
def processar_imagem_para_redondo_b64(url, size=(300, 300)):
    if not url or pd.isna(url):
        return None
    try:
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            img = Image.open(io.BytesIO(response.content)).convert("RGBA")
            img = ImageOps.fit(img, size, centering=(0.5, 0.5))
            mask = Image.new("L", size, 0)
            from PIL import ImageDraw
            draw = ImageDraw.Draw(mask)
            draw.ellipse((0, 0) + size, fill=255)
            output = Image.new("RGBA", size, (255, 255, 255, 0))
            output.paste(img, (0, 0), mask)
            buffer = io.BytesIO()
            output.save(buffer, format="PNG")
            return base64.b64encode(buffer.getvalue()).decode("utf-8")
    except Exception:
        pass
    return None


def processar_imagem_para_redondo_word(url, size=(150, 150)):
    if not url or pd.isna(url):
        return None
    try:
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            img = Image.open(io.BytesIO(response.content)).convert("RGBA")
            img = ImageOps.fit(img, size, centering=(0.5, 0.5))
            mask = Image.new("L", size, 0)
            from PIL import ImageDraw
            draw = ImageDraw.Draw(mask)
            draw.ellipse((0, 0) + size, fill=255)
            output = Image.new("RGBA", size, (255, 255, 255, 0))
            output.paste(img, (0, 0), mask)
            background = Image.new("RGB", size, (255, 255, 255))
            background.paste(output, mask=output.split()[3])
            buffer = io.BytesIO()
            background.save(buffer, format="PNG")
            return buffer.getvalue()
    except Exception:
        pass
    return None


# ==============================================================================
# 📘 MOTOR 1: GERAÇÃO DO WORD NATIVO (.DOCX)
# ==============================================================================
def gerar_cartaz_word_core(df_mes, titulo, subtitulo="", mensagem_cartaz=""):
    if not DOCX_DISPONIVEL:
        raise RuntimeError(f"python-docx indisponível: {_DOCX_ERR}")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    _niver_cfg = _get_id_cfg()
    _niver_logo = _niver_cfg.get("logo_principal", "logo-imbra.png")
    if os.path.exists(_niver_logo):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(_niver_logo, width=Inches(1.5))
        except Exception:
            pass

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run(f"{titulo}\n")
    run_h.bold = True
    run_h.font.size = Pt(22)
    run_h.font.color.rgb = RGBColor(30, 136, 229)

    if subtitulo:
        run_sub = p_header.add_run(f"{subtitulo}\n")
        run_sub.bold = True
        run_sub.font.size = Pt(16)
        run_sub.font.color.rgb = RGBColor(100, 116, 139)

    if mensagem_cartaz.strip():
        p_msg = doc.add_paragraph()
        p_msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_m = p_msg.add_run(f'"{mensagem_cartaz.strip()}"\n')
        run_m.font.size = Pt(11)
        run_m.font.italic = True
        run_m.font.color.rgb = RGBColor(71, 85, 105)

    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    col_idx = 0
    for _, r in df_mes.iterrows():
        if col_idx == 0:
            row_cells = table.add_row().cells
        cell = row_cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        img_bytes = processar_imagem_para_redondo_word(r.get("url_foto"))
        if img_bytes:
            run_img = p.add_run()
            run_img.add_picture(io.BytesIO(img_bytes), width=Inches(1.2))
            p.add_run("\n")
        else:
            run_no_img = p.add_run("[ SEM FOTO ]\n")
            run_no_img.font.color.rgb = RGBColor(148, 163, 184)
            run_no_img.font.size = Pt(8)

        run_nome = p.add_run(f"{r['nome'].upper()}\n")
        run_nome.bold = True
        run_nome.font.size = Pt(11)
        run_data = p.add_run(f"{r['dia']:02d}/{r['mes']:02d}\n")
        run_data.font.color.rgb = RGBColor(220, 38, 38)
        run_data.bold = True
        col_idx = (col_idx + 1) % 3

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ==============================================================================
# 📕 MOTOR 2: GERAÇÃO DO CARTAZ ECO-PDF (LAYOUT GRIDO DE 2 COLUNAS)
# ==============================================================================
def gerar_cartaz_pdf_core(df_mes, titulo, subtitulo="", mensagem_cartaz=""):
    if not XHTML_DISPONIVEL:
        return None

    cfg = _get_id_cfg()
    logo_p_url = _get_logo_url(cfg.get("logo_principal", "logo-imbra.png"))
    logo_s_url = _get_logo_url(cfg.get("logo_secundaria", "logo-secretaria.png"))

    html_logo_p = f'<img src="{logo_p_url}" style="height: 60px; width: auto;">' if logo_p_url else ""
    html_logo_s = f'<img src="{logo_s_url}" style="height: 60px; width: auto;">' if logo_s_url else ""

    if not mensagem_cartaz.strip():
        mensagem_cartaz = "Celebrando os aniversariantes! Muita saúde e vida ativa para todos!"

    linhas_colunas_html = ""
    registros = df_mes.reset_index(drop=True)

    for i in range(0, len(registros), 2):
        aluno_esq = registros.iloc[i]
        nome_esq = str(aluno_esq["nome"]).upper().strip()
        dia_esq = f"{int(aluno_esq['dia']):02d}/{int(aluno_esq['mes']):02d}"
        b64_img_esq = processar_imagem_para_redondo_b64(aluno_esq.get("url_foto"))

        foto_html_esq = f'<img src="data:image/png;base64,{b64_img_esq}" class="foto-perfil">' if b64_img_esq else '<div class="no-foto"></div>'

        celula_esquerda = f"""
            <td class="celula-aluno">
                <table style="width: 100%; border: none;">
                    <tr>
                        <td style="width: 125px; border: none; text-align: center;">{foto_html_esq}</td>
                        <td style="border: none; text-align: left; padding-left: 10px;">
                            <div class="nome-aluno">{nome_esq}</div>
                            <div class="data-aluno">🎂 {dia_esq}</div>
                        </td>
                    </tr>
                </table>
            </td>
        """

        if i + 1 < len(registros):
            aluno_dir = registros.iloc[i + 1]
            nome_dir = str(aluno_dir["nome"]).upper().strip()
            dia_dir = f"{int(aluno_dir['dia']):02d}/{int(aluno_dir['mes']):02d}"
            b64_img_dir = processar_imagem_para_redondo_b64(aluno_dir.get("url_foto"))

            foto_html_dir = f'<img src="data:image/png;base64,{b64_img_dir}" class="foto-perfil">' if b64_img_dir else '<div class="no-foto"></div>'

            celula_direita = f"""
                <td class="celula-aluno">
                    <table style="width: 100%; border: none;">
                        <tr>
                            <td style="width: 125px; border: none; text-align: center;">{foto_html_dir}</td>
                            <td style="border: none; text-align: left; padding-left: 10px;">
                                <div class="nome-aluno">{nome_dir}</div>
                                <div class="data-aluno">🎂 {dia_dir}</div>
                            </td>
                        </tr>
                    </table>
                </td>
            """
        else:
            celula_direita = '<td class="celula-vazia"></td>'

        linhas_colunas_html += f"<tr>{celula_esquerda}{celula_direita}</tr>"

    html_content = f"""
    <html><head><meta charset="UTF-8"><style>
        @page {{ size: A4 portrait; margin: 1.2cm; }}
        body {{ font-family: Helvetica, Arial, sans-serif; color: #1E293B; text-align: center; }}
        .tb-header {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; border-bottom: 3px solid #1E88E5; padding-bottom: 12px; }}
        .tb-header td {{ border: none; vertical-align: middle; }}
        .header-center h1 {{ font-size: 24px; color: #1E88E5; margin: 0; font-weight: bold; }}
        .header-center h2 {{ font-size: 14px; color: #64748B; margin: 4px 0 0 0; font-weight: bold; }}
        .msg-box {{ background-color: #F0F9FF; border-left: 4px solid #1E88E5; padding: 12px; text-align: center; font-size: 12px; font-style: italic; color: #0369A1; margin-bottom: 25px; }}
        .grid-table {{ width: 100%; border-collapse: separate; border-spacing: 12px; }}
        .celula-aluno {{ width: 50%; background-color: #F8FAFC; border: 1px solid #E2E8F0; padding: 10px; vertical-align: middle; border-radius: 6px; }}
        .foto-perfil {{ width: 110px; height: 110px; border-radius: 55px; border: 3px solid #1E88E5; object-fit: cover; }}
        .no-foto {{ width: 110px; height: 110px; border-radius: 55px; border: 2px dashed #94A3B8; background: #F1F5F9; display: inline-block; }}
        .nome-aluno {{ font-size: 13px; font-weight: bold; color: #0F172A; line-height: 1.3; }}
        .data-aluno {{ font-size: 12px; font-weight: bold; color: #DC2626; margin-top: 4px; }}
        .celula-vazia {{ width: 50%; border: none; background: none; }}
        .footer-tag {{ text-align: center; font-size: 8px; color: #94A3B8; margin-top: 35px; border-top: 1px solid #E2E8F0; padding-top: 8px; }}
    </style></head><body>
        <table class="tb-header">
            <tr>
                <td style="width: 20%; text-align: left;">{html_logo_s}</td>
                <td style="width: 60%;" class="header-center">
                    <h1>{titulo}</h1>
                    {f'<h2>{subtitulo}</h2>' if subtitulo else ''}
                </td>
                <td style="width: 20%; text-align: right;">{html_logo_p}</td>
            </tr>
        </table>
        <div class="msg-box">"{mensagem_cartaz.strip()}"</div>
        <table class="grid-table">
            <tbody>{linhas_colunas_html}</tbody>
        </table>
        <div class="footer-tag">Moveright™ Sistema de Gestão Integrada • {cfg.get("nome_organizacao", "Instituto Muda Brasil")}</div>
    </body></html>
    """

    result = io.BytesIO()
    pisa.pisaDocument(io.StringIO(html_content), result)
    return result.getvalue()


# ==============================================================================
# 🚀 MOTOR DE DISPARO — executa automação se habilitada no horário configurado
# ==============================================================================
def _verificar_disparo_automatico(df_hoje):
    """Verifica se deve disparar e-mail/Z-API automaticamente pelo horário."""
    if df_hoje is None or len(df_hoje) == 0:
        return
    cfg = get_config_niver()
    agora = datetime.datetime.now()
    horario_str = cfg.get("zapi_horario", "08:00")

    try:
        h, m = [int(x) for x in horario_str.split(":")]
        horario_alvo = agora.replace(hour=h, minute=m, second=0, microsecond=0)
        diff_min = abs((agora - horario_alvo).total_seconds() / 60)
    except Exception:
        return

    chave_sess = f"niver_auto_disparado_{agora.date()}"
    if st.session_state.get(chave_sess):
        return
    if diff_min > 30:
        return

    # janela de ±30 min — dispara
    st.session_state[chave_sess] = True

    if cfg["email_habilitado"]:
        ok, msg = enviar_email_aniversariantes(df_hoje, cfg)
        if ok:
            st.toast(f"📧 E-mail de aniversários enviado!", icon="🎂")

    if cfg["zapi_habilitado"]:
        resultados = disparar_zapi_aniversariantes(df_hoje, cfg)
        enviados = sum(1 for r in resultados if r["sucesso"])
        if enviados:
            st.toast(f"📱 {enviados} mensagem(ns) enviada(s) via WhatsApp!", icon="🎉")


# ==============================================================================
# 🧭 INTERFACE PRINCIPAL DO DASHBOARD (STREAMLIT RENDER)
# ==============================================================================
def renderizar_aba_niver():
    df_alunos = buscar_alunos_geral("")
    if df_alunos.empty:
        st.warning("A base de alunos está vazia.")
        return

    df_alunos["data_nascimento"] = pd.to_datetime(df_alunos["data_nascimento"], errors="coerce")
    df_validos = df_alunos.dropna(subset=["data_nascimento"]).copy()
    hoje = datetime.date.today()
    df_validos["dia"] = df_validos["data_nascimento"].dt.day
    df_validos["mes"] = df_validos["data_nascimento"].dt.month

    meses = {
        1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
    }
    c_mes, _, _ = st.columns([4, 1, 1], vertical_alignment="bottom")
    meses_selecionados = c_mes.multiselect("Selecionar Mês(es):", list(meses.values()), default=[meses[hoje.month]])

    if not meses_selecionados:
        st.warning("⚠️ Selecione pelo menos um mês para gerar o relatório.")
        return

    meses_inv = {v: k for k, v in meses.items()}
    meses_nums = [meses_inv[m] for m in meses_selecionados]
    df_mes = df_validos[df_validos["mes"].isin(meses_nums)].sort_values(by=["mes", "dia"]).copy()

    # Filtrar aniversariantes dentro da janela configurada (hoje + N dias à frente)
    cfg_niver_pre = get_config_niver()
    _aviso_dias = int(cfg_niver_pre.get("aviso_dias", 0))

    datas_janela = set()
    for _d in range(_aviso_dias + 1):
        _dt = hoje + datetime.timedelta(days=_d)
        datas_janela.add((_dt.day, _dt.month))

    df_hoje = df_validos[
        df_validos.apply(lambda r: (int(r["dia"]), int(r["mes"])) in datas_janela, axis=1)
    ].copy()

    # Adiciona coluna com quantos dias faltam para cada aniversariante da janela
    def _dias_para(row):
        for _d in range(_aviso_dias + 1):
            _dt = hoje + datetime.timedelta(days=_d)
            if int(row["dia"]) == _dt.day and int(row["mes"]) == _dt.month:
                return _d
        return 0
    if not df_hoje.empty:
        df_hoje["dias_para_niver"] = df_hoje.apply(_dias_para, axis=1)
        df_hoje = df_hoje.sort_values("dias_para_niver")

    # ── DISPARO AUTOMÁTICO (verificação silenciosa — só para os de HOJE) ─────
    df_so_hoje = df_hoje[df_hoje.get("dias_para_niver", pd.Series(dtype=int)) == 0] if not df_hoje.empty else df_hoje
    if not df_so_hoje.empty:
        _verificar_disparo_automatico(df_so_hoje)

    if len(meses_selecionados) == 1:
        titulo_doc = f"ANIVERSARIANTES DE {meses_selecionados[0].upper()}"
        subtitulo_doc = ""
        nome_arq = meses_selecionados[0]
        nome_meses_tela = meses_selecionados[0]
    else:
        titulo_doc = "ANIVERSARIANTES"
        subtitulo_doc = f"{meses_selecionados[0].upper()} A {meses_selecionados[-1].upper()}"
        nome_arq = f"{meses_selecionados[0]}_A_{meses_selecionados[-1]}"
        nome_meses_tela = f"{meses_selecionados[0]} a {meses_selecionados[-1]}"

    # ── PAINEL DE DISPARO DO DIA ─────────────────────────────────────────────
    if not df_hoje.empty:
        cfg_niver = get_config_niver()
        _aviso_dias_cfg = int(cfg_niver.get("aviso_dias", 0))
        n_hoje = len(df_hoje)
        parab_dict = get_parabenizados_dict()
        n_parab = sum(1 for _, r in df_hoje.iterrows() if str(r.get("id", "")) in parab_dict)

        # Título dinâmico conforme janela configurada
        if _aviso_dias_cfg == 0:
            _titulo_painel = f"Hoje fazem aniversário: {n_hoje} aluno(s)"
        else:
            _n_exato_hoje = sum(1 for _, r in df_hoje.iterrows() if int(r.get("dias_para_niver", 0)) == 0)
            _n_proximos = n_hoje - _n_exato_hoje
            partes = []
            if _n_exato_hoje:
                partes.append(f"{_n_exato_hoje} hoje")
            if _n_proximos:
                partes.append(f"{_n_proximos} nos próximos {_aviso_dias_cfg} dia(s)")
            _titulo_painel = f"Aniversários na janela: {' · '.join(partes)}"

        with st.container(border=True):
            st.markdown(
                f"<div style='display:flex;align-items:center;gap:10px;'>"
                f"<span style='font-size:28px;'>🎂</span>"
                f"<div><strong style='font-size:15px;'>{_titulo_painel}</strong>"
                f"<br><span style='font-size:12px;color:#64748B;'>"
                f"{n_parab} parabenizado(s) · {n_hoje - n_parab} pendente(s)"
                f"{f' · aviso com até {_aviso_dias_cfg}d de antecedência' if _aviso_dias_cfg > 0 else ''}"
                f"</span></div>"
                f"</div>",
                unsafe_allow_html=True,
            )
            st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)

            col_wa, col_em, col_zapi = st.columns(3)

            with col_wa:
                if st.button(
                    "💬 Abrir WhatsApp p/ Todos",
                    use_container_width=True,
                    type="primary",
                    key="btn_wa_todos",
                    help="Abre uma aba do WhatsApp para cada aniversariante com a mensagem pronta",
                ):
                    links_gerados = []
                    sem_wap = []
                    for _, r in df_hoje.iterrows():
                        nome = str(r.get("nome", "")).strip()
                        wap = str(r.get("whatsapp", "") or "").strip()
                        # Texto idêntico ao painel admin conforme status.
                        # Futuros (status None) não recebem parabéns antecipado.
                        _status = status_niver_por_delta(int(r.get("dias_para_niver", 0) or 0))
                        if _status is None:
                            continue
                        msg = montar_mensagem_niver(_status, nome)
                        link = montar_link_whatsapp(wap, msg) if wap else None
                        if link:
                            links_gerados.append((nome, link))
                        elif wap == "":
                            sem_wap.append(nome)

                    if links_gerados:
                        # Abre todos os links via JS
                        js_links = "\n".join(
                            [f'window.open("{lnk}", "_blank");' for _, lnk in links_gerados]
                        )
                        st.markdown(
                            f"<script>{js_links}</script>",
                            unsafe_allow_html=True,
                        )
                        # Mostra botões individuais também
                        st.success(f"✅ {len(links_gerados)} link(s) aberto(s)! Clique abaixo se algum não abriu:")
                        for nome_al, lnk in links_gerados:
                            st.markdown(
                                f'<a href="{lnk}" target="_blank" style="display:inline-block;'
                                f'margin:3px 4px;background:#25D366;color:white;padding:5px 12px;'
                                f'border-radius:6px;text-decoration:none;font-size:12px;font-weight:700;">'
                                f'💬 {nome_al.split()[0]}</a>',
                                unsafe_allow_html=True,
                            )
                    if sem_wap:
                        st.warning(f"Sem WhatsApp: {', '.join(sem_wap)}")

            with col_em:
                if st.button(
                    "📧 Enviar E-mail Agora",
                    use_container_width=True,
                    key="btn_email_hoje",
                    help="Envia e-mail com a lista de hoje para os endereços configurados",
                ):
                    if not cfg_niver["email_habilitado"]:
                        st.warning("⚠️ E-mail não habilitado. Configure em ⚙️ Config → 🎂 Aniversários.")
                    else:
                        with st.spinner("Enviando e-mail..."):
                            ok, msg = enviar_email_aniversariantes(df_hoje, cfg_niver)
                        if ok:
                            st.success(f"✅ {msg}")
                        else:
                            st.error(f"❌ {msg}")

            with col_zapi:
                if st.button(
                    "📱 Disparar via Z-API",
                    use_container_width=True,
                    key="btn_zapi_hoje",
                    help="Envia mensagem automática via WhatsApp para todos os aniversariantes de hoje",
                ):
                    if not cfg_niver["zapi_habilitado"]:
                        st.warning("⚠️ Z-API não habilitada. Configure em ⚙️ Config → 🎂 Aniversários.")
                    else:
                        with st.spinner("Disparando via Z-API..."):
                            resultados = disparar_zapi_aniversariantes(df_hoje, cfg_niver)
                        ok_list = [r for r in resultados if r["sucesso"]]
                        err_list = [r for r in resultados if not r["sucesso"]]
                        if ok_list:
                            st.success(f"✅ {len(ok_list)} mensagem(ns) enviada(s)!")
                        if err_list:
                            for e in err_list:
                                st.warning(f"⚠️ {e['nome']}: {e['msg']}")

    st.markdown("<hr style='margin: 10px 0;'>", unsafe_allow_html=True)

    c_msg, c_botoes = st.columns([3, 1], vertical_alignment="bottom")
    with c_msg:
        st.markdown(f"**💌 Mensagem Temática ({nome_meses_tela})**")
        st.selectbox("Tom da mensagem:", ["🏃‍♀️ Energia & Movimento", "👨‍👩‍👧 Acolhedora", "✍️ Personalizada"], label_visibility="collapsed", key="tom_msg")
        msg_base = f"Celebrando os aniversariantes de {nome_meses_tela}! Muita saúde e vida ativa para todos!"
        mensagem_digitada = st.text_area("Ajuste o texto:", value=msg_base, height=70)

    with c_botoes:
        if not df_mes.empty:
            if st.button("📕 GERAR PDF", use_container_width=True, type="primary"):
                st.session_state.pdf_niver = gerar_cartaz_pdf_core(df_mes, titulo_doc, subtitulo_doc, mensagem_digitada)

            if "pdf_niver" in st.session_state:
                st.download_button("📥 BAIXAR PDF", st.session_state.pdf_niver, f"Cartaz_{nome_arq}.pdf", "application/pdf", use_container_width=True)

    st.markdown(
        """<style>
        .zoom-niver { width: 50px; height: 50px; border-radius: 50%; object-fit: cover; border: 2px solid #1E88E5; transition: transform 0.3s ease; cursor: zoom-in; }
        .zoom-niver:hover { transform: scale(3.5); z-index: 999; position: relative; }
        .badge-hoje { background: #10B981; color: white; padding: 4px 10px; border-radius: 12px; font-size: 11px; font-weight: 800; }
        .badge-passou { background: #F1F5F9; color: #64748B; padding: 4px 10px; border-radius: 12px; font-size: 11px; font-weight: 800; border: 1px solid #E2E8F0; }
        .badge-chegando { background: #FEF3C7; color: #D97706; padding: 4px 10px; border-radius: 12px; font-size: 11px; font-weight: 800; border: 1px solid #FDE68A; }
        .badge-parab { background: #DCFCE7; color: #16A34A; padding: 4px 10px; border-radius: 12px; font-size: 11px; font-weight: 800; border: 1px solid #BBF7D0; }
    </style>""",
        unsafe_allow_html=True,
    )

    # ── LISTA DE ANIVERSARIANTES ─────────────────────────────────────────────
    cfg_niver = get_config_niver()
    parab_dict = get_parabenizados_dict()

    for _, r in df_mes.iterrows():
        _mes_r, _dia_r = int(r["mes"]), int(r["dia"])
        try:
            aniv_data = datetime.date(hoje.year, _mes_r, _dia_r)
        except ValueError:
            # 29/02 em ano não bissexto → considera 28/02 para o cálculo.
            aniv_data = datetime.date(hoje.year, _mes_r, 28)
        delta = (aniv_data - hoje).days
        aluno_id = str(r.get("id", ""))
        ja_parab = aluno_id in parab_dict

        with st.container(border=True):
            c_av, c_info, c_status, c_whats, c_parab, c_ficha = st.columns(
                [1, 3, 2, 0.7, 1.4, 0.7], vertical_alignment="center"
            )
            with c_av:
                if pd.notna(r.get("url_foto")) and str(r.get("url_foto")).strip() != "":
                    st.markdown(f'<img src="{r["url_foto"]}" class="zoom-niver">', unsafe_allow_html=True)
                else:
                    st.markdown("👤", unsafe_allow_html=True)

            with c_info:
                st.markdown(f"**{r['nome'].upper()}**")
                st.markdown(f"<span style='font-size:12px; color:#64748B;'>🎂 {r['dia']:02d}/{r['mes']:02d}</span>", unsafe_allow_html=True)

            with c_status:
                if delta == 0:
                    st.markdown('<span class="badge-hoje">🎈 É HOJE!</span>', unsafe_allow_html=True)
                elif delta > 0:
                    st.markdown(f'<span class="badge-chegando">⏳ Faltam {delta}d</span>', unsafe_allow_html=True)
                else:
                    st.markdown('<span class="badge-passou">✔️ Passou</span>', unsafe_allow_html=True)
                _turma = str(r.get("turma") or "").strip()[:10]
                if _turma:
                    st.markdown(f"<span style='font-size:12px;color:#475569;font-weight:600;'>📍 {_turma}</span>", unsafe_allow_html=True)

            with c_whats:
                if ja_parab:
                    # Já parabenizado — ícone apagado, sem link
                    st.markdown(
                        '<span title="Já parabenizado" '
                        'style="font-size:20px;opacity:0.22;cursor:default;">💬</span>',
                        unsafe_allow_html=True,
                    )
                else:
                    # Texto idêntico ao painel admin conforme status do aniversário:
                    # Dia Exato (hoje) ou Atrasado (passou). Futuros não recebem link.
                    _status_w = status_niver_por_delta(delta)
                    if _status_w is not None:
                        msg_pessoal = montar_mensagem_niver(_status_w, str(r.get("nome", "")))
                        link_w = montar_link_whatsapp(str(r.get("whatsapp", "") or ""), msg_pessoal)
                        if link_w:
                            st.markdown(
                                f'<a href="{link_w}" target="_blank" title="Enviar parabéns via WhatsApp" '
                                f'style="font-size:20px;text-decoration:none;">💬</a>',
                                unsafe_allow_html=True,
                            )

            with c_parab:
                if ja_parab:
                    ts_raw = parab_dict.get(aluno_id, "")
                    ts_parte = ts_raw.split("|")[0][:16].replace("T", " ") if ts_raw else ""
                    _ts_html = (
                        f'<br><span style="font-size:10px;color:#64748B;">{ts_parte}</span>'
                        if ts_parte else ""
                    )
                    st.markdown(
                        f'<div style="background:#D1FAE5;border:1px solid #6EE7B7;border-radius:8px;'
                        f'padding:4px 8px;text-align:center;">'
                        f'<span style="font-size:12px;font-weight:800;color:#065F46;">✅ Parabenizado</span>'
                        f'{_ts_html}</div>',
                        unsafe_allow_html=True,
                    )
                    if st.button("↩️", key=f"desparab_{aluno_id}",
                                 help="Desfazer marcação de parabenizado"):
                        desmarcar_parabenizado(aluno_id)
                        st.rerun()
                else:
                    if st.button(
                        "🎉 Parabenizei",
                        key=f"parab_{aluno_id}",
                        use_container_width=True,
                        help="Clique após enviar os parabéns — registra e desativa o link do WhatsApp",
                    ):
                        nome_op = st.session_state.get("usuario_nome", "")
                        marcar_parabenizado(aluno_id, nome_op)
                        st.rerun()

            with c_ficha:
                if st.button("🩺", key=f"n_{r['id']}"):
                    st.session_state.aluno_prontuario = r.to_dict()
                    st.session_state.menu_atual = "Portal do Aluno"
                    st.rerun()
