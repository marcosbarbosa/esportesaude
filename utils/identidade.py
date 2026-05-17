# ==============================================================================
# 📄 utils/identidade.py
# ⚙️ Módulo de Identidade Visual — cabeçalho e rodapé partilhados por todos
#    os relatórios do sistema (HTML, Word, Excel).
# ==============================================================================

import json
import os
import base64

_CONFIG_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "identidade.json")

_DEFAULTS = {
    "titulo_projeto":    "ESPORTE E SAÚDE NA COMUNIDADE - FASE 2",
    "subtitulo_projeto": "Projeto de Atividade Física, Saúde e Bem-Estar",
    "nome_organizacao":  "Instituto Muda Brasil",
    "cnpj":              "08.817.519/0001-79",
    "endereco":          "R. Sapoti, 20 - Campo Belo, São Paulo - SP, 04615-040",
    "telefone":          "",
    "email_contato":     "",
    "site":              "imbra.org.br",
    "instagram":         "@institutomudabrasil",
    "logo_principal":    "logo-imbra.png",
    "logo_secundaria":   "logo-secretaria.png",
}


# ── Leitura / escrita ──────────────────────────────────────────────────────────

def get_config() -> dict:
    """Retorna configuração actual (lê do ficheiro; usa defaults se falhar)."""
    try:
        with open(_CONFIG_PATH, encoding="utf-8") as f:
            data = json.load(f)
        # garante campos ausentes com defaults
        for k, v in _DEFAULTS.items():
            data.setdefault(k, v)
        return data
    except Exception:
        return dict(_DEFAULTS)


def salvar_config(data: dict) -> None:
    """Persiste configuração no ficheiro JSON."""
    merged = dict(_DEFAULTS)
    merged.update(data)
    with open(_CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(merged, f, ensure_ascii=False, indent=2)


# ── Logos em base64 ────────────────────────────────────────────────────────────

def get_logo_b64(nome_ficheiro: str) -> str:
    """Devolve string base64 da imagem ou '' se não existir."""
    if not nome_ficheiro:
        return ""
    path = nome_ficheiro if os.path.isabs(nome_ficheiro) else os.path.join(
        os.path.dirname(os.path.dirname(__file__)), nome_ficheiro
    )
    try:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except Exception:
        return ""


def get_logo_data_url(nome_ficheiro: str) -> str:
    """Devolve data:image/png;base64,... para uso directo em HTML src=."""
    b64 = get_logo_b64(nome_ficheiro)
    return f"data:image/png;base64,{b64}" if b64 else ""


def _img_tag(b64: str, alt: str, width: int = 120) -> str:
    if not b64:
        return f'<span style="color:#94A3B8;font-size:8pt;border:1px dashed #94A3B8;padding:8px;">{alt}</span>'
    ext = "png"
    return f'<img src="data:image/{ext};base64,{b64}" style="max-width:{width}px;height:auto;max-height:80px;" alt="{alt}">'


# ── Renders HTML (usados em PDF/Word/Streamlit) ────────────────────────────────

def render_cabecalho_html(cfg: dict | None = None, extra: str = "") -> str:
    """
    Devolve bloco HTML completo do cabeçalho do relatório.
    extra: texto adicional (ex: 'Período: Jan/2026')
    """
    if cfg is None:
        cfg = get_config()

    logo_p = get_logo_b64(cfg.get("logo_principal", ""))
    logo_s = get_logo_b64(cfg.get("logo_secundaria", ""))

    html_logo_p = _img_tag(logo_p, cfg.get("nome_organizacao", ""), 120)
    html_logo_s = _img_tag(logo_s, "Parceiro Institucional", 140)

    extra_html = f'<p style="font-size:9pt;color:#64748B;margin-top:5px;">{extra}</p>' if extra else ""

    return f"""
<table style="width:100%;border-bottom:3px solid #0056b3;margin-bottom:20px;border-collapse:collapse;">
  <tr>
    <td style="width:22%;text-align:left;vertical-align:middle;padding:8px 4px;">{html_logo_s}</td>
    <td style="width:56%;text-align:center;vertical-align:middle;padding:8px 4px;">
      <p style="margin:0;font-size:10pt;color:#0A2540;text-transform:uppercase;font-weight:900;letter-spacing:-0.3px;line-height:1.2;">
        {cfg.get("titulo_projeto","")}
      </p>
      <p style="margin:2px 0 0 0;font-size:9pt;color:#475569;font-weight:600;line-height:1.2;">
        {cfg.get("subtitulo_projeto","")}
      </p>
      {extra_html}
    </td>
    <td style="width:22%;text-align:right;vertical-align:middle;padding:8px 4px;">{html_logo_p}</td>
  </tr>
</table>
""".strip()


def render_rodape_html(cfg: dict | None = None) -> str:
    """Devolve bloco HTML do rodapé do relatório."""
    if cfg is None:
        cfg = get_config()

    partes = [cfg.get("nome_organizacao", "")]
    if cfg.get("cnpj"):
        partes.append(f"CNPJ: {cfg['cnpj']}")
    if cfg.get("site"):
        partes.append(f"🌐 {cfg['site']}")
    if cfg.get("instagram"):
        partes.append(f"📸 {cfg['instagram']}")
    if cfg.get("telefone"):
        partes.append(f"📞 {cfg['telefone']}")
    if cfg.get("endereco"):
        partes.append(cfg["endereco"])

    linha = " &nbsp;|&nbsp; ".join(partes)
    return f"""
<div style="margin-top:40px;text-align:center;font-size:8pt;color:#94A3B8;
            border-top:1px solid #E2E8F0;padding-top:10px;">
  {linha}
</div>
""".strip()


# ── Render Word (python-docx) ──────────────────────────────────────────────────

def render_cabecalho_docx(doc, cfg: dict | None = None) -> None:
    """Adiciona cabeçalho padrão (logo + título) a um documento python-docx."""
    try:
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
    except ImportError:
        return

    if cfg is None:
        cfg = get_config()

    logo_path = cfg.get("logo_principal", "logo-imbra.png")
    if not os.path.isabs(logo_path):
        logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), logo_path)

    if os.path.exists(logo_path):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(1.5))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cfg.get("titulo_projeto", "") + "\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 86, 179)

    if cfg.get("subtitulo_projeto"):
        run2 = p.add_run(cfg["subtitulo_projeto"])
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(100, 116, 139)


def render_rodape_docx(doc, cfg: dict | None = None) -> None:
    """Adiciona rodapé padrão ao documento python-docx."""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
    except ImportError:
        return

    if cfg is None:
        cfg = get_config()

    partes = [cfg.get("nome_organizacao", "")]
    if cfg.get("cnpj"):
        partes.append(f"CNPJ: {cfg['cnpj']}")
    if cfg.get("site"):
        partes.append(cfg["site"])
    if cfg.get("instagram"):
        partes.append(cfg["instagram"])
    if cfg.get("telefone"):
        partes.append(cfg["telefone"])

    doc.add_paragraph()
    p = doc.add_paragraph(" | ".join(partes))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(148, 163, 184)


# ── Render Excel (xlsxwriter) ─────────────────────────────────────────────────

def render_cabecalho_excel(worksheet, workbook, cfg: dict | None = None) -> None:
    """Insere logos e título nas primeiras linhas da planilha Excel."""
    if cfg is None:
        cfg = get_config()

    logo_muda = cfg.get("logo_principal", "logo-imbra.png")
    logo_sec  = cfg.get("logo_secundaria", "logo-secretaria.png")
    if not os.path.isabs(logo_muda):
        base = os.path.dirname(os.path.dirname(__file__))
        logo_muda = os.path.join(base, logo_muda)
        logo_sec  = os.path.join(base, cfg.get("logo_secundaria", "logo-secretaria.png"))

    try:
        if os.path.exists(logo_muda):
            worksheet.insert_image("A1", logo_muda, {"x_scale": 0.16, "y_scale": 0.16})
    except Exception:
        pass
    try:
        if os.path.exists(logo_sec):
            worksheet.insert_image("C1", logo_sec, {"x_scale": 0.35, "y_scale": 0.35})
    except Exception:
        pass

    f_tit = workbook.add_format({"bold": True, "font_size": 13, "align": "left", "valign": "vcenter"})
    f_sub = workbook.add_format({"font_size": 10, "align": "left", "valign": "vcenter", "font_color": "#555555"})
    worksheet.write("E2", cfg.get("titulo_projeto", ""), f_tit)
    worksheet.write("E3", cfg.get("subtitulo_projeto", ""), f_sub)


def render_rodape_excel(worksheet, workbook, cfg: dict | None = None,
                        linha_rodape: int = 20) -> None:
    """Insere rodapé institucional no Excel a partir de linha_rodape."""
    if cfg is None:
        cfg = get_config()

    f_bold = workbook.add_format({"bold": True, "font_size": 8, "font_color": "#475569"})
    f_norm = workbook.add_format({"font_size": 8, "font_color": "#94A3B8"})

    worksheet.write(linha_rodape,     0, cfg.get("nome_organizacao", ""), f_bold)
    worksheet.write(linha_rodape + 1, 0, f"CNPJ: {cfg.get('cnpj','')}", f_norm)
    if cfg.get("site"):
        worksheet.write(linha_rodape + 2, 0, cfg["site"], f_norm)
    if cfg.get("instagram"):
        worksheet.write(linha_rodape + 3, 0, cfg["instagram"], f_norm)
    if cfg.get("endereco"):
        worksheet.write(linha_rodape + 4, 0, cfg["endereco"], f_norm)
    if cfg.get("telefone"):
        worksheet.write(linha_rodape + 5, 0, f"Tel: {cfg['telefone']}", f_norm)
